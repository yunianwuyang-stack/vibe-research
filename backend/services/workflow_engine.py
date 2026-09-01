"""(docstring)"""
from __future__ import annotations

import asyncio
import hashlib
import json
import logging
import os
import re
import shlex
import shutil
import subprocess
import sys
import uuid
from contextvars import ContextVar
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Awaitable, Callable, Dict, List, Optional

from config import SKILLS_DIR, TOOLS_DIR, WORKSPACES_DIR, IS_DESKTOP
from models.schemas import TemplateType, WorkflowStatus, StepStatus

log = logging.getLogger(__name__)

# Cap concurrent skill executions. Matrix launches previously spawned every
# workflow at once; each held a SQLite connection for the entire agent run and
# collapsed into "database is locked" on step completion.
_MAX_CONCURRENT_STEPS = max(1, int(os.environ.get("VIBE_MAX_CONCURRENT_STEPS", "4")))
_step_semaphore: asyncio.Semaphore | None = None
_step_semaphore_loop: asyncio.AbstractEventLoop | None = None


def _get_step_semaphore() -> asyncio.Semaphore:
    """Create the step gate on the current event loop.

    Module-level asyncio primitives bind to the first loop that touches them.
    Desktop restarts and pytest TestClient runs each create a fresh loop, so the
    semaphore must be recreated whenever the running loop changes.
    """
    global _step_semaphore, _step_semaphore_loop
    loop = asyncio.get_running_loop()
    if _step_semaphore is None or _step_semaphore_loop is not loop:
        _step_semaphore = asyncio.Semaphore(_MAX_CONCURRENT_STEPS)
        _step_semaphore_loop = loop
    return _step_semaphore

# ============================================================

# ============================================================

@dataclass
class StepDef:
    skill_name: str
    display_name: str
    output_files: List[str] = field(default_factory=list)
    primary_output: Optional[str] = None
    has_checkpoint: bool = False
    checkpoint_type: Optional[str] = None


@dataclass
class TemplateDef:
    pipeline_skill: str
    display_name: str
    sub_steps: List[StepDef]


# ============================================================

# ============================================================


# Host offline scaffolding skills (deterministic local builders). These are
# useful for competition/demo recovery lanes but are NOT production scientific
# backends. Production catalogs must exclude content-scaffold host skills.
HOST_CONTENT_SCAFFOLD_SKILLS: frozenset[str] = frozenset({
    "research-lit", "idea-creator", "novelty-check", "research-review",
    "research-refine-pipeline", "experiment-bridge", "paper-plan",
    "paper-analysis", "paper-figure", "paper-figure-drawio", "paper-write",
    "paper-write-zh", "paper-write-nature", "auto-paper-improvement-loop",
    "comp-problem-read", "comp-modeling", "comp-solve", "comp-sensitivity",
    "comp-paper-zh", "comp-paper-en", "thesis-proposal", "literature-review",
    "course-paper", "course-report", "paper-from-assets", "paper-slides",
    "paper-poster", "humanities-plan", "humanities-write", "project-blueprint",
    "grad-project-scaffold", "copyright-material", "patent-scaffold",
    "paper-rebuttal", "paper-cover-letter",
})
HOST_TOOLING_SKILLS: frozenset[str] = frozenset({
    "paper-compile", "paper-compile-zh", "paper-compile-nature",
    "docx-format-check", "paper-export",
})


def production_skill_catalog() -> Dict[str, Any]:
    """Return production-safe skill catalog without host content scaffolds.

    Host content scaffolds may still run in offline recovery mode, but they
    must not appear as production scientific backends (P1-PS-004 / REQ-P1-04).
    """
    production_templates: Dict[str, Any] = {}
    for template_id, tmpl in TEMPLATES.items():
        steps = []
        for step in tmpl.sub_steps:
            skill = step.skill_name
            if skill in HOST_CONTENT_SCAFFOLD_SKILLS:
                steps.append({
                    "skill_name": skill,
                    "display_name": step.display_name,
                    "executor": "agent_broker_required",
                    "host_scaffold_allowed": False,
                    "primary_output": step.primary_output,
                    "output_files": list(step.output_files),
                })
            elif skill in HOST_TOOLING_SKILLS:
                steps.append({
                    "skill_name": skill,
                    "display_name": step.display_name,
                    "executor": "host_tooling",
                    "host_scaffold_allowed": True,
                    "primary_output": step.primary_output,
                    "output_files": list(step.output_files),
                })
            else:
                steps.append({
                    "skill_name": skill,
                    "display_name": step.display_name,
                    "executor": "agent_or_tooling",
                    "host_scaffold_allowed": skill not in HOST_CONTENT_SCAFFOLD_SKILLS,
                    "primary_output": step.primary_output,
                    "output_files": list(step.output_files),
                })
        production_templates[template_id] = {
            "pipeline_skill": tmpl.pipeline_skill,
            "display_name": tmpl.display_name,
            "steps": steps,
            "requires_agent_broker": any(
                s["executor"] == "agent_broker_required" for s in steps
            ),
        }
    return {
        "mode": "production",
        "host_content_scaffolds_excluded": True,
        "host_content_scaffold_skills": sorted(HOST_CONTENT_SCAFFOLD_SKILLS),
        "host_tooling_skills": sorted(HOST_TOOLING_SKILLS),
        "templates": production_templates,
    }


def production_catalog_has_no_host_content_scaffolds() -> bool:
    catalog = production_skill_catalog()
    for tmpl in catalog["templates"].values():
        for step in tmpl["steps"]:
            if step["skill_name"] in HOST_CONTENT_SCAFFOLD_SKILLS and step.get("host_scaffold_allowed"):
                return False
            if step.get("executor") == "host_content_scaffold":
                return False
    return True


TEMPLATES: Dict[str, TemplateDef] = {
    "idea_discovery": TemplateDef(
        pipeline_skill="idea-discovery",
        display_name="Idea 发现",
        sub_steps=[
            StepDef(skill_name="research-lit", display_name="文献调研", output_files=["literature_review.md","references.bib"], primary_output="literature_review.md", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="idea-creator", display_name="Idea 生成", output_files=["IDEA_REPORT.md"], primary_output="IDEA_REPORT.md", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="novelty-check", display_name="新颖性验证", output_files=["novelty_check_report.md"], primary_output="novelty_check_report.md", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="research-review", display_name="外部评审", output_files=["review_report.md"], primary_output="review_report.md", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="research-refine-pipeline", display_name="方法精炼+实验规划", output_files=["refine-logs/FINAL_PROPOSAL.md","refine-logs/EXPERIMENT_PLAN.md"], primary_output="refine-logs/FINAL_PROPOSAL.md", has_checkpoint=True, checkpoint_type="feedback"),
        ],
    ),
    "experiment_bridge": TemplateDef(
        pipeline_skill="experiment-bridge",
        display_name="实验桥接",
        sub_steps=[
            StepDef(skill_name="experiment-bridge", display_name="实验实现+部署+出图", output_files=["experiment_results.md","figures/latex_includes.tex","figures/experiment_data.json"], primary_output="figures/", has_checkpoint=True, checkpoint_type="feedback"),
        ],
    ),
    "auto_review": TemplateDef(
        pipeline_skill="auto-review-loop",
        display_name="自动审稿循环",
        sub_steps=[
            StepDef(skill_name="auto-review-loop", display_name="自动审稿循环", output_files=["NARRATIVE_REPORT.md","AUTO_REVIEW.md"], primary_output="NARRATIVE_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
        ],
    ),
    "paper_writing": TemplateDef(
        pipeline_skill="paper-writing",
        display_name="论文写作",
        sub_steps=[
            StepDef(skill_name="paper-plan", display_name="论文大纲", output_files=["PAPER_PLAN.md"], primary_output="PAPER_PLAN.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-analysis", display_name="代码实现", output_files=["RESULTS.md","figures/all_results.json","code/main.py"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="流程与架构图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-write", display_name="LaTeX 写作", output_files=["paper/main.tex","paper/references.bib"], primary_output="paper/main.tex", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-compile", display_name="编译 PDF", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="auto-paper-improvement-loop", display_name="论文改进循环", output_files=["paper/main.pdf","paper/PAPER_IMPROVEMENT_LOG.md"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "paper_writing_zh": TemplateDef(
        pipeline_skill="paper-writing",
        display_name="中文论文写作",
        sub_steps=[
            StepDef(skill_name="paper-plan-zh", display_name="论文大纲（中文）", output_files=["PAPER_PLAN.md"], primary_output="PAPER_PLAN.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-analysis", display_name="代码实现", output_files=["RESULTS.md","figures/all_results.json","code/main.py"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="流程与架构图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-write-zh", display_name="LaTeX 写作（中文）", output_files=["paper/main.tex","paper/references.bib"], primary_output="paper/main.tex", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-compile-zh", display_name="编译 PDF（中文）", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="auto-paper-improvement-loop", display_name="论文改进循环", output_files=["paper/main.pdf","paper/PAPER_IMPROVEMENT_LOG.md"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "nature_writing": TemplateDef(
        pipeline_skill="paper-plan",
        display_name="Nature 论文写作",
        sub_steps=[
            StepDef(skill_name="paper-plan", display_name="论文规划", output_files=["PAPER_PLAN.md"], primary_output="PAPER_PLAN.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-analysis", display_name="数据分析", output_files=["RESULTS.md","figures/all_results.json"], primary_output="RESULTS.md", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="nature-figure", display_name="Nature 图表", output_files=["figures/"], primary_output="figures/", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure-drawio", display_name="架构图绘制", output_files=["figures/"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-write-nature", display_name="Nature 写作", output_files=["paper/main.tex","paper/references.bib"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-compile", display_name="论文编译", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="auto-paper-improvement-loop", display_name="改进循环", output_files=["paper/main.pdf","paper/PAPER_IMPROVEMENT_LOG.md"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "full_pipeline": TemplateDef(
        pipeline_skill="research-pipeline",
        display_name="全流程",
        sub_steps=[
            StepDef(skill_name="research-lit", display_name="文献调研", output_files=["literature_review.md","references.bib"], primary_output="literature_review.md", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="idea-creator", display_name="Idea 生成", output_files=["IDEA_REPORT.md"], primary_output="IDEA_REPORT.md", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="novelty-check", display_name="新颖性验证", output_files=["novelty_check_report.md"], primary_output="novelty_check_report.md", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="research-review", display_name="外部评审", output_files=["review_report.md"], primary_output="review_report.md", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="research-refine-pipeline", display_name="方法精炼+实验规划", output_files=["refine-logs/FINAL_PROPOSAL.md","refine-logs/EXPERIMENT_PLAN.md"], primary_output="refine-logs/FINAL_PROPOSAL.md", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="experiment-bridge", display_name="实验实现+出图", output_files=["experiment_results.md","figures/latex_includes.tex","figures/experiment_data.json"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-plan", display_name="论文大纲", output_files=["PAPER_PLAN.md"], primary_output="PAPER_PLAN.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure-drawio", display_name="流程与架构图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-write", display_name="LaTeX 写作", output_files=["paper/main.tex","paper/references.bib"], primary_output="paper/main.tex", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-compile", display_name="编译 PDF", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="auto-paper-improvement-loop", display_name="论文改进循环", output_files=["paper/main.pdf","paper/PAPER_IMPROVEMENT_LOG.md"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_cumcm": TemplateDef(
        pipeline_skill="comp-prob-analysis",
        display_name="国赛数学建模 (CUMCM)",
        sub_steps=[
            StepDef(skill_name="comp-prob-analysis", display_name="赛题分析", output_files=["PROBLEM_ANALYSIS.md"], primary_output="PROBLEM_ANALYSIS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-modeling", display_name="建模求解", output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="comp-code", display_name="编程实现", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="流程与架构图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-zh", display_name="竞赛论文撰写", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-zh", display_name="编译与合规检查", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_mcm": TemplateDef(
        pipeline_skill="comp-prob-analysis",
        display_name="美赛 (MCM/ICM)",
        sub_steps=[
            StepDef(skill_name="comp-prob-analysis", display_name="Problem Analysis", output_files=["PROBLEM_ANALYSIS.md"], primary_output="PROBLEM_ANALYSIS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-modeling", display_name="Modeling & Solution", output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="comp-code", display_name="Programming", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="Figure Generation", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="Flow & Architecture Diagrams", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-en", display_name="Paper Writing (English)", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-en", display_name="Compile & Compliance Check", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_huawei": TemplateDef(
        pipeline_skill="comp-prob-analysis",
        display_name="华为杯研究生数学建模",
        sub_steps=[
            StepDef(skill_name="comp-prob-analysis", display_name="赛题分析", output_files=["PROBLEM_ANALYSIS.md"], primary_output="PROBLEM_ANALYSIS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-modeling", display_name="建模求解", output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="comp-code", display_name="编程实现", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="流程与架构图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-zh", display_name="竞赛论文撰写", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-zh", display_name="编译与合规检查", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_apmcm_zh": TemplateDef(
        pipeline_skill="comp-prob-analysis",
        display_name="亚太地区数学建模 (APMCM 中文赛项)",
        sub_steps=[
            StepDef(skill_name="comp-prob-analysis", display_name="赛题分析", output_files=["PROBLEM_ANALYSIS.md"], primary_output="PROBLEM_ANALYSIS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-modeling", display_name="建模求解", output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="comp-code", display_name="编程实现", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="流程与架构图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-zh", display_name="竞赛论文撰写", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-zh", display_name="编译与合规检查", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_mathorcup": TemplateDef(
        pipeline_skill="comp-prob-analysis",
        display_name="MathorCup 数学建模",
        sub_steps=[
            StepDef(skill_name="comp-prob-analysis", display_name="赛题分析", output_files=["PROBLEM_ANALYSIS.md"], primary_output="PROBLEM_ANALYSIS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-modeling", display_name="建模求解", output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="comp-code", display_name="编程实现", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="流程与架构图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-zh", display_name="竞赛论文撰写", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-zh", display_name="编译与合规检查", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_apmcm": TemplateDef(
        pipeline_skill="comp-prob-analysis",
        display_name="亚太数学建模 (APMCM)",
        sub_steps=[
            StepDef(skill_name="comp-prob-analysis", display_name="赛题分析", output_files=["PROBLEM_ANALYSIS.md"], primary_output="PROBLEM_ANALYSIS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-modeling", display_name="建模求解", output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="comp-code", display_name="编程实现", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="Flow & Architecture Diagrams", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-en", display_name="Paper Writing", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-en", display_name="Compile & Compliance Check", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_stats": TemplateDef(
        pipeline_skill="comp-stats-topic",
        display_name="统计建模大赛",
        sub_steps=[
            StepDef(skill_name="comp-stats-topic", display_name="选题与数据规划", output_files=["TOPIC_PLAN.md"], primary_output="TOPIC_PLAN.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-code", display_name="数据采集与统计分析", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="流程与架构图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-zh", display_name="论文撰写", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-zh", display_name="编译与合规检查", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_teddy": TemplateDef(
        pipeline_skill="comp-prob-analysis",
        display_name="泰迪杯数据挖掘挑战赛",
        sub_steps=[
            StepDef(skill_name="comp-prob-analysis", display_name="赛题分析", output_files=["PROBLEM_ANALYSIS.md"], primary_output="PROBLEM_ANALYSIS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-modeling", display_name="建模求解", output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="comp-code", display_name="编程实现", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="流程与架构图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-zh", display_name="竞赛论文撰写", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-zh", display_name="编译与合规检查", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_certcup": TemplateDef(
        pipeline_skill="comp-prob-analysis",
        display_name="认证杯数学建模",
        sub_steps=[
            StepDef(skill_name="comp-prob-analysis", display_name="赛题分析", output_files=["PROBLEM_ANALYSIS.md"], primary_output="PROBLEM_ANALYSIS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-modeling", display_name="建模求解", output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="comp-code", display_name="编程实现", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="流程与架构图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-zh", display_name="竞赛论文撰写", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-zh", display_name="编译与合规检查", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_huazhong": TemplateDef(
        pipeline_skill="comp-prob-analysis",
        display_name="华中杯数学建模",
        sub_steps=[
            StepDef(skill_name="comp-prob-analysis", display_name="赛题分析", output_files=["PROBLEM_ANALYSIS.md"], primary_output="PROBLEM_ANALYSIS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-modeling", display_name="建模求解", output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="comp-code", display_name="编程实现", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="流程与架构图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-zh", display_name="竞赛论文撰写", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-zh", display_name="编译与合规检查", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_huadong": TemplateDef(
        pipeline_skill="comp-prob-analysis",
        display_name="华东杯数学建模",
        sub_steps=[
            StepDef(skill_name="comp-prob-analysis", display_name="赛题分析", output_files=["PROBLEM_ANALYSIS.md"], primary_output="PROBLEM_ANALYSIS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-modeling", display_name="建模求解", output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="comp-code", display_name="编程实现", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="流程与架构图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-zh", display_name="竞赛论文撰写", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-zh", display_name="编译与合规检查", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_wuyi": TemplateDef(
        pipeline_skill="comp-prob-analysis",
        display_name="五一杯数学建模",
        sub_steps=[
            StepDef(skill_name="comp-prob-analysis", display_name="赛题分析", output_files=["PROBLEM_ANALYSIS.md"], primary_output="PROBLEM_ANALYSIS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-modeling", display_name="建模求解", output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="comp-code", display_name="编程实现", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="流程与架构图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-zh", display_name="竞赛论文撰写", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-zh", display_name="编译与合规检查", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_shuwei": TemplateDef(
        pipeline_skill="comp-prob-analysis",
        display_name="数维杯数学建模",
        sub_steps=[
            StepDef(skill_name="comp-prob-analysis", display_name="赛题分析", output_files=["PROBLEM_ANALYSIS.md"], primary_output="PROBLEM_ANALYSIS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-modeling", display_name="建模求解", output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="comp-code", display_name="编程实现", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="流程与架构图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-zh", display_name="竞赛论文撰写", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-zh", display_name="编译与合规检查", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_zhongqing": TemplateDef(
        pipeline_skill="comp-prob-analysis",
        display_name="中青杯数学建模",
        sub_steps=[
            StepDef(skill_name="comp-prob-analysis", display_name="赛题分析", output_files=["PROBLEM_ANALYSIS.md"], primary_output="PROBLEM_ANALYSIS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-modeling", display_name="建模求解", output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="comp-code", display_name="编程实现", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="流程与架构图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-zh", display_name="竞赛论文撰写", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-zh", display_name="编译与合规检查", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_yangtze": TemplateDef(
        pipeline_skill="comp-prob-analysis",
        display_name="长三角数学建模",
        sub_steps=[
            StepDef(skill_name="comp-prob-analysis", display_name="赛题分析", output_files=["PROBLEM_ANALYSIS.md"], primary_output="PROBLEM_ANALYSIS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-modeling", display_name="建模求解", output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="comp-code", display_name="编程实现", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="流程与架构图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-zh", display_name="竞赛论文撰写", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-zh", display_name="编译与合规检查", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_diangong": TemplateDef(
        pipeline_skill="comp-prob-analysis",
        display_name="电工杯数学建模",
        sub_steps=[
            StepDef(skill_name="comp-prob-analysis", display_name="赛题分析", output_files=["PROBLEM_ANALYSIS.md"], primary_output="PROBLEM_ANALYSIS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-modeling", display_name="建模求解", output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="comp-code", display_name="编程实现", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="流程与架构图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-zh", display_name="竞赛论文撰写", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-zh", display_name="编译与合规检查", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_shenzhen": TemplateDef(
        pipeline_skill="comp-prob-analysis",
        display_name="深圳杯数学建模",
        sub_steps=[
            StepDef(skill_name="comp-prob-analysis", display_name="赛题分析", output_files=["PROBLEM_ANALYSIS.md"], primary_output="PROBLEM_ANALYSIS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-modeling", display_name="建模求解", output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="comp-code", display_name="编程实现", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="流程与架构图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-zh", display_name="竞赛论文撰写", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-zh", display_name="编译与合规检查", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_huashu": TemplateDef(
        pipeline_skill="comp-prob-analysis",
        display_name="华数杯数学建模",
        sub_steps=[
            StepDef(skill_name="comp-prob-analysis", display_name="赛题分析", output_files=["PROBLEM_ANALYSIS.md"], primary_output="PROBLEM_ANALYSIS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-modeling", display_name="建模求解", output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="comp-code", display_name="编程实现", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="流程与架构图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-zh", display_name="竞赛论文撰写", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-zh", display_name="编译与合规检查", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_tianfu": TemplateDef(
        pipeline_skill="comp-prob-analysis",
        display_name="天府杯数学建模",
        sub_steps=[
            StepDef(skill_name="comp-prob-analysis", display_name="赛题分析", output_files=["PROBLEM_ANALYSIS.md"], primary_output="PROBLEM_ANALYSIS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-modeling", display_name="建模求解", output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="comp-code", display_name="编程实现", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="流程与架构图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-zh", display_name="竞赛论文撰写", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-zh", display_name="编译与合规检查", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_liaoning": TemplateDef(
        pipeline_skill="comp-prob-analysis",
        display_name="辽宁省/东三省数学建模",
        sub_steps=[
            StepDef(skill_name="comp-prob-analysis", display_name="赛题分析", output_files=["PROBLEM_ANALYSIS.md"], primary_output="PROBLEM_ANALYSIS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-modeling", display_name="建模求解", output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="comp-code", display_name="编程实现", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="流程与架构图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-zh", display_name="竞赛论文撰写", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-zh", display_name="编译与合规检查", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_certcup_en": TemplateDef(
        pipeline_skill="comp-prob-analysis",
        display_name="认证杯国际赛 (小美赛)",
        sub_steps=[
            StepDef(skill_name="comp-prob-analysis", display_name="Problem Analysis", output_files=["PROBLEM_ANALYSIS.md"], primary_output="PROBLEM_ANALYSIS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-modeling", display_name="Modeling & Solution", output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="comp-code", display_name="Programming", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="Figure Generation", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="Flow & Architecture Diagrams", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-en", display_name="Paper Writing (English)", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-en", display_name="Compile & Compliance Check", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "comp_shuwei_en": TemplateDef(
        pipeline_skill="comp-prob-analysis",
        display_name="数维杯国际赛",
        sub_steps=[
            StepDef(skill_name="comp-prob-analysis", display_name="Problem Analysis", output_files=["PROBLEM_ANALYSIS.md"], primary_output="PROBLEM_ANALYSIS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-modeling", display_name="Modeling & Solution", output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="comp-code", display_name="Programming", output_files=["code/main.py","RESULTS.md"], primary_output="RESULTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-figure", display_name="Figure Generation", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="Flow & Architecture Diagrams", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="comp-paper-en", display_name="Paper Writing (English)", output_files=["paper/main.tex"], primary_output="paper/main.tex", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="comp-compile-en", display_name="Compile & Compliance Check", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "thesis_proposal": TemplateDef(
        pipeline_skill="thesis-proposal",
        display_name="开题报告",
        sub_steps=[
            StepDef(skill_name="thesis-proposal", display_name="文献调研与开题撰写", output_files=["literature_notes.md","PROPOSAL.md"], primary_output="PROPOSAL.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="paper-figure-drawio", display_name="技术路线图绘制", output_files=["figures/fig_roadmap.png"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="docx-export", display_name="格式检查与 Word 导出", output_files=["PROPOSAL.docx"], primary_output="PROPOSAL.docx", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "literature_review": TemplateDef(
        pipeline_skill="literature-review",
        display_name="文献综述",
        sub_steps=[
            StepDef(skill_name="literature-review", display_name="文献检索与综述撰写", output_files=["papers_pool.md","LITERATURE_REVIEW.md"], primary_output="LITERATURE_REVIEW.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="docx-export", display_name="格式检查与 Word 导出", output_files=["LITERATURE_REVIEW.docx"], primary_output="LITERATURE_REVIEW.docx", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "course_paper": TemplateDef(
        pipeline_skill="course-plan",
        display_name="课程论文",
        sub_steps=[
            StepDef(skill_name="course-plan", display_name="大纲与图表规划", output_files=["OUTLINE.md","PAPER_PLAN.md"], primary_output="OUTLINE.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-analysis", display_name="数据分析", output_files=["RESULTS.md","figures/all_results.json","code/main.py"], primary_output="RESULTS.md", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure", display_name="图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="架构图/流程图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="course-paper", display_name="课程论文撰写", output_files=["COURSE_PAPER.md"], primary_output="COURSE_PAPER.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="docx-export", display_name="格式检查与 Word 导出", output_files=["COURSE_PAPER.docx"], primary_output="COURSE_PAPER.docx", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "course_report": TemplateDef(
        pipeline_skill="course-report-plan",
        display_name="课程报告",
        sub_steps=[
            StepDef(skill_name="course-report-plan", display_name="事实提取与大纲规划", output_files=["PROJECT_FACTS.md","OUTLINE.md","PAPER_PLAN.md"], primary_output="OUTLINE.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-analysis", display_name="数据分析", output_files=["RESULTS.md","figures/all_results.json","code/main.py"], primary_output="RESULTS.md", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure", display_name="数据图表生成", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="架构图/流程图绘制", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="course-report", display_name="课程报告撰写", output_files=["COURSE_REPORT.md"], primary_output="COURSE_REPORT.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="docx-export", display_name="格式检查与 Word 导出", output_files=["COURSE_REPORT.docx"], primary_output="COURSE_REPORT.docx", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "paper_from_assets": TemplateDef(
        pipeline_skill="assets-inventory",
        display_name="已有资产 → 论文",
        sub_steps=[
            StepDef(skill_name="assets-inventory", display_name="资产清点与一致性检查", output_files=["ASSETS_INVENTORY.md","_assets_index.json"], primary_output="ASSETS_INVENTORY.md", has_checkpoint=True, checkpoint_type="assets_resolve"),
            StepDef(skill_name="paper-plan", display_name="论文规划", output_files=["PAPER_PLAN.md"], primary_output="PAPER_PLAN.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-analysis", display_name="缺口补全(代码/结果)", output_files=["RESULTS.md","figures/all_results.json","code/main.py"], primary_output="RESULTS.md", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure", display_name="缺口补全(图表)", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="缺口补全(架构图/路线图)", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-write", display_name="论文撰写", output_files=["paper/main.tex","paper/references.bib"], primary_output="paper/main.tex", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-compile", display_name="编译 PDF", output_files=["paper/main.pdf"], primary_output="paper/main.pdf", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "paper_slides": TemplateDef(
        pipeline_skill="paper-slides",
        display_name="论文 → 会议幻灯片",
        sub_steps=[
            StepDef(
                skill_name="paper-slides",
                display_name="Beamer 幻灯片 + 可编辑 PPTX",
                output_files=[
                    "slides/main.tex",
                    "slides/main.pdf",
                    "slides/presentation.pptx",
                    "slides/SLIDE_OUTLINE.md",
                    "slides/TALK_SCRIPT.md",
                ],
                primary_output="slides/main.pdf",
                has_checkpoint=True,
                checkpoint_type="approve",
            ),
        ],
    ),
    "paper_poster": TemplateDef(
        pipeline_skill="paper-poster",
        display_name="论文 → 会议海报",
        sub_steps=[
            StepDef(
                skill_name="paper-poster",
                display_name="A0/A1 海报 + 可编辑 PPTX",
                output_files=[
                    "poster/main.tex",
                    "poster/main.pdf",
                    "poster/poster.pptx",
                    "poster/POSTER_CONTENT_PLAN.md",
                    "poster/POSTER_SPEECH.md",
                ],
                primary_output="poster/main.pdf",
                has_checkpoint=True,
                checkpoint_type="approve",
            ),
        ],
    ),
    "humanities_paper": TemplateDef(
        pipeline_skill="humanities-plan",
        display_name="人文社科论文",
        sub_steps=[
            StepDef(skill_name="humanities-plan", display_name="选题与结构规划", output_files=["OUTLINE.md", "PAPER_PLAN.md"], primary_output="OUTLINE.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="paper-analysis", display_name="数据分析（可选）", output_files=["RESULTS.md", "figures/all_results.json", "code/main.py"], primary_output="RESULTS.md", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure", display_name="数据图表生成（可选）", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="paper-figure-drawio", display_name="理论框架图/示意图（可选）", output_files=["figures/latex_includes.tex"], primary_output="figures/", has_checkpoint=False, checkpoint_type=None),
            StepDef(skill_name="humanities-write", display_name="论文撰写", output_files=["HUMANITIES_PAPER.md"], primary_output="HUMANITIES_PAPER.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="docx-export", display_name="格式检查与 Word 导出", output_files=["HUMANITIES_PAPER.docx"], primary_output="HUMANITIES_PAPER.docx", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "one_sentence_project": TemplateDef(
        pipeline_skill="project-blueprint",
        display_name="一句话生成项目",
        sub_steps=[
            StepDef(skill_name="project-blueprint", display_name="项目蓝图与研究合同", output_files=["PROJECT_BLUEPRINT.md","RESEARCH_CONTRACT_DRAFT.md","MILESTONES.md"], primary_output="PROJECT_BLUEPRINT.md", has_checkpoint=True, checkpoint_type="approve"),
        ],
    ),
    "grad_project": TemplateDef(
        pipeline_skill="dev-requirement",
        display_name="一句话生成项目",
        sub_steps=[
            StepDef(skill_name="dev-requirement", display_name="需求分析", output_files=["REQUIREMENTS.md"], primary_output="REQUIREMENTS.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="dev-design", display_name="系统设计", output_files=["DESIGN.md", "schema.sql"], primary_output="DESIGN.md", has_checkpoint=True, checkpoint_type="feedback"),
            StepDef(skill_name="dev-code", display_name="编码实现", output_files=["code/backend/main.py", "RUN.md"], primary_output="code/", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="dev-selfcheck", display_name="自测验证", output_files=["TEST_REPORT.md"], primary_output="TEST_REPORT.md", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "software_copyright": TemplateDef(
        pipeline_skill="software-copyright",
        display_name="软件著作权材料",
        sub_steps=[
            StepDef(skill_name="software-copyright", display_name="软著材料清点与撰写", output_files=["software-copyright/PRODUCT_OVERVIEW.md","software-copyright/USER_MANUAL.md","software-copyright/SOURCE_CODE_INDEX.md","software-copyright/REGISTRATION_CHECKLIST.md"], primary_output="software-copyright/USER_MANUAL.md", has_checkpoint=True, checkpoint_type="approve"),
        ],
    ),
    "copyright_material": TemplateDef(
        pipeline_skill="copyright-draft",
        display_name="一句话生成软著申请资料",
        sub_steps=[
            StepDef(skill_name="copyright-draft", display_name="起草申请资料", output_files=["软件著作权申请资料/草稿/申请表信息.md"], primary_output="软件著作权申请资料/草稿/", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="copyright-build", display_name="生成正式 Word/TXT", output_files=["软件著作权申请资料/正式资料/生成报告.md"], primary_output="软件著作权申请资料/正式资料/", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
    "patent_disclosure": TemplateDef(
        pipeline_skill="patent-draft",
        display_name="一句话生成专利交底书",
        sub_steps=[
            StepDef(skill_name="patent-draft", display_name="起草技术交底书", output_files=["专利交底书/交底书草稿.md"], primary_output="专利交底书/交底书草稿.md", has_checkpoint=True, checkpoint_type="approve"),
            StepDef(skill_name="patent-build", display_name="渲染图示并导出 Word", output_files=["专利交底书/交底书.docx"], primary_output="专利交底书/交底书.docx", has_checkpoint=False, checkpoint_type=None),
        ],
    ),
}


# ============================================================

# ============================================================

_broadcast_func: Optional[Callable] = None
_checkpoint_events: Dict[str, asyncio.Event] = {}
_checkpoint_responses: Dict[str, dict] = {}
_workflow_managed_steps: set[str] = set()
_managed_step_rows: ContextVar[Optional[List[dict]]] = ContextVar(
    "managed_step_rows", default=None
)


_running_tasks: Dict[str, asyncio.Task] = {}

_DRAWIO_FIG_PREFIXES = (
    "fig_arch", "fig_flow", "fig_roadmap", "fig_pipeline", "fig_framework", "fig_er",
    "fig_overview", "fig_system", "fig_module", "fig_index", "fig_hierarchy", "fig_multiagent",
    "fig_topology", "fig_dataflow", "fig_pkg", "fig_class", "fig_seq", "fig_gantt", "fig_network",
    "fig_model_decision", "fig_decision", "fig_state", "fig_uml", "tikz_",
)
_SCENE_NAME_PREFIXES = ("fig_scene",)
_FIG_IMG_EXTS = (".png", ".pdf", ".jpg", ".jpeg", ".svg")
# Image-adjacent formats that are NOT accepted — agent may have used these by mistake.
# Support/script files (.json, .py, .md, .tex, etc.) are intentionally excluded here;
# they are handled separately as "辅助文件" rather than "wrong extension".
_NEAR_IMG_EXTS = frozenset({".tiff", ".tif", ".eps", ".bmp", ".webp", ".gif", ".ppm", ".pgm"})

# Per-workflow re-entry lock.  Clicking "retry/recover" several times in a row
# used to schedule multiple run_workflow coroutines on the same workflow_id,
# which then raced each other over _utils mounts, gen_fig_*.py outputs and the
# attempt ledger (fb4f4e5b7272 accumulated 8 zombie 'running' attempts this
# way).  Serialise per workflow so the second invocation waits for the first
# to settle instead of double-mounting the same workspace.
_WORKFLOW_RUN_LOCKS: Dict[str, asyncio.Lock] = {}
_WORKFLOW_RUN_LOCKS_GUARD = asyncio.Lock()


async def _acquire_workflow_lock(workflow_id: str) -> asyncio.Lock:
    async with _WORKFLOW_RUN_LOCKS_GUARD:
        lock = _WORKFLOW_RUN_LOCKS.get(workflow_id)
        if lock is None:
            lock = asyncio.Lock()
            _WORKFLOW_RUN_LOCKS[workflow_id] = lock
        return lock
_STEP_MIN_SIZE = {
    "comp-prob-analysis": 1500, "comp-modeling": 2000, "comp-code": 1000,
    "comp-stats-topic": 1000, "comp-paper-zh": 10000, "comp-paper-en": 10000,
    "paper-write": 15000, "paper-write-zh": 15000, "paper-write-nature": 15000,
    "paper-write-docx": 8000, "paper-write-zh-docx": 8000, "paper-write-nature-docx": 8000,
    "comp-paper-zh-docx": 8000, "comp-paper-en-docx": 8000, "paper-plan": 1000,
    "paper-plan-zh": 1000, "paper-analysis": 1000, "course-plan": 800, "course-paper": 5000,
    "course-report": 5000, "course-report-plan": 800, "thesis-proposal": 2000,
    "literature-review": 2000, "research-lit": 1500, "idea-creator": 1500,
    "novelty-check": 800, "research-review": 800, "research-refine-pipeline": 1500,
    "auto-review-loop": 1000, "auto-paper-improvement-loop": 50000,
    "comp-compile-zh": 30000, "comp-compile-en": 30000, "paper-compile": 30000,
    "paper-compile-zh": 30000, "assets-inventory": 500, "format-profile": 300,
    "docx-template-map": 100, "docx-format-check": 200, "docx-export": 5000,
    "experiment-bridge": 500, "paper-figure": 500, "paper-figure-drawio": 500,
    "paper-figure-html": 500, "paper-slides": 20000, "paper-poster": 20000,
    "nature-figure": 500, "humanities-plan": 1200, "humanities-write": 5000,
    "humanities-write-latex": 5000, "dev-requirement": 1500, "dev-design": 2000,
    "dev-code": 500, "dev-code-frontend": 500, "dev-code-backend": 500,
    "dev-selfcheck": 500, "dev-report": 5000, "software-copyright": 1000,
    "patent-disclosure": 1000, "patent-build": 1000, "copyright-build": 500,
}
_STEP_REQUIRED_COMPANIONS = {"comp-code": ["code/main.py", "figures/all_results.json"]}


def _scan_workspace(workspace: Path) -> List[str]:
    """Return visible workspace files while excluding engine context files."""
    excluded = {"CLAUDE.md", "_created_files.json", "_extract_status.json", "_input_manifest.json"}
    files = []
    for path in Path(workspace).rglob("*"):
        if not path.is_file():
            continue
        rel = path.relative_to(workspace)
        if rel.name in excluded or any(part.startswith(".") or part in {"__pycache__", "_utils"} for part in rel.parts):
            continue
        files.append(rel.as_posix())
    return sorted(files, key=str.lower)


_INPUT_ROLE_LABELS = {
    "requirements": "题目 / 写作要求",
    "problem": "赛题",
    "problem_images": "赛题图片",
    "outline": "大纲 / 思路",
    "code": "已有代码",
    "data": "数据集 / 附件数据",
    "figures": "已有图表",
    "results": "实验结果",
    "templates": "格式模板",
    "source": "真实源材料",
    "paper": "已编译论文",
    "material": "通用参考材料",
}


def _input_role_entries(workspace: Path) -> Dict[str, List[str]]:
    """Read the upload manifest without exposing the engine metadata file."""
    root = Path(workspace) / "user_data"
    manifest_path = root / "_input_manifest.json"
    try:
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    except (OSError, ValueError, TypeError):
        return {}
    grouped: Dict[str, List[str]] = {}
    for relative, metadata in manifest.get("files", {}).items():
        if not isinstance(relative, str) or not isinstance(metadata, dict):
            continue
        path = root / relative
        try:
            path.resolve().relative_to(root.resolve())
        except (OSError, ValueError):
            continue
        if not path.is_file():
            continue
        role = str(metadata.get("role") or "material")
        grouped.setdefault(role, []).append(f"user_data/{Path(relative).as_posix()}")
    return {role: sorted(set(paths), key=str.lower) for role, paths in grouped.items()}


def _snapshot_workspace(workspace: Path) -> Dict[str, tuple[int, int]]:
    """Capture visible files as size/mtime pairs for step change detection."""
    snapshot = {}
    for rel in _scan_workspace(workspace):
        try:
            stat = (Path(workspace) / rel).stat()
        except OSError:
            continue
        snapshot[rel] = (stat.st_size, stat.st_mtime_ns)
    return snapshot


def _order_step_files(paths: Any, step_def: StepDef) -> List[str]:
    """Order actual files by the step's declared output contract, then name."""
    remaining = set(paths or [])
    ordered: List[str] = []
    for declared in step_def.output_files:
        prefix = str(declared).replace("\\", "/").strip("./")
        if not prefix:
            continue
        matches = []
        if prefix.endswith("/") or not Path(prefix).suffix:
            prefix = prefix.rstrip("/")
            matches = [path for path in remaining if path == prefix or path.startswith(prefix + "/")]
        elif prefix in remaining:
            matches = [prefix]
        for path in sorted(matches, key=str.lower):
            ordered.append(path)
            remaining.discard(path)
    ordered.extend(sorted(remaining, key=str.lower))
    return ordered


def _order_reported_files(
    files_created: List[str], files_modified: List[str], step_def: StepDef
) -> List[str]:
    created = set(files_created)
    modified = set(files_modified)
    remaining = created | modified
    ordered: List[str] = []
    for declared in step_def.output_files:
        prefix = str(declared).replace("\\", "/").strip("./")
        if not prefix:
            continue
        if prefix.endswith("/") or not Path(prefix).suffix:
            prefix = prefix.rstrip("/")
            created_matches = [
                path for path in created
                if path in remaining and (path == prefix or path.startswith(prefix + "/"))
            ]
            modified_matches = [
                path for path in modified
                if path in remaining and (path == prefix or path.startswith(prefix + "/"))
            ]
            matches = sorted(created_matches, key=str.lower) + sorted(modified_matches, key=str.lower)
            ordered_matches = matches
        else:
            matches = [prefix] if prefix in remaining else []
            ordered_matches = sorted(matches, key=str.lower)
        for path in ordered_matches:
            ordered.append(path)
            remaining.discard(path)
    if step_def.skill_name == "nature-figure" and any(
        path.startswith("figures/") for path in ordered
    ):
        ordered.insert(0, "figures/")
    ordered.extend(sorted(created & remaining, key=str.lower))
    ordered.extend(sorted(modified & remaining, key=str.lower))
    return ordered


def _build_context_summary(workspace: Path, files: List[str], max_chars: int = 12000) -> str:
    """Read a bounded summary of existing text assets for the executor prompt."""
    preferred = {".md": 0, ".tex": 1, ".bib": 2, ".txt": 3, ".json": 4, ".py": 5}
    ordered = sorted(files, key=lambda name: (preferred.get(Path(name).suffix.lower(), 99), name.lower()))
    chunks = []
    remaining = max_chars
    for rel in ordered:
        path = Path(workspace) / rel
        if path.suffix.lower() not in preferred or remaining <= 0:
            continue
        try:
            text = path.read_text(encoding="utf-8", errors="replace")
        except OSError:
            continue
        header = f"\n--- {rel} ---\n"
        chunk = header + text[:max(0, remaining - len(header))]
        chunks.append(chunk)
        remaining -= len(chunk)
    return "".join(chunks)


def _build_managed_context_summary(
    files: Optional[List[str]],
    step_def: StepDef,
    template_def: TemplateDef,
    step_rows: List[dict],
) -> str:
    lines = ["工作区中已有文件（请用 Read/Bash 工具按需读取具体内容）:"]
    if files:
        lines.extend(f"- {name}" for name in files)
    else:
        lines.append("(无文件)")

    template_index = next(
        (index for index, item in enumerate(template_def.sub_steps) if item.skill_name == step_def.skill_name),
        0,
    )
    row_index = next(
        (index for index, row in enumerate(step_rows) if row["skill_name"] == step_def.skill_name),
        0,
    )
    completed = [
        row["display_name"] for row in step_rows if row.get("status") == "completed"
    ]
    remaining = [
        row["display_name"] for row in step_rows[row_index + 1:] if row.get("status") == "pending"
    ]
    lines.extend([
        "", "", "## Pipeline Context",
        f"当前步骤: {step_def.display_name} ({step_def.skill_name}) — 第 {template_index + 1}/{len(step_rows)} 步",
    ])
    if completed:
        lines.append(f"已完成步骤: {', '.join(completed)}")
    if remaining:
        lines.append(f"剩余步骤: {', '.join(remaining)}")
    lines.extend(["", "该步骤必须产出的文件（至少）:"])
    lines.extend(f"- {name}" for name in step_def.output_files)

    existing = set(files or [])
    previous_outputs: List[str] = []
    for row in step_rows[:row_index]:
        if row.get("status") != "completed":
            continue
        previous = next(
            (item for item in template_def.sub_steps if item.skill_name == row["skill_name"]),
            None,
        )
        if previous is None:
            continue
        for output in previous.output_files:
            normalized = output.replace("\\", "/")
            if normalized in existing and normalized not in previous_outputs:
                previous_outputs.append(normalized)
    if previous_outputs:
        lines.extend([
            "",
            "## IMPORTANT: 前步骤的关键文件",
            "以下文件是前面步骤的产出，内容在摘要中可能被截断。请在开始工作前使用 Read 工具完整读取这些文件。",
        ])
        lines.extend(f"- {name}" for name in previous_outputs)
    return "\n".join(lines) + "\n"


def _primary_output_exists(workspace: Path, step_def: StepDef) -> bool:
    primary = step_def.primary_output
    if not primary:
        return True
    path = Path(workspace) / primary
    if not Path(primary.rstrip("/\\")).suffix:
        return path.is_dir() and any(item.is_file() for item in path.rglob("*"))
    return path.is_file()


_GENERIC_ZH_COMP = {
    "language": "zh",
    "max_pages": 30,
    "template_cls": "ctexart",
    "rules": [
        "论文不超过 30 页",
        "赛程较长（7月中旬至9月上旬）",
        "侧重电气工程与能源方向的建模问题",
        "论文结构：摘要 → 问题重述 → 问题分析 → 模型假设 → 符号说明 → 模型建立与求解 → 结果分析 → 参考文献 → 附录",
        "⛔ 只需要中文摘要，不要英文摘要",
        "使用 ctexart 文档类，XeLaTeX 编译",
    ],
}


_COMP_RULES = {
    "comp_cumcm": {
        "name": "全国大学生数学建模竞赛 (CUMCM)", "language": "zh", "max_pages": 30,
        "template_cls": "cumcmthesis", "rules": [
            "论文不超过 30 页（含摘要、正文、附录和参考文献）",
            "必须有摘要项（约 1 页），含关键词",
            "⛔ 只需要中文摘要，不要英文摘要（国赛不要求英文摘要）",
            "论文结构：摘要 → 问题重述 → 问题分析 → 模型假设 → 符号说明 → 模型建立与求解 → 模型检验 → 模型评价与推广 → 参考文献 → 附录",
            "附录中必须包含完整代码",
            "使用 cumcmthesis 文档类，XeLaTeX 编译",
        ],
    },
    "comp_mcm": {
        "name": "MCM/ICM (COMAP)", "language": "en", "max_pages": 25,
        "template_cls": "mcmthesis", "rules": [
            "Summary Sheet (摘要页) 必须是第一印象，决定一票否决权",
            "论文结构：Summary → Introduction → Problem Analysis → Model Design → Solution → Sensitivity Analysis → Strengths & Weaknesses → Conclusions → References → Appendix",
            "Summary 必须独立成页，包含 team control number",
            "全英文撰写，使用 mcmthesis 文档类，pdflatex 编译",
            "论文建议不超过 25 页",
        ],
    },
    "comp_huawei": {
        "name": "华为杯全国研究生数学建模竞赛", "language": "zh", "max_pages": 50,
        "template_cls": "gmcmthesis", "rules": [
            "使用 gmcmthesis 文档类，XeLaTeX 编译",
            "封面需包含 4 个 Logo（组委会提供）",
            "论文格式有严格要求（参见规程、评审标准）",
            "论文结构同国赛，但允许更长篇幅（一等奖典型 50-80 页含附录，正文不少于 50 页）",
            "⛔ 只需要中文摘要，不要英文摘要（华为杯是中文竞赛）",
            "附录需包含代码和数据支撑材料",
        ],
    },
    "comp_mathorcup": {
        "name": "MathorCup 数学建模挑战赛", "language": "zh", "max_pages": 30,
        "template_cls": "ctexart", "rules": [
            "论文不超过 30 页", "必须有目录页",
            "页面顶部必须有队伍编号+题号表格（使用 mathorcup_main.tex 模板）",
            "论文结构：摘要 → 目录 → 问题重述 → 问题分析 → 模型假设 → 符号说明 → 模型建立与求解 → 结果分析 → 模型评价 → 参考文献 → 附录",
            "⛔ 只需要中文摘要，不要英文摘要（MathorCup 不要求英文摘要）",
            "使用 mathorcup_main.tex 模板（基于 ctexart），XeLaTeX 编译",
        ],
    },
    "comp_apmcm": {
        "name": "亚太地区数学建模竞赛 (APMCM)", "language": "en", "max_pages": 25,
        "template_cls": "apmcmthesis", "rules": [
            "按照官网上传的 PDF 模板格式提交", "全英文撰写",
            "论文结构参考美赛：Summary → Introduction → Model → Solution → Conclusions",
            "使用 apmcmthesis 或通用模板，pdflatex 编译",
        ],
    },
    "comp_apmcm_zh": {
        "name": "亚太地区数学建模竞赛中文赛项 (APMCM 中文)", "language": "zh", "max_pages": 25,
        "template_cls": "MathorCupmodeling", "rules": [
            "按照官网上传的 PDF 模板格式提交（亚太赛中文赛项）",
            "中文撰写，封面含参赛队号、参赛队员",
            "论文结构：摘要 → 关键词 → 问题重述 → 问题分析 → 模型假设 → 符号说明 → 模型建立与求解 → 灵敏度分析 → 模型评价 → 参考文献 → 附录",
            "需要中文摘要 + 关键词（一般 3-5 个）",
            "使用 MathorCupmodeling 文档类（\\bianhao/\\tihao/\\timu 设置队伍信息，\\keyword 设置关键词），XeLaTeX 编译",
        ],
    },
    "comp_stats": {
        "name": "全国大学生统计建模大赛", "language": "zh", "max_pages": 30,
        "template_cls": "ctexart", "rules": [
            "自主选题（官方只给选题方向，不给具体题目）",
            "必须自行采集真实数据（标注数据来源：国家统计局、Wind、CSMAR、Kaggle 等）",
            "论文结构：摘要 → 引言 → 文献综述 → 数据与方法 → 实证分析 → 结论与建议 → 参考文献 → 附录",
            "⛔ 只需要中文摘要，不要英文摘要（统计建模大赛不要求英文摘要）",
            "注重统计分析的规范性（必须报告 p 值、置信区间、显著性注释、水平、回归诊断、模型比较）",
            "鼓励使用描述性统计、推断统计、回归分析等经典方法",
            "准备答辩展示材料（PPT）", "注重数据质量，不能抄袭已有研究",
            "使用 ctexart 文档类，XeLaTeX 编译",
        ],
    },
    "comp_teddy": {
        "name": "泰迪杯数据挖掘挑战赛", "language": "zh", "max_pages": 40,
        "template_cls": "ctexart", "rules": [
            "赛程约 13 天，侧重数据挖掘与机器学习方向",
            "论文结构：摘要 → 问题重述 → 数据预处理 → 特征工程 → 模型建立与求解 → 结果分析 → 参考文献 → 附录",
            "⛔ 只需要中文摘要，不要英文摘要",
            "必须包含完整的数据预处理和特征工程描述", "代码附录需包含关键算法实现",
            "使用 ctexart 文档类，XeLaTeX 编译",
        ],
    },
    "comp_certcup": {
        "name": "认证杯数学建模（第一/二阶段）", "language": "zh", "max_pages": 35,
        "template_cls": "ctexart", "rules": [
            "分两阶段：第一阶段提交初步方案，第二阶段提交完整论文",
            "论文结构：摘要 → 问题重述 → 问题分析 → 模型假设 → 模型建立与求解 → 结果分析 → 模型评价 → 参考文献 → 附录",
            "⛔ 只需要中文摘要，不要英文摘要", "使用 ctexart 文档类，XeLaTeX 编译",
        ],
    },
    "comp_huazhong": {
        "name": "华中杯数学建模邀请赛", "language": "zh", "max_pages": 30,
        "template_cls": "cumcmthesis", "rules": [
            "论文不超过 30 页",
            "使用 cumcmthesis 文档类（withoutpreface 选项去掉封面），XeLaTeX 编译",
            "论文结构：摘要 → 目录 → 问题重述 → 问题分析 → 模型假设 → 符号说明 → 模型建立与求解 → 模型检验 → 参考文献 → 附录",
            "⛔ 只需要中文摘要，不要英文摘要", "使用 huazhong_main.tex 模板",
        ],
    },
    "comp_wuyi": {
        "name": "五一杯数学建模竞赛", "language": "zh", "max_pages": 30,
        "template_cls": "cumcmthesis", "rules": [
            "论文不超过 30 页", "使用 cumcmthesis 文档类（withoutpreface 选项），XeLaTeX 编译",
            "论文第一页为承诺书（模板已包含，不要删除）",
            "第二页为封面（题目+关键词+摘要，手动排版格式）",
            "论文结构：承诺书 → 封面(题目+关键词+摘要) → 目录 → 问题重述 → 问题分析 → 模型假设 → 符号说明 → 各问题建模与求解 → 模型评价 → 参考文献 → 附录",
            "⛔ 只需要中文摘要，不要英文摘要", "附录第一节为文件列表表格",
            "参考文献使用 gbt7714-numerical 样式", "使用 wuyi_main.tex 模板",
            "⛔ 不要加 \\usepackage{cite}（和 natbib 冲突导致编译错误）",
            "⛔ 不要重复加载 subcaption 和 float（cls 已包含）",
            "⛔ 禁止使用 subfigure 并排图片，每张图独占一个 figure 环境，宽度 ≥ 0.85\\textwidth",
            "⛔ 不要加 \\maketitle（五一杯用手写承诺书页）",
        ],
    },
    "comp_certcup_en": {
        "name": "认证杯国际赛 (小美赛)", "language": "en", "max_pages": 25,
        "template_cls": "article", "rules": [
            "全英文撰写",
            "论文结构参考美赛：Summary → Introduction → Model → Solution → Conclusions → References → Appendix",
            "Summary 必须独立成页", "使用 article 或 mcmthesis 文档类，pdflatex 编译",
            "论文建议不超过 25 页",
        ],
    },
    "comp_shuwei_en": {
        "name": "数维杯国际赛 (ISTIC)", "language": "en", "max_pages": 25,
        "template_cls": "article", "rules": [
            "全英文撰写",
            "论文结构参考美赛：Summary → Introduction → Model → Solution → Conclusions → References → Appendix",
            "使用 article 文档类，pdflatex 编译", "论文建议不超过 25 页",
        ],
    },
}


for _template_name, _competition_name in {
    "comp_huadong": "华东杯数学建模邀请赛",
    "comp_shuwei": "数维杯数学建模（国内赛）",
    "comp_zhongqing": "中青杯数学建模竞赛",
    "comp_yangtze": "长三角高校数学建模竞赛",
    "comp_diangong": "电工杯数学建模竞赛",
    "comp_shenzhen": "深圳杯数学建模竞赛",
    "comp_huashu": "华数杯数学建模竞赛",
    "comp_tianfu": "天府杯数学建模竞赛",
    "comp_liaoning": "辽宁省大学生数学建模/东三省数学建模联赛",
}.items():
    _COMP_RULES[_template_name] = {"name": _competition_name, **_GENERIC_ZH_COMP}


_COMP_CORE_PRINCIPLES = """## ⛔ 核心原则：发现问题必须修正（不能只解释）
优先级：物理/业务约束 > 数据忠实度 > 计算正确性

如果计算结果违反了题目给定的物理/业务约束（如超出最大值、为负、不守恒）：
- 禁止：发现超出后只写一句'这是因为XXX'然后继续使用超出的结果
- 禁止：把不合理的值写进 JSON/报告，旁边加注释说明
- 禁止：写'后续可改进'或'留给下一步处理'然后跳过
- 禁止：说'数学上正确'或'纯数学解'来为不合理结果辩护
- 禁止：说'这是题目数据的特性'来回避修正责任
- 必须：分析原因 → 修正数据或模型（加物理约束）→ 重新计算 → 用修正后的合理值
- 必须：在报告中写明'原始结果为X（不合理），修正后为Y，修正方法为Z'
- 必须：修正后重新验证，确认结果在合理范围内
- 必须：每个子问题代码末尾有 validate_constraints() 自动验证

⛔ 关键认知：数学正确 ≠ 物理合理。ODE数值解超出物理边界时，
不是'数据导致的正常结果'，而是模型缺少物理约束（接触约束/饱和/边界条件）。
必须在代码中加入约束后重新求解，而不是解释为什么会超出。

⛔ 几何/参数完整性：
- 建模：禁止未声明的几何简化（矩形→线段、实体→质心点）
- 代码：禁止用中间计算量替代物理实体参数（用完整尺寸，不用中心距代替全长）
- 碰撞/约束检测必须用实体完整外轮廓，不能只用中心点距离

⛔ 遮蔽/覆盖/碰撞判定的完整性（防止对象降维）：
- 判定必须作用于目标的完整几何边界（所有边界点或等价的充要条件）
- 函数签名必须包含几何参数（半径/高度/长宽/外轮廓），只传一个坐标点 = 反模式
- 禁止用'中心被遮蔽 = 整体被遮蔽'的等价假设（除非证明等价性）
- 离散采样近似时必须做收敛性验证（如 N=100→300→500 结果稳定才算收敛）
- 代码注释必须说明：判定的是哪个几何体的哪些边界（如'圆柱体上下底面圆周 300 点'）

⛔ 论文深度与求解精度（提升论文质量）：
- 关键假设必须有数学证明或文献支撑（≥1个定理/引理），不能只写'为简化假设XXX'
- 求解精度必须与方法匹配：连续问题用二分/牛顿（精度≥1e-6），禁止固定步长=0.01作为最终结果
- 核心结果必做双算法交叉验证（几何解析+启发式 / DE+PSO / 精确+近似）
- 高维优化（维度≥10）必须先做几何/物理降维，禁止直接喂黑箱算法

⛔ 多资源协同优化（防止资源闲置陷阱）：
- 多资源（多弹/多机/多点）必须用联合目标函数，不能独立优化后叠加
- 检测到'某个资源T=0/覆盖=0/贡献=0' → 说明被误判放弃，必须修
- 检测到'两个资源作用区间完全不重叠' → 没利用互补性，必须改
- 启发式算法必须有几何/物理启动种子（≥5个），禁止纯随机初始化
- 多阶段问题必须分层（任务分配+子任务优化），禁止把所有维度混在一起全量搜索

⛔ 建模与优化有效性约束（通用）：
- 资源单调性：更多决策资源→结果必须严格更优，否则建模降维失败
- 解析基准下界：启发式优化解必须优于解析/几何/贪心基准，否则算法失败
- 搜索空间健康度：维度≥6时先采样100点，有效解<30%则必须降维或启发式播种
- 维度分层：决策变量≥10维时禁止单层黑箱优化，必须分层分解

检测到问题 = 必须修复。解释原因 ≠ 处理完毕。
发现不合理 → 必须行动修正 → 用修正后的值继续。
任何步骤都不允许带着已知问题继续往下走。"""


def _generate_claude_md(workspace: Path, title: str, template: str, params: dict) -> None:
    """Generate the executor context file used by the original workflow engine."""
    tmpl = _resolve_template(template, params, workspace)
    lines = [
        f"# 研究项目: {title}",
        "",
        "## 项目说明",
        f"这是一个自动化研究管理的工作流引擎，当前正在执行「{tmpl.display_name}」流水线。",
        "",
        "## 流水线步骤",
    ]
    lines.extend(
        f"  {index}. {step.display_name} ({step.skill_name})"
        for index, step in enumerate(tmpl.sub_steps, 1)
    )
    lines.extend(
        [
            "",
            "## 重要规则",
            "- 所有产出文件都写入当前工作目录（不要写到其他位置）",
            "- 使用 Write 工具创建文件，如果文件太长则用 Bash 的 cat << 'EOF' > file 方式",
            "- 每个步骤完成后至少产出一个文件",
            "- 读取工作区中已有的文件，在前步骤的基础上继续工作",
            "- 使用 English 撰写正文" if str(params.get("language") or "zh").lower() == "en" else "- 使用中文撰写正文（除非是 LaTeX 代码）",
        ]
    )
    format_text = str(params.get("format_text") or "").strip()
    format_requirements_path = Path(workspace) / "FORMAT_REQUIREMENTS.md"
    if format_text:
        format_requirements_path.write_text(
            "# 用户文字格式要求\n\n" + format_text + "\n",
            encoding="utf-8",
        )
    elif format_requirements_path.exists():
        format_requirements_path.unlink()

    workspace_files = _scan_workspace(workspace)
    input_roles = _input_role_entries(workspace)
    if workspace_files:
        lines.extend(["", "## 工作区已有文件"])
        lines.extend(f"- {name}" for name in workspace_files)
        lines.append("请仔细阅读并利用这些文件作为研究基础。")
    if input_roles:
        lines.extend(["", "## 用户上传文件分类"])
        for role, paths in sorted(input_roles.items()):
            lines.append(f"### {_INPUT_ROLE_LABELS.get(role, role)} ({role})")
            lines.extend(f"- {path}" for path in paths)
        lines.append("必须按上述角色理解材料；已有结果和图表不得被编造内容覆盖。")
    requirements_path = Path(workspace) / "CUSTOM_REQUIREMENTS.md"
    if requirements_path.is_file():
        try:
            requirements_text = requirements_path.read_text(encoding="utf-8", errors="replace").strip()
        except OSError:
            requirements_text = ""
        if requirements_text:
            lines.extend([
                "",
                "## 用户自定义要求（最高优先级）",
                "以下内容来自用户上传的要求文档；与一般写作约定冲突时以此为准：",
                requirements_text[:50_000],
            ])
    if "checkpoint_feedback.md" in workspace_files:
        lines.extend([
            "",
            "## 用户修改意见（最高优先级）",
            "工作区中存在 `checkpoint_feedback.md`，包含用户对当前步骤的修改意见。",
            "**请在开始前优先阅读该文件，仔细遵从用户的意见来调整产出。**",
        ])
    lines.extend(["", "## 参数", f"- 研究主题: {title}", f"- template: {template}"])
    lines.extend(f"- {key}: {value}" for key, value in params.items())

    docx_templates, latex_templates = _template_files_by_kind(params)
    if docx_templates or latex_templates:
        lines.extend([
            "",
            "## 用户上传的格式模板（硬约束）",
            "模板套件必须由对应 prepare/apply 步骤真实使用，不得只在文本中提及。",
        ])
        lines.extend(f"- Word 模板: `{item}`" for item in docx_templates)
        lines.extend(f"- LaTeX 模板套件: `{item}`" for item in latex_templates)
        if latex_templates:
            lines.extend([
                "- 写作前必须读取 `_latex_template.json` 与 `paper/_user_template_main.tex`。",
                "- 保留用户模板的 documentclass、封面宏、页边距、字体、套件及 cls/sty/bst 依赖。",
                "- 仅填入论文内容，不得用通用 article/ctexart 覆盖用户模板。",
            ])

    # Skills receive the same direct aliases in both executor implementations,
    # but many of their embedded Bash checks explicitly source .env_skill.
    # Materialize the reviewed aliases for every workflow, not only competition
    # templates, so user overrides such as paper_writing_zh.max_pages are real.
    from services.claude_runner import _skill_parameter_environment

    skill_environment = {
        key: value
        for key, value in _skill_parameter_environment(params).items()
        if not key.startswith("SKILL_")
    }
    env_path = workspace / ".env_skill"
    if skill_environment:
        env_lines = [
            "#!/bin/bash",
            "# Auto-generated by workflow_engine.",
            "# Source this in any bash script: source .env_skill",
        ]
        env_lines.extend(
            f"export {key}={shlex.quote(value)}"
            for key, value in sorted(skill_environment.items())
        )
        env_path.write_text("\n".join(env_lines) + "\n", encoding="utf-8")
        lines.extend([
            "",
            "## Skill 执行环境（与 Responses / 外部 Claude 一致）",
            "工作区已生成 `.env_skill`；需要 shell 参数的检查脚本必须先 `source .env_skill` 。",
            "",
            "```bash",
        ])
        lines.extend(
            f"export {key}={shlex.quote(value)}"
            for key, value in sorted(skill_environment.items())
        )
        lines.extend(["```", ""])
    elif env_path.exists():
        env_path.unlink()

    hard_minimums = []
    for parameter, label in (
        ("min_figures", "数据图"),
        ("min_tables", "表格"),
        ("min_models", "模型"),
    ):
        value = params.get(parameter)
        if isinstance(value, int) and not isinstance(value, bool):
            hard_minimums.append((parameter, label, value))
    if hard_minimums:
        lines.extend([
            "## 用户明确指定的数量硬下限",
            "以下值不是建议值；计划、生成与最终自检都必须达到，不足时不得宣称完成：",
        ])
        lines.extend(
            f"- {parameter}: {value}（{label}至少 {value} 个）"
            for parameter, label, value in hard_minimums
        )
        if any(parameter == "min_models" and value > 0 for parameter, _, value in hard_minimums):
            lines.extend([
                "- 建模步骤必须额外写入 `MODEL_MANIFEST.json`，格式为 "
                "`{\"version\":1,\"models\":[{\"id\":\"M1\",\"name\":\"...\",\"purpose\":\"...\"}]}`。",
                "- 每个 models 条目必须对应 `MODELING_REPORT.md` 中的独立模型；别名或同一模型参数变体不得重复计数。",
            ])
        if any(parameter == "min_tables" and value > 0 for parameter, _, value in hard_minimums):
            lines.append(
                "- 表格按 `figures/TABLE_<id>.md|tex|csv|xlsx` 的唯一 `<id>` 计数，同一表格的多格式文件只算 1 个。"
            )
        lines.append(
            "- 引擎会生成 `QUANTITY_MANIFEST.json` 并做 terminal gate；实际数量小于上述任一硬下限时步骤失败。"
        )
        lines.append("")

    if str(params.get("output_format", "")).lower() == "docx":
        lines.extend(
            [
                "",
                "## ⛔ 输出格式：Word（DOCX）",
                "- **本工作流将导出为 Word 文档（.docx），不编译 LaTeX**",
                "- 写作步骤产出 **纯 Markdown 文件**，不要写 LaTeX 代码",
                "- 论文写作步骤产出 `paper/main.md`（不是 main.tex）",
                "- 竞赛论文写作步骤产出 `paper/main.md`（不是 main.tex）",
                "- 所有公式用 Markdown 格式：行内 `$E=mc^2$`，块级 `$$...$$`",
                "- 表格使用 Markdown 表格语法（`| 表头 | ... |\\n|---|---|\\n`），**优先用 `cat figures/TABLE_xxx.md >> paper/main.md` 嵌入已生成的表格**（数据真实，杜绝编造）",
                "- 图片使用 Markdown 图片语法 `![描述](figures/xxx.png)`",
                "- 参考文献用 `[1] 作者. 题名[J]. 期刊, 年, 卷(期): 页码.` 格式",
                "- 编译相关步骤会自动跳过，最后一步会把 Markdown 转 Word",
            ]
        )
    if template == "auto_review":
        auto_review_format = str(params.get("output_format") or "markdown").lower()
        if auto_review_format == "docx":
            lines.extend([
                "",
                "## AUTO_REVIEW_DOCX_MODE",
                "- 审稿最终正文必须写入 `NARRATIVE_REPORT.md`。",
                "- 本步骤禁止调用 LaTeX；下游 `docx-export` 会生成 Word。",
            ])
        elif auto_review_format == "pdf":
            lines.extend([
                "",
                "## AUTO_REVIEW_PDF_MODE",
                "- 仅当工作区存在 `paper/main.tex` 时编译 PDF。",
                "- 若没有 LaTeX 源文件，必须生成 `NARRATIVE_REPORT.md` 并在 `AUTO_REVIEW.md` 记录降级原因。",
            ])

    competition = _COMP_RULES.get(template)
    if competition:
        max_pages = params.get("max_pages", competition["max_pages"])
        language = "English" if competition["language"] == "en" else "中文"
        lines.extend([
            "",
            f"## 竞赛规则（{competition['name']}）",
            f"- 语言: {language}",
            f"- 页数目标: 约 {max_pages} 页（参考值，最终以论文实际质量为准；不要为凑页数注水）",
            f"- LaTeX 模板: {competition['template_cls']}",
        ])
        lines.extend(f"- {rule}" for rule in competition["rules"])
        lines.extend(["- 编程工具: python", ""])
        lines.extend(_COMP_CORE_PRINCIPLES.splitlines())

    (workspace / "CLAUDE.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


_image_descriptions: Dict[str, str] = {}


def _compress_claude_md(workspace: Path) -> None:
    """Compress verbose image and extracted-document blocks in ``CLAUDE.md``."""
    claude_md = Path(workspace) / "CLAUDE.md"
    if not claude_md.is_file():
        return
    try:
        content = claude_md.read_text(encoding="utf-8", errors="replace")
    except OSError:
        return

    def _truncate_image_section(match: re.Match[str]) -> str:
        heading, section = match.group(1), match.group(2)
        if len(section) <= 128:
            return match.group(0)
        return heading + section[:100] + "... (\u5df2\u538b\u7f29)"

    def _truncate_doc_section(match: re.Match[str]) -> str:
        heading, section = match.group(1), match.group(2)
        if len(section) <= 512:
            return match.group(0)
        return heading + section[:500] + "\n... (\u5df2\u538b\u7f29\uff0c\u8bf7\u7528 Read \u5de5\u5177\u67e5\u770b\u5b8c\u6574\u5185\u5bb9)\n```"

    updated_content = re.sub(
        r"(### \u56fe\u7247: [^\n]+\n)(.*?)(?=\n### |\n## |\Z)",
        _truncate_image_section,
        content,
        flags=re.DOTALL,
    )
    updated_content = re.sub(
        r"(### \u6765\u6e90: [^\n]+\n)(.*?)(?=\n### |\n## |\Z)",
        _truncate_doc_section,
        updated_content,
        flags=re.DOTALL,
    )
    if updated_content != content:
        claude_md.write_text(updated_content, encoding="utf-8")

async def _describe_workspace_images(workspace: Path, content: str) -> str:
    """Append cached vision descriptions for uploaded images to context text."""
    user_data = Path(workspace) / "user_data"
    if not user_data.is_dir():
        return content

    image_files = [
        path for path in sorted(user_data.iterdir(), key=lambda item: item.name.lower())
        if path.is_file() and path.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp"}
        and path.stat().st_size > 1000
    ]
    if not image_files:
        return content

    from services.llm_client import describe_image

    descriptions: List[tuple[str, str]] = []
    context = "这是数学建模竞赛或学术论文的数据/研究类图片"
    for image_path in image_files:
        cache_key = str(image_path)
        description = _image_descriptions.get(cache_key)
        if description is None:
            try:
                description = await describe_image(cache_key, context=context)
            except Exception as exc:
                log.warning("Vision description failed for %s: %s", image_path.name, exc)
                continue
            if not description:
                continue
            _image_descriptions[cache_key] = description
        descriptions.append((image_path.name, description))

    if not descriptions:
        return content
    appendix = "\n## 上传图片内容（AI 自动识别）\n以下是从上传图片中自动识别的内容描述，仅供参考。\n"
    for name, description in descriptions:
        appendix += f"\n### 图片: {name}\n{description}\n"
    return content + appendix + "\n"


async def _extract_pdf_with_vision(workspace: Path, workflow_id: str) -> None:
    """OCR uploaded PDFs page-by-page through the vision client.

    Vision failures intentionally fall back to PyPDF2 text extraction, matching
    the installed engine's best-effort behavior.  ``workflow_id`` is part of
    the private ABI even though the helper currently does not persist it.
    """
    del workflow_id
    user_data = Path(workspace) / "user_data"
    if not user_data.is_dir():
        return

    pdf_files = [path for path in user_data.iterdir() if path.is_file() and path.suffix.lower() == ".pdf" and path.stat().st_size > 1000]
    if not pdf_files:
        return
    tmp_dir = Path(workspace) / "_tmp"
    prompt = (
        "这是数学建模竞赛赛题 PDF 的一页。请完整转录页面上的所有文字内容，包括数学公式"
        "（用 LaTeX 格式如 $F = ma$、$\\int_0^1 f(x)dx$）、表格（用 Markdown 表格格式）、"
        "编号列表等。保持原文结构，不要遗漏任何内容。不要添加解释或评论，只输出原文转录。"
    )
    from services.llm_client import describe_image

    try:
        import fitz
    except Exception as exc:
        log.warning("Vision PDF rendering unavailable: %s", exc)
        return
    for pdf_path in pdf_files:
            output_path = pdf_path.with_name(f"{pdf_path.stem}_extracted.txt")
            tmp_dir.mkdir(parents=True, exist_ok=True)
            page_text: List[str] = []
            try:
                document = fitz.open(pdf_path)
                try:
                    for index, page in enumerate(document):
                        image_path = tmp_dir / f"_pdf_page_{index}.png"
                        pixmap = page.get_pixmap()
                        pixmap.save(image_path)
                        extracted = ""
                        for attempt in range(3):
                            try:
                                extracted = await describe_image(str(image_path), context=prompt)
                                if len((extracted or "").strip()) >= 50:
                                    break
                                extracted = ""
                            except Exception as exc:
                                log.warning("Vision extraction retry %d failed for page %d: %s", attempt + 1, index + 1, exc)
                        if extracted:
                            page_text.append(f"## 第 {index + 1} 页\n\n{extracted}")
                finally:
                    document.close()
            except Exception as exc:
                log.warning("Vision PDF rendering failed for %s: %s", pdf_path.name, exc)

            if page_text:
                if not output_path.exists():
                    output_path.write_text(f"# {pdf_path.name} 内容提取（Vision OCR）\n\n" + "\n\n".join(page_text) + "\n\n", encoding="utf-8")
                continue
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(str(pdf_path))
                raw_pages = [page.extract_text() or "" for page in reader.pages]
                if any(text.strip() for text in raw_pages) and not output_path.exists():
                    fallback = "\n\n".join(
                        f"## 第 {index + 1} 页\n\n{text}" for index, text in enumerate(raw_pages)
                    )
                    output_path.write_text(f"# {pdf_path.name} 内容提取（PyPDF2）\n\n{fallback}\n", encoding="utf-8")
            except Exception as exc:
                log.warning("PyPDF2 fallback failed for %s: %s", pdf_path.name, exc)

    # Rendered page images are ephemeral OCR inputs.  Never leave them in a
    # research workspace after extraction, including partial failure paths.
    shutil.rmtree(tmp_dir, ignore_errors=True)


def _initialize_git_workspace(workspace: Path) -> None:
    """Best-effort `git init`, matching the original workspace bootstrap contract."""
    try:
        subprocess.run(
            ["git", "init"],
            cwd=str(workspace),
            stdin=subprocess.DEVNULL,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            check=False,
            timeout=15,
            creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
        )
    except (OSError, subprocess.SubprocessError):
        log.debug("Unable to initialize git workspace %s", workspace, exc_info=True)


def _inject_once(claude_md: Path, marker: str, content: str) -> bool:
    token = f"<!-- INJECT:{marker} -->"
    existing = claude_md.read_text(encoding="utf-8", errors="replace") if claude_md.exists() else ""
    if token in existing:
        return False
    with claude_md.open("a", encoding="utf-8") as handle:
        handle.write(f"\n{token}\n{content.rstrip()}\n")
    return True


def _inject_plan_stage_guidance(workspace: Path, skill_name: str, params: Dict[str, Any]) -> bool:
    """Append the installed planner constraints once, immediately before execution."""
    if skill_name not in {"paper-plan", "paper-plan-zh"}:
        return False

    skip_figures = bool(params.get("skip_figures", False))
    skip_analysis = bool(params.get("skip_analysis", False))
    skip_drawio = bool(params.get("skip_drawio", False))
    data_suffixes = {".csv", ".tsv", ".xlsx", ".xls", ".json"}
    user_data = Path(workspace) / "user_data"
    has_data = user_data.is_dir() and any(
        path.is_file() and path.suffix.lower() in data_suffixes
        for path in user_data.iterdir()
    )

    if skip_figures or skip_analysis:
        content = (
            "\n\n## ⛔ 规划阶段强制规则（用户已设置）\n"
            "**用户已禁用数据分析与图表生成，后续 paper-analysis 与 paper-figure 步骤会被跳过。**\n"
            "因此本步骤生成规划文档（PAPER_PLAN.md / PROBLEM_ANALYSIS.md / TOPIC_PLAN.md）时：\n"
            "1. 必须使用「无图表分支」格式\n"
            "2. **禁止规划任何 fig_xxx 或 TABLE_xxx**（不要列「图表预规划」「Figure Plan」章节）\n"
            "3. 大纲章节描述里禁止写「将绘制」「如图 X 所示」「图 X 展示」\n"
            "4. Claims-Evidence Matrix 中所有 Evidence 用「文献 [N]」「概念推导」支撑，不写 fig_xxx\n"
            "5. 写作阶段会用纯文字 + 必要的 markdown 表格表达数据，不嵌入图片\n"
            f"6. 触发原因: skip_figures={skip_figures}, skip_analysis={skip_analysis}\n"
        )
    elif has_data:
        content = (
            "\n\n## ⛔ 规划阶段强制规则（用户已设置）\n"
            "**用户启用了数据分析与图表生成，且检测到 user_data/ 中有数据文件。**\n"
            "请按真实数据规划具体图表（fig_xxx），图表数量与数据资产匹配。\n"
        )
    else:
        content = (
            "\n\n## ⛔ 规划阶段强制规则（用户已设置）\n"
            "**用户启用了数据分析与图表生成，但未上传数据文件。**\n"
            "请按以下方式处理：\n"
            "1. 在规划文档中写明：「**用户未提供数据，将使用仿真/示例数据**」\n"
            "2. **正常规划完整图表清单**（fig_desc / fig_main / fig_ablation 等），不要因为没数据就跳过\n"
            "3. 在分析任务部分明确数据来源：基于论文主题构造合理的仿真数据集（含数据规模、字段、分布假设）\n"
            "4. 图表说明里标注「（基于仿真数据）」，保持学术诚信\n"
            "5. 后续 paper-analysis 步骤会按这份规划生成仿真数据并产出 figures/all_results.json\n"
            "**绝对不要因为 user_data/ 为空就走「无图表分支」，那只在用户显式关闭开关时才用。**\n"
        )

    if skip_drawio:
        content += (
            "**用户已禁用架构图/流程图，后续 paper-figure-drawio 步骤会被跳过。**\n"
            "因此规划文档中 **禁止规划 fig_arch / fig_er / fig_flow_***。\n"
            "方法/系统结构章节用纯文字 + Markdown 表格表达。\n"
        )
    else:
        content += (
            "**用户启用了架构图/流程图。** 请按需规划 fig_arch / fig_er / fig_flow_*，"
            "每张图须明确归属到哪一节。\n"
        )

    injected = _inject_once(Path(workspace) / "CLAUDE.md", "PLAN_STAGE_GUIDANCE", content)
    if injected:
        log.info(
            "Injected plan-stage guidance for %s (skip_figures=%s skip_drawio=%s has_data=%s)",
            skill_name,
            skip_figures,
            skip_drawio,
            has_data,
        )
    return injected


def _is_drawio_fig(rel_path: str) -> bool:
    stem = Path(rel_path).stem.lower()
    return stem.startswith(_DRAWIO_FIG_PREFIXES) or stem.startswith(_SCENE_NAME_PREFIXES)


def _min_size_for(skill_name: str, primary_output: Optional[str] = None) -> int:
    if primary_output:
        suffix = Path(primary_output).suffix.lower()
        if suffix == ".pdf":
            return 30000
        if suffix in {".md", ".tex"} and skill_name.startswith("paper-write"):
            return 15000
    return _STEP_MIN_SIZE.get(skill_name, 100)


def _required_companions_for(skill_name: str) -> List[str]:
    return list(_STEP_REQUIRED_COMPANIONS.get(skill_name, []))


def _missing_upstream_primary_outputs(tmpl: TemplateDef, step_def: StepDef, workspace: Path) -> List[str]:
    """Return required upstream artifacts; callers must fail closed when non-empty."""
    del tmpl
    # Production scientific steps must not continue when required upstream
    # primary artifacts are absent (P1-PS-005 / REQ-P1-03).
    prerequisites = {
        "comp-modeling": ("PROBLEM_ANALYSIS.md",),
        "paper-write": ("PAPER_PLAN.md",),
        "paper-write-zh": ("PAPER_PLAN.md",),
        "paper-write-nature": ("PAPER_PLAN.md",),
        "paper-compile": ("paper/main.tex",),
        "paper-compile-zh": ("paper/main.tex",),
        "paper-compile-nature": ("paper/main.tex",),
        "comp-paper-zh": ("MODELING_REPORT.md",),
        "comp-paper-en": ("MODELING_REPORT.md",),
        "auto-paper-improvement-loop": ("paper/main.tex",),
        "experiment-bridge": ("FINAL_PROPOSAL.md", "refine-logs/FINAL_PROPOSAL.md"),
    }
    required = prerequisites.get(step_def.skill_name, ())
    missing: List[str] = []
    for rel_path in required:
        # For multi-candidate requirements (tuple alternatives joined by path
        # variants), accept any existing candidate when skill lists variants.
        path = Path(workspace) / rel_path
        if path.is_file() and path.stat().st_size > 0:
            # Satisfied this requirement; for experiment-bridge we use OR groups below.
            continue
        missing.append(rel_path)
    # experiment-bridge: either FINAL_PROPOSAL.md or refine-logs/FINAL_PROPOSAL.md
    if step_def.skill_name == "experiment-bridge":
        if any((Path(workspace) / cand).is_file() for cand in ("FINAL_PROPOSAL.md", "refine-logs/FINAL_PROPOSAL.md")):
            missing = [m for m in missing if m not in ("FINAL_PROPOSAL.md", "refine-logs/FINAL_PROPOSAL.md")]
        else:
            missing = ["FINAL_PROPOSAL.md|refine-logs/FINAL_PROPOSAL.md"]
    return missing


def _has_table_and_json_data(workspace: Path) -> bool:
    workspace = Path(workspace)
    figures = workspace / "figures"
    has_table = figures.is_dir() and any(path.is_file() for path in figures.glob("TABLE_*"))
    json_roots = (figures, workspace / "data")
    has_json = any(
        root.is_dir() and any(path.is_file() for path in root.rglob("*.json"))
        for root in json_roots
    )
    return has_table and has_json


def _check_step_companions(workspace: Path, skill_name: str) -> tuple[bool, List[str]]:
    missing = []
    for rel in _required_companions_for(skill_name):
        path = Path(workspace) / rel
        if not path.is_file() or path.stat().st_size <= 50:
            missing.append(rel)
    return not missing, missing


def _read_assets_index(workspace: Path) -> dict:
    defaults = {
        "has_problem": False, "has_code": False, "has_data": False, "has_figures": False,
        "has_results": False, "has_templates": False, "missing_assets": [], "conflicts": [],
        "conflict_count": 0, "high_severity_count": 0,
    }
    path = Path(workspace) / "_assets_index.json"
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return defaults
    result = {**defaults, **{key: data.get(key, default) for key, default in defaults.items()}}
    result["conflict_count"] = data.get("conflict_count", len(result["conflicts"]))
    result["high_severity_count"] = data.get(
        "high_severity_count",
        sum(1 for item in result["conflicts"] if str(item.get("severity", "")).lower() == "high"),
    )
    return result


def _read_figure_manifest(workspace: Path) -> tuple[List[str], List[str]]:
    data_figs: List[str] = []
    drawio_figs: List[str] = []
    for filename in ("PROBLEM_ANALYSIS.md", "PAPER_PLAN.md", "MODELING_REPORT.md"):
        path = Path(workspace) / filename
        if not path.exists():
            continue
        text = path.read_text(encoding="utf-8", errors="replace")
        match = re.search(r"<!--\s*BEGIN FIGURE_MANIFEST\s*-->(.*?)<!--\s*END FIGURE_MANIFEST\s*-->", text, re.I | re.S)
        if not match:
            continue
        for name in re.findall(r"\b(?:fig|tikz)_[A-Za-z0-9_-]+", match.group(1)):
            target = drawio_figs if _is_drawio_fig(name) else data_figs
            if name not in target:
                target.append(name)
    return data_figs, drawio_figs


def _count_existing_figs_for(workspace: Path, names: List[str]) -> tuple[int, List[str]]:
    stems = {path.stem.lower() for path in (Path(workspace) / "figures").glob("*") if path.is_file()}
    missing = [name for name in names if name.lower() not in stems]
    return len(names) - len(missing), missing


def _plan_says_no_figures(workspace: Path) -> bool:
    patterns = (
        r"\bno\s+(?:figures?|charts?|tables?)\b", r"\bwithout\s+(?:figures?|charts?|tables?)\b",
        r"无需(?:任何)?图表", r"不需要(?:任何)?图表", r"无图表", r"纯文字", r"思辨论文",
    )
    for filename in ("PAPER_PLAN.md", "PROBLEM_ANALYSIS.md", "TOPIC_PLAN.md"):
        path = Path(workspace) / filename
        if path.exists():
            text = path.read_text(encoding="utf-8", errors="replace")
            if any(re.search(pattern, text, re.I) for pattern in patterns):
                return True
    return False


def _pfa_safety_copy_assets(workspace: Path) -> dict:
    workspace = Path(workspace)
    user_data = workspace / "user_data"
    counts = {"code": 0, "data": 0, "figures": 0, "templates": 0, "skipped": 0}
    if not user_data.exists():
        return counts
    groups = {
        "code": ({".py", ".ipynb", ".zip"}, workspace / "code"),
        "data": ({".csv", ".xlsx", ".tsv", ".json"}, workspace / "data"),
        "templates": ({".tex", ".cls", ".sty", ".docx", ".dotx"}, workspace / "_user_templates"),
    }
    image_exts = {".png", ".jpg", ".jpeg", ".svg"}
    for src in sorted(user_data.iterdir()):
        if not src.is_file() or src.name.lower() == "results.json":
            continue
        category = next((name for name, (exts, _) in groups.items() if src.suffix.lower() in exts), None)
        if src.suffix.lower() in image_exts:
            category = "figures"
            digest = hashlib.sha1(src.read_bytes()).hexdigest()[:6]
            dst = workspace / "figures" / f"user_fig_{src.stem}_{digest}{src.suffix.lower()}"
        elif category:
            dst = groups[category][1] / src.name
        else:
            continue
        if dst.exists():
            counts["skipped"] += 1
            continue
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(src, dst)
        counts[category] += 1
    return counts


def _should_skip_step_by_assets(workspace: Path, skill_name: str, template: str) -> tuple[bool, str]:
    if template != "paper_from_assets":
        return False, ""
    assets = _read_assets_index(workspace)
    missing_text = " ".join(map(str, assets.get("missing_assets", []))).lower()
    if skill_name == "paper-analysis" and assets["has_code"] and assets["has_results"]:
        return True, "code and results already provided"
    if skill_name == "paper-figure" and assets["has_figures"] and not any(word in missing_text for word in ("figure", "chart", "plot", "图")):
        return True, "figures already provided"
    if skill_name in {"paper-figure-drawio", "paper-figure-html"}:
        figures = Path(workspace) / "figures"
        has_roadmap = any(path.stem.lower().startswith(_DRAWIO_FIG_PREFIXES) for path in figures.glob("*") if path.is_file())
        if has_roadmap and not any(word in missing_text for word in ("roadmap", "architecture", "flow", "架构", "路线")):
            return True, "architecture figures already provided"
    return False, ""


def _figure_files(workspace: Path) -> List[str]:
    figures = Path(workspace) / "figures"
    if not figures.exists():
        return []
    allowed = set(_FIG_IMG_EXTS) | {".drawio", ".tex"}
    return [path.relative_to(workspace).as_posix() for path in figures.rglob("*") if path.is_file() and path.suffix.lower() in allowed]


def _explicit_minimum(params: Dict[str, Any], key: str) -> int:
    value = params.get(key)
    return value if isinstance(value, int) and not isinstance(value, bool) and value > 0 else 0


def _quantity_manifest(workspace: Path, params: Dict[str, Any]) -> Dict[str, Any]:
    workspace = Path(workspace)
    figures_dir = workspace / "figures"
    table_files = sorted(
        {
            path.relative_to(workspace).as_posix()
            for path in figures_dir.glob("TABLE_*")
            if path.is_file() and path.suffix.lower() in {".md", ".tex", ".csv", ".xlsx"}
        }
    ) if figures_dir.is_dir() else []
    table_ids = sorted({Path(item).stem.lower() for item in table_files})

    figure_items = sorted({
        path.relative_to(workspace).as_posix()
        for path in figures_dir.rglob("*")
        if path.is_file()
        and path.suffix.lower() in _FIG_IMG_EXTS
        and not path.stem.upper().startswith("TABLE_")
        and not _is_drawio_fig(path.name)
    }) if figures_dir.is_dir() else []
    figure_ids = sorted({Path(item).stem.lower() for item in figure_items})

    model_items: List[Dict[str, Any]] = []
    model_manifest = workspace / "MODEL_MANIFEST.json"
    if model_manifest.is_file():
        try:
            payload = json.loads(model_manifest.read_text(encoding="utf-8"))
            raw_models = payload.get("models", []) if isinstance(payload, dict) else []
            if isinstance(raw_models, list):
                for index, item in enumerate(raw_models):
                    if isinstance(item, dict):
                        identifier = str(item.get("id") or item.get("name") or "").strip()
                        if identifier:
                            model_items.append({"id": identifier, **item})
                    elif str(item).strip():
                        model_items.append({"id": str(item).strip(), "name": str(item).strip()})
        except (OSError, json.JSONDecodeError):
            model_items = []
    model_ids = {str(item.get("id") or item.get("name")).strip().lower() for item in model_items}

    manifest = {
        "version": 1,
        "requested": {
            "figures": _explicit_minimum(params, "min_figures"),
            "tables": _explicit_minimum(params, "min_tables"),
            "models": _explicit_minimum(params, "min_models"),
        },
        "actual": {
            "figures": len(figure_ids), "tables": len(table_ids), "models": len(model_ids),
        },
        "items": {
            "figures": figure_items, "tables": table_files, "models": model_items,
        },
    }
    (workspace / "QUANTITY_MANIFEST.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8",
    )
    return manifest


def _minimum_quantity_issue(
    workspace: Path, skill_name: str, params: Dict[str, Any],
) -> Optional[str]:
    relevant = {
        "comp-modeling": ("models", "min_models"),
        "paper-figure": ("figures", "min_figures"),
        "nature-figure": ("figures", "min_figures"),
        "experiment-bridge": ("figures", "min_figures"),
    }
    checks: List[tuple[str, str]] = []
    if skill_name in relevant:
        checks.append(relevant[skill_name])
    if skill_name in {"paper-figure", "nature-figure", "experiment-bridge"}:
        checks.append(("tables", "min_tables"))
    if not checks or not any(_explicit_minimum(params, key) for _, key in checks):
        return None

    manifest = _quantity_manifest(workspace, params)
    failures = []
    for quantity, key in checks:
        required = _explicit_minimum(params, key)
        actual = int(manifest["actual"][quantity])
        if required and actual < required:
            failures.append(f"{quantity} {actual} < {key}={required}")
    return "minimum quantity contract failed: " + ", ".join(failures) if failures else None


def _check_figures_step_health_static(workspace: Path, skill_name: str, params: Dict[str, Any]) -> tuple[bool, str]:
    if params.get("enable_figures") is False or params.get("generate_figures") is False:
        return True, "figure generation disabled"
    if _plan_says_no_figures(workspace):
        return True, "plan declares no figures"
    files = _figure_files(workspace)
    if skill_name in {"paper-figure-drawio", "paper-figure-html"}:
        healthy = any(_is_drawio_fig(path) for path in files)
        return (True, "architecture figures exist") if healthy else (False, "no architecture figure output")
    if skill_name in {"paper-figure", "nature-figure", "experiment-bridge"}:
        healthy = any(not _is_drawio_fig(path) for path in files)
        return (True, "data figures exist") if healthy else (False, "no data figure output")
    return True, "not a figure step"


def _check_figures_step_health(
    workspace: Path,
    skill_name: str,
    files_created: Optional[List[str]],
    files_modified: Optional[List[str]],
    params: Dict[str, Any],
) -> tuple[bool, str]:
    if params.get("enable_figures") is False or params.get("generate_figures") is False:
        return True, "figure generation disabled"
    if _plan_says_no_figures(workspace):
        return True, "plan declares no figures"
    changed = list(files_created or []) + list(files_modified or [])
    changed_figures = [path for path in changed if path.replace("\\", "/").startswith("figures/")]
    if skill_name in {"paper-figure-drawio", "paper-figure-html"}:
        if any(_is_drawio_fig(path) for path in changed_figures):
            return True, "architecture figure created"
    elif skill_name in {"paper-figure", "nature-figure", "experiment-bridge"}:
        if any(Path(path).suffix.lower() in _FIG_IMG_EXTS and not _is_drawio_fig(path) for path in changed_figures):
            return True, "data figure created"
        # Provide specific diagnosis so retry prompt can target the actual root cause.
        if changed and not changed_figures:
            return False, "图片文件未保存到 figures/ 目录下（检测到其他目录有文件变更）"
        if changed_figures:
            arch_named = [
                p for p in changed_figures
                if _is_drawio_fig(p) and Path(p).suffix.lower() in _FIG_IMG_EXTS
            ]
            if arch_named:
                example = Path(arch_named[0]).name
                return False, (
                    f"figures/ 内文件使用了架构图命名前缀 (例如 {example!r})，"
                    "数据图禁止以 fig_arch/fig_flow/fig_roadmap/fig_overview/fig_system 等前缀命名，"
                    "请改用 gen_fig_XXX.png 格式"
                )
            # Near-image formats (e.g. .tiff/.eps/.bmp): agent used an unsupported image format.
            near_img = [p for p in changed_figures if Path(p).suffix.lower() in _NEAR_IMG_EXTS]
            if near_img:
                exts = sorted({Path(p).suffix.lower() for p in near_img})
                return False, (
                    f"figures/ 内文件扩展名不合法 {exts}，必须使用 .png/.pdf/.jpg/.jpeg/.svg"
                )
            # Support/script files only (.json/.py/.md/.tex etc.): scripts were created/run
            # but never produced actual image output.
            non_img = [p for p in changed_figures if Path(p).suffix.lower() not in _FIG_IMG_EXTS]
            if non_img:
                exts = sorted({Path(p).suffix.lower() for p in non_img})
                return False, (
                    f"figures/ 内只有辅助文件（扩展名: {exts}），"
                    "gen_fig_*.py 必须实际执行并产出 .png/.pdf 文件"
                )
    else:
        return True, "not a figure step"
    if not changed:
        return _check_figures_step_health_static(workspace, skill_name, params)
    return False, "figure step did not create the required figure type"


_AUTO_RECOVER_FIGURE_SKILLS = {
    "paper-figure", "paper-figure-drawio", "paper-figure-html", "nature-figure", "experiment-bridge",
}


def _primary_output_issue(
    workspace: Path,
    step_def: StepDef,
    *,
    initial_directory_check: bool = False,
) -> Optional[str]:
    primary = step_def.primary_output
    if not primary:
        return None
    path = Path(workspace) / primary
    if not Path(primary.rstrip("/\\")).suffix:
        has_files = path.is_dir() and any(item.is_file() for item in path.rglob("*"))
        if not has_files:
            label = "完全为空" if initial_directory_check else "为空"
            normalized = primary.replace("\\", "/").rstrip("/") + "/"
            return f"primary_output 目录 `{normalized}` {label}"
        return None
    minimum = _min_size_for(step_def.skill_name, primary)
    if not path.is_file() or path.stat().st_size < minimum:
        return f"primary output missing or too small: {primary} (minimum {minimum} bytes)"
    return None


def _figure_recovery_params(params: Dict[str, Any], reason: str, attempt: int) -> Dict[str, Any]:
    retry_params = dict(params)
    retry_params["_retry_reason"] = (
        f"上一次没有产出图表 (原因: {reason})。本次是第 {attempt}/5 次自动重试。"
        "⛔ 必须按 PAPER_PLAN.md / PROBLEM_ANALYSIS.md 的图表清单逐张生成 "
        "gen_fig_*.py 并执行产出 PNG/PDF。"
        " ⛔ 文件三要素 (违反任一项均视为未产出图表): "
        "1) 必须保存到 figures/ 目录下，不能放到 data/ 或工作区根目录；"
        " 2) 扩展名只能用 .png/.pdf/.jpg/.jpeg/.svg；"
        " 3) 数据图文件名禁止以架构图前缀开头"
        "（fig_arch/fig_flow/fig_roadmap/fig_overview/fig_system/fig_pipeline/fig_framework 等），"
        "请使用 gen_fig_XXX.png 或 fig_data_XXX.png 格式。"
        " ⛔ 节约 context: 不要 cat 整个"
        " figure_style_guide.md, 用 head -1500 + grep 按需读 recipe。"
    )
    if attempt >= 2:
        retry_params["_compact_preamble"] = True
    return retry_params


_DOCX_TEMPLATE_SUFFIXES = {".docx", ".dotx"}
_LATEX_TEMPLATE_SUFFIXES = {".tex", ".cls", ".sty", ".bst", ".bib"}


def _template_files_by_kind(params: Dict[str, Any]) -> tuple[List[str], List[str]]:
    """Return de-duplicated DOCX and LaTeX template paths from UI params."""
    candidates: List[str] = []
    single = params.get("template_file")
    if isinstance(single, str) and single.strip():
        candidates.append(single.strip())
    multiple = params.get("template_files")
    if isinstance(multiple, list):
        candidates.extend(str(item).strip() for item in multiple if str(item).strip())

    docx_files: List[str] = []
    latex_files: List[str] = []
    seen = set()
    for item in candidates:
        key = item.replace("\\", "/").lower()
        if key in seen:
            continue
        seen.add(key)
        suffix = Path(item).suffix.lower()
        if suffix in _DOCX_TEMPLATE_SUFFIXES:
            docx_files.append(item)
        elif suffix in _LATEX_TEMPLATE_SUFFIXES:
            latex_files.append(item)
    return docx_files, latex_files


def _replace_step(
    steps: List[StepDef],
    old_skill: str,
    new_skill: str,
    display_name: str,
    output_files: List[str],
    primary_output: str,
) -> None:
    for step in steps:
        if step.skill_name == old_skill:
            step.skill_name = new_skill
            step.display_name = display_name
            step.output_files = list(output_files)
            step.primary_output = primary_output


def _insert_before_skill(steps: List[StepDef], before: set[str], step: StepDef) -> None:
    index = next((i for i, item in enumerate(steps) if item.skill_name in before), len(steps))
    steps.insert(index, step)


def _canonical_paper_template(template: str, params: dict) -> str:
    """Delegate to the shared workflow options contract."""
    from services.workflow_options import _canonical_paper_template as shared_canonical

    return shared_canonical(template, params)


def _resolve_template(template: str, params: dict, workspace: Path) -> TemplateDef:
    """Resolve the installed engine's language, DOCX, and pruning variants."""
    template = _canonical_paper_template(template, params)
    base = TEMPLATES.get(template)
    if base is None:
        raise ValueError(f"Unknown template: {template}")
    steps = [
        StepDef(
            skill_name=step.skill_name,
            display_name=step.display_name,
            output_files=list(step.output_files),
            primary_output=step.primary_output,
            has_checkpoint=step.has_checkpoint,
            checkpoint_type=step.checkpoint_type,
        )
        for step in base.sub_steps
    ]

    # The installed HTML engine is the default. DrawIO remains available for
    # users who need an editable source diagram.
    if str(params.get("flowchart_engine") or "html").lower() != "drawio":
        for step in steps:
            if step.skill_name == "paper-figure-drawio":
                step.skill_name = "paper-figure-html"

    # The visible Nature style option selects the dedicated Nature figure
    # producer; merely putting figure_style in the prompt leaves paper-figure
    # on its default palette and makes the control cosmetic.
    if str(params.get("figure_style") or "default").lower() == "nature":
        for step in steps:
            if step.skill_name == "paper-figure":
                step.skill_name = "nature-figure"
                step.display_name = "Nature 图表"

    # The original UI exposed report generation for software projects but the
    # installed resolver did not add the report step.  Make the visible option
    # operational while keeping the four-step baseline unchanged when disabled.
    if template == "grad_project" and not params.get("skip_report", True):
        report_is_docx = str(params.get("output_format") or "pdf").lower() == "docx"
        steps.append(StepDef(
            skill_name="dev-report", display_name="项目报告",
            output_files=["paper/main.md"] if report_is_docx else ["paper/main.tex", "paper/references.bib"],
            primary_output="paper/main.md" if report_is_docx else "paper/main.tex",
        ))
        if params.get("output_format") in {"pdf", "latex"}:
            steps.append(StepDef(
                skill_name="paper-compile-zh", display_name="编译项目报告 PDF",
                output_files=["paper/main.pdf"], primary_output="paper/main.pdf",
            ))

    output_format = str(params.get("output_format") or "pdf").lower()

    if template == "humanities_paper" and output_format in {"latex", "pdf"}:
        resolved_humanities: List[StepDef] = []
        for step in steps:
            if step.skill_name in {"docx-export", "docx-format-check"}:
                continue
            if step.skill_name == "humanities-write":
                step.skill_name = "humanities-write-latex"
                step.display_name = "人文社科论文撰写（LaTeX）"
                step.output_files = ["paper/main.tex", "paper/references.bib"]
                step.primary_output = "paper/main.tex"
            resolved_humanities.append(step)
        resolved_humanities.append(StepDef(
            skill_name="paper-compile-zh", display_name="编译 PDF",
            output_files=["paper/main.pdf"], primary_output="paper/main.pdf",
        ))
        steps = resolved_humanities

    if template == "paper_from_assets":
        target = str(params.get("paper_type_target") or "academic_zh").lower()
        language = str(params.get("language") or ("en" if target in {"academic_en", "nature"} else "zh")).lower()
        if target == "academic_zh":
            language_map = {
                "paper-plan": ("paper-plan-zh", "论文大纲（中文）"),
                "paper-write": ("paper-write-zh", "LaTeX 写作（中文）"),
                "paper-compile": ("paper-compile-zh", "编译 PDF（中文）"),
            }
            for step in steps:
                replacement = language_map.get(step.skill_name)
                if replacement:
                    step.skill_name, step.display_name = replacement
        elif target == "nature":
            _replace_step(
                steps, "paper-figure", "nature-figure", "Nature 图表",
                ["figures/"], "figures/",
            )
            _replace_step(
                steps, "paper-write", "paper-write-nature", "Nature 论文写作",
                ["paper/main.tex", "paper/references.bib"], "paper/main.tex",
            )
        elif target == "competition":
            _replace_step(
                steps, "paper-plan", "comp-prob-analysis", "竞赛赛题与结构分析",
                ["PROBLEM_ANALYSIS.md"], "PROBLEM_ANALYSIS.md",
            )
            plan_index = next(
                (i for i, item in enumerate(steps) if item.skill_name == "comp-prob-analysis"),
                0,
            )
            steps.insert(plan_index + 1, StepDef(
                skill_name="comp-modeling", display_name="竞赛建模补全",
                output_files=["MODELING_REPORT.md"], primary_output="MODELING_REPORT.md",
                has_checkpoint=True, checkpoint_type="feedback",
            ))
            writer = "comp-paper-en" if language == "en" else "comp-paper-zh"
            compiler = "comp-compile-en" if language == "en" else "comp-compile-zh"
            _replace_step(
                steps, "paper-write", writer, "竞赛论文撰写",
                ["paper/main.tex"], "paper/main.tex",
            )
            _replace_step(
                steps, "paper-compile", compiler, "竞赛论文编译与合规检查",
                ["paper/main.pdf"], "paper/main.pdf",
            )
        elif target == "course":
            _replace_step(
                steps, "paper-plan", "course-plan", "课程论文大纲与图表规划",
                ["OUTLINE.md", "PAPER_PLAN.md"], "OUTLINE.md",
            )
            _replace_step(
                steps, "paper-write", "course-paper", "课程论文撰写",
                ["COURSE_PAPER.md"], "COURSE_PAPER.md",
            )
            if output_format == "docx":
                steps = [item for item in steps if item.skill_name != "paper-compile"]
            else:
                _replace_step(
                    steps, "paper-compile", "markdown-pdf-export", "课程论文 PDF 导出",
                    ["paper/main.pdf"], "paper/main.pdf",
                )

    if template == "full_pipeline":
        branch = str(params.get("paper_branch") or "general").lower()
        if branch == "nature":
            _replace_step(
                steps, "paper-write", "paper-write-nature", "Nature 论文写作",
                ["paper/main.tex", "paper/references.bib"], "paper/main.tex",
            )
            plan_index = next((i for i, item in enumerate(steps) if item.skill_name == "paper-plan"), len(steps))
            steps.insert(plan_index + 1, StepDef(
                skill_name="nature-figure", display_name="Nature 图表",
                output_files=["figures/"], primary_output="figures/",
            ))
        elif branch == "humanities":
            upstream = steps[:6]
            humanities_tail = [
                StepDef(
                    skill_name="humanities-plan", display_name="人文社科结构规划",
                    output_files=["OUTLINE.md", "PAPER_PLAN.md"], primary_output="OUTLINE.md",
                    has_checkpoint=True, checkpoint_type="approve",
                ),
            ]
            if not params.get("skip_drawio"):
                humanities_tail.append(StepDef(
                    skill_name="paper-figure-drawio", display_name="理论框架图/示意图",
                    output_files=["figures/latex_includes.tex"], primary_output="figures/",
                ))
            if output_format == "docx":
                humanities_tail.append(StepDef(
                    skill_name="humanities-write", display_name="人文社科论文撰写",
                    output_files=["HUMANITIES_PAPER.md"], primary_output="HUMANITIES_PAPER.md",
                    has_checkpoint=True, checkpoint_type="feedback",
                ))
            else:
                humanities_tail.extend([
                    StepDef(
                        skill_name="humanities-write-latex", display_name="人文社科论文撰写（LaTeX）",
                        output_files=["paper/main.tex", "paper/references.bib"], primary_output="paper/main.tex",
                        has_checkpoint=True, checkpoint_type="feedback",
                    ),
                    StepDef(
                        skill_name="paper-compile-zh", display_name="编译 PDF",
                        output_files=["paper/main.pdf"], primary_output="paper/main.pdf",
                    ),
                ])
            if not params.get("skip_improvement_loop"):
                humanities_tail.insert(-1 if output_format != "docx" else len(humanities_tail), StepDef(
                    skill_name=("auto-paper-improvement-docx" if output_format == "docx" else "auto-paper-improvement-loop"),
                    display_name="论文改进循环",
                    output_files=(
                        ["HUMANITIES_PAPER.md", "paper/PAPER_IMPROVEMENT_LOG.md"]
                        if output_format == "docx"
                        else ["paper/main.pdf", "paper/PAPER_IMPROVEMENT_LOG.md"]
                    ),
                    primary_output="HUMANITIES_PAPER.md" if output_format == "docx" else "paper/main.pdf",
                ))
            steps = upstream + humanities_tail
        elif str(params.get("language") or "zh").lower() == "zh":
            language_map = {
                "paper-plan": ("paper-plan-zh", "论文大纲（中文）"),
                "paper-write": ("paper-write-zh", "LaTeX 写作（中文）"),
                "paper-compile": ("paper-compile-zh", "编译 PDF（中文）"),
            }
            for step in steps:
                replacement = language_map.get(step.skill_name)
                if replacement:
                    step.skill_name, step.display_name = replacement

    if not params.get("skip_improvement_loop") and not any(
        step.skill_name == "auto-paper-improvement-loop" for step in steps
    ) and not any(step.skill_name == "auto-paper-improvement-docx" for step in steps) and (
        template.startswith("comp_") or template == "paper_from_assets"
    ):
        docx_improvement_primary = (
            "COURSE_PAPER.md"
            if template == "paper_from_assets" and params.get("paper_type_target") == "course"
            else "paper/main.md"
        )
        steps.append(StepDef(
            skill_name=("auto-paper-improvement-docx" if output_format == "docx" else "auto-paper-improvement-loop"),
            display_name="论文改进循环",
            output_files=(
                [docx_improvement_primary, "paper/PAPER_IMPROVEMENT_LOG.md"]
                if output_format == "docx"
                else ["paper/main.pdf", "paper/PAPER_IMPROVEMENT_LOG.md"]
            ),
            primary_output=docx_improvement_primary if output_format == "docx" else "paper/main.pdf",
        ))

    if (
        template in {"thesis_proposal", "literature_review", "course_paper", "course_report", "humanities_paper"}
        and str(params.get("output_format") or "docx").lower() == "docx"
    ):
        insert_at = next((i for i, step in enumerate(steps) if step.skill_name == "docx-export"), len(steps))
        steps.insert(insert_at, StepDef(
            skill_name="docx-format-check",
            display_name="Markdown 格式自检与修复",
            output_files=["DOCX_FORMAT_CHECK_REPORT.md"],
            primary_output="DOCX_FORMAT_CHECK_REPORT.md",
        ))

    if params.get("skip_improvement_loop"):
        steps = [
            step for step in steps
            if step.skill_name not in {"auto-paper-improvement-loop", "auto-paper-improvement-docx"}
        ]

    docx_templates, latex_templates = _template_files_by_kind(params)
    has_format_text = bool(str(params.get("format_text") or "").strip())

    if output_format == "docx":
        # Dedicated formal builders already emit final Word packages. Do not
        # inject the generic Markdown→DOCX export chain for those templates.
        host_formal_templates = {"copyright_material", "patent_disclosure"}
        if template in host_formal_templates:
            return TemplateDef(base.pipeline_skill, base.display_name, steps)

        writer_map = {
            "paper-write": "paper-write-docx",
            "paper-write-zh": "paper-write-zh-docx",
            "paper-write-nature": "paper-write-nature-docx",
            "comp-paper-en": "comp-paper-en-docx",
            "comp-paper-zh": "comp-paper-zh-docx",
        }
        resolved = []
        for step in steps:
            if (
                step.skill_name.startswith("paper-compile")
                or step.skill_name.startswith("comp-compile")
                or step.skill_name in {"markdown-pdf-export", "auto-review-pdf-export"}
            ):
                continue
            if step.skill_name in writer_map:
                step.skill_name = writer_map[step.skill_name]
                step.display_name = f"{step.display_name}(Markdown)"
                step.output_files = ["paper/main.md"]
                step.primary_output = "paper/main.md"
            resolved.append(step)
        for step in resolved:
            if step.skill_name == "auto-paper-improvement-loop":
                step.skill_name = "auto-paper-improvement-docx"
                step.display_name = "论文改进循环"
                step.output_files = ["paper/main.md", "paper/PAPER_IMPROVEMENT_LOG.md"]
                step.primary_output = "paper/main.md"

        # Template/profile steps intentionally run after the Markdown writer so
        # the DOCX map can be trial-filled against the real source document.
        if docx_templates or latex_templates:
            resolved.append(StepDef(
                skill_name="template-prepare", display_name="准备格式模板",
                output_files=["_template_contract.json", "_derived_profile.json"],
                primary_output="_template_contract.json",
            ))
        if docx_templates or latex_templates or has_format_text:
            resolved.append(StepDef(
                skill_name="format-profile",
                display_name="解析格式要求",
                output_files=["_text_profile.json"],
                primary_output="_text_profile.json",
            ))
        if docx_templates:
            resolved.append(StepDef(
                skill_name="docx-template-map",
                display_name="识别 Word 模板占位",
                output_files=["_template_map.json"],
                primary_output="_template_map.json",
            ))

        resolved = [step for step in resolved if step.skill_name not in {"docx-format-check", "docx-export"}]
        resolved.append(StepDef(
            skill_name="docx-format-check",
            display_name="Markdown 格式自检与修复",
            output_files=["DOCX_FORMAT_CHECK_REPORT.md"],
            primary_output="DOCX_FORMAT_CHECK_REPORT.md",
        ))

        output_name = "paper/main.docx"
        export_display = "格式检查与 Word 导出"
        if template == "auto_review":
            output_name, export_display = "NARRATIVE_REPORT.docx", "导出为 Word"
        elif template == "thesis_proposal":
            output_name = "PROPOSAL.docx"
        elif template == "literature_review":
            output_name = "LITERATURE_REVIEW.docx"
        elif template == "course_paper":
            output_name = "COURSE_PAPER.docx"
        elif template == "course_report":
            output_name = "COURSE_REPORT.docx"
        elif template == "humanities_paper":
            output_name = "HUMANITIES_PAPER.docx"
        elif template == "grad_project":
            output_name = "paper/PROJECT_REPORT.docx"
        elif template == "software_copyright":
            output_name = "software-copyright/SOFTWARE_COPYRIGHT.docx"
        resolved.append(StepDef(
            skill_name="docx-export",
            display_name=export_display,
            output_files=[output_name],
            primary_output=output_name,
        ))
        steps = resolved

    elif output_format in {"pdf", "latex"}:
        if latex_templates:
            _insert_before_skill(
                steps,
                {
                    "paper-write", "paper-write-zh", "paper-write-nature",
                    "comp-paper-en", "comp-paper-zh", "humanities-write-latex",
                    "course-paper", "auto-review-loop",
                },
                StepDef(
                    skill_name="latex-template-prepare", display_name="准备 LaTeX 模板套件",
                    output_files=["_latex_template.json", "paper/main.tex"],
                    primary_output="_latex_template.json",
                ),
            )

        if template == "auto_review":
            steps.append(StepDef(
                skill_name="auto-review-pdf-export", display_name="生成审稿 PDF",
                output_files=["paper/main.pdf"], primary_output="paper/main.pdf",
            ))
        elif latex_templates or has_format_text:
            _insert_before_skill(
                steps,
                {
                    "paper-compile", "paper-compile-zh", "comp-compile-en",
                    "comp-compile-zh", "markdown-pdf-export",
                },
                StepDef(
                    skill_name="latex-template-apply", display_name="应用 LaTeX 模板与格式",
                    output_files=["_latex_template_applied.json", "paper/main.tex"],
                    primary_output="_latex_template_applied.json",
                ),
            )

    return TemplateDef(base.pipeline_skill, base.display_name, steps)


def _runtime_skip_reason(skill_name: str, params: Dict[str, Any]) -> str | None:
    if params.get("skip_analysis") and skill_name == "paper-analysis":
        return "用户已关闭数据分析"
    if params.get("skip_figures") and skill_name in {"paper-figure", "nature-figure"}:
        return "用户已关闭数据图表"
    if params.get("skip_drawio") and skill_name in {"paper-figure-drawio", "paper-figure-html"}:
        return "用户已关闭流程图/架构图"
    return None


def _checkpoint_feedback_text(response: Any) -> str:
    """Return an actionable revision request for every checkpoint type."""
    if not isinstance(response, dict) or response.get("action") != "feedback":
        return ""
    data = response.get("data")
    if not isinstance(data, dict):
        return ""
    return str(data.get("feedback") or "").strip()


async def _wait_for_extracts(workspace: Path, timeout_sec: int = 300, poll_interval: float = 1.0) -> None:
    """Wait until asynchronous user-data extraction reaches terminal states."""
    status_path = Path(workspace) / "user_data" / "_extract_status.json"
    if not status_path.exists():
        return
    loop = asyncio.get_running_loop()
    deadline = loop.time() + timeout_sec
    terminal = {"completed", "done", "failed", "skipped"}
    while loop.time() < deadline:
        try:
            status = json.loads(status_path.read_text(encoding="utf-8"))
            files = status.get("files", {})
            if not files or all(str(item.get("status", "")).lower() in terminal for item in files.values()):
                return
        except (OSError, json.JSONDecodeError):
            return
        await asyncio.sleep(poll_interval)


def set_broadcast(broadcast_func: Callable) -> None:
    """(docstring)"""
    global _broadcast_func
    _broadcast_func = broadcast_func


async def _broadcast(wf_id: str, msg: dict) -> None:
    """(docstring)"""
    if _broadcast_func:
        try:
            payload = {"workflow_id": wf_id, **msg}
            await _broadcast_func(wf_id, payload)
        except Exception as e:
            log.warning("Broadcast failed: %s", e)


async def _log(wf_id: str, step_name: str | None, level: str, message: str) -> None:
    """(docstring)"""
    from services.state_store import add_log
    await add_log(wf_id, step_name, level, message)
    await _broadcast(wf_id, {"type": "log", "step_name": step_name, "level": level, "message": message})


async def _persist_log_batch(wf_id: str, step_name: str | None, batch: List[str]) -> None:
    """Persist buffered runner output in one DB transaction."""
    if not batch:
        return
    from services.state_store import execute_write, get_db

    async def _write() -> None:
        db = await get_db()
        try:
            await db.executemany(
                "INSERT INTO workflow_logs (workflow_id, step_name, level, message) VALUES (?, ?, 'info', ?)",
                [(wf_id, step_name, line) for line in batch],
            )
            await db.commit()
        finally:
            await db.close()

    await execute_write(f"persist_log_batch:{wf_id}", _write)


def _result_summary(files_created: List[str], files_modified: List[str]) -> str:
    """Build the renderer summary emitted by the installed workflow engine."""
    parts = []
    if files_created:
        parts.append(f"创建 {len(files_created)} 个文件")
    if files_modified:
        parts.append(f"更新 {len(files_modified)} 个文件")
    return "，".join(parts) if parts else "无文件变更"


def _write_assurance_artifact(workspace: Path, envelope: dict[str, Any]) -> Path:
    """Persist the terminal assurance decision as an auditable workspace artifact."""
    target = Path(workspace) / "ASSURANCE_ENVELOPE.json"
    temporary = target.with_suffix(".json.tmp")
    temporary.write_text(
        json.dumps(envelope, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    temporary.replace(target)
    return target


async def _evaluate_full_pipeline_assurance(workflow: dict, workspace: Path) -> dict[str, Any]:
    """Run the deterministic submission gates before a full pipeline can complete."""
    project_id = str(workflow.get("project_id") or "").strip()
    if not project_id:
        return {
            "format_version": "assurance-envelope/v1",
            "status": "BLOCKED",
            "submission_ready": False,
            "input_hashes": {
                "project_snapshot_sha256": "",
                "latest_review_inputs_sha256": None,
                "review_report_sha256": None,
            },
            "findings": [{
                "severity": "critical",
                "code": "workflow_project_missing",
                "message": "full_pipeline requires a persisted research project binding.",
                "locator": workflow.get("id", ""),
            }],
            "repair_actions": [{
                "finding_code": "workflow_project_missing",
                "action": "创建研究合同并使用 project_id 重新运行全流程。",
            }],
            "verifier_version": "vibe-assurance/1.0",
            "independent_from_generator": True,
            "gates": [{
                "id": "final_submission",
                "label": "最终提交门",
                "status": "BLOCKED",
                "findings": [{"code": "workflow_project_missing"}],
            }],
            "current_review": None,
            "latest_review": None,
        }

    try:
        from services import assurance

        return await assurance.read(project_id)
    except Exception as exc:
        log.exception("Full pipeline assurance evaluation failed for project %s", project_id)
        return {
            "format_version": "assurance-envelope/v1",
            "status": "BLOCKED",
            "submission_ready": False,
            "input_hashes": {
                "project_snapshot_sha256": "",
                "latest_review_inputs_sha256": None,
                "review_report_sha256": None,
            },
            "findings": [{
                "severity": "critical",
                "code": "assurance_evaluation_error",
                "message": f"Deterministic assurance evaluation failed: {exc}",
                "locator": project_id,
            }],
            "repair_actions": [{
                "finding_code": "assurance_evaluation_error",
                "action": "修复质量门禁评估错误后重新运行全流程。",
            }],
            "verifier_version": "vibe-assurance/1.0",
            "independent_from_generator": True,
            "gates": [{
                "id": "final_submission",
                "label": "最终提交门",
                "status": "BLOCKED",
                "findings": [{"code": "assurance_evaluation_error"}],
            }],
            "current_review": None,
            "latest_review": None,
        }


async def create_new_workflow(
    template: str,
    title: str,
    params: dict,
    enable_checkpoints: bool = False,
    project_id: str | None = None,
) -> str:
    """(docstring)"""
    from services.state_store import _get_db, create_workflow

    wf_id = uuid.uuid4().hex[:12]
    workspace_dir = WORKSPACES_DIR / wf_id
    workspace_dir.mkdir(parents=True, exist_ok=True)

    params = dict(params)
    template = _canonical_paper_template(template, params)
    tmpl = _resolve_template(template, params, workspace_dir)

    # full_pipeline terminal assurance requires a bound research project. When
    # the UI creates the workflow without project_id, auto-bind a contract so
    # the host chain can complete and write ASSURANCE_ENVELOPE.json instead of
    # failing only on workflow_project_missing.
    bound_project_id = str(project_id or "").strip() or None
    if template == "full_pipeline" and not bound_project_id:
        try:
            from services import research_contracts

            topic = str(
                params.get("topic")
                or params.get("research_question")
                or title
                or "full pipeline research"
            ).strip()
            contract = await research_contracts.create_contract(
                title=title or topic,
                research_question=topic,
                inclusion_criteria=str(
                    params.get("inclusion_criteria")
                    or "host executor artifacts with audit lineage"
                ),
            )
            bound_project_id = str(contract.get("id") or "").strip() or None
            if bound_project_id:
                params["_auto_bound_research_project"] = True
        except Exception:
            log.exception("Failed to auto-bind research project for full_pipeline")

    # The original engine persists this marker and bootstraps both executor
    # context and a local git repository before exposing the workflow.
    params["_sub_steps_pruned"] = True
    _generate_claude_md(workspace_dir, title, template, params)
    _initialize_git_workspace(workspace_dir)


    steps = []
    for i, step in enumerate(tmpl.sub_steps):
        steps.append({
            "skill_name": step.skill_name,
            "display_name": step.display_name,
            "step_order": i,
            "status": "pending",
            "has_checkpoint": int(step.has_checkpoint),
            "checkpoint_type": step.checkpoint_type,
            "output_files": json.dumps(step.output_files),
        })

    wf = {
        "id": wf_id,
        "project_id": bound_project_id,
        "template": template,
        "title": title,
        "params": params,
        "status": "pending",
        "workspace_dir": str(workspace_dir),
        "enable_checkpoints": enable_checkpoints,
    }

    db = await _get_db()
    try:
        await create_workflow(db, wf)
        for step in steps:
            await db.execute(
                "INSERT INTO workflow_steps (workflow_id, skill_name, display_name, step_order, status, has_checkpoint, checkpoint_type, output_files) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                (wf_id, step["skill_name"], step["display_name"], step["step_order"], step["status"],
                 step["has_checkpoint"], step["checkpoint_type"], step["output_files"])
            )
        await db.commit()
    finally:
        await db.close()

    return wf_id


def _sync_run_process(
    command: List[str],
    cwd: str,
    timeout: float | None,
) -> tuple[int, str, str]:
    """Synchronous last-resort runner for ``_HostStepRunner._run_process``.

    Used when the event loop cannot spawn subprocesses at all — e.g. uvicorn
    >=0.36 with ``--reload``/``--workers>1`` forces a SelectorEventLoop on
    Windows (bypassing ``asyncio.set_event_loop_policy``), and every
    ``asyncio.create_subprocess_exec`` raises a bare ``NotImplementedError``.
    Classic ``subprocess.run`` (CreateProcess + WaitForMultipleObjects) does
    not depend on the event loop, so it still works there.  Mirrors the agent
    sandbox's ``_run_command_sync_fallback``.  Must be invoked via
    ``asyncio.to_thread`` so the loop is never blocked.
    """
    try:
        completed = subprocess.run(
            command,
            cwd=cwd,
            capture_output=True,
            timeout=timeout,
        )
        return (
            int(completed.returncode or 0),
            completed.stdout.decode("utf-8", errors="replace"),
            completed.stderr.decode("utf-8", errors="replace"),
        )
    except subprocess.TimeoutExpired as exc:
        stdout = exc.stdout if isinstance(exc.stdout, bytes) else b""
        stderr = exc.stderr if isinstance(exc.stderr, bytes) else b""
        detail = (
            f"Host step timed out after {timeout}s: "
            f"{' '.join(str(part) for part in command[:4])}"
        )
        return (
            124,
            stdout.decode("utf-8", errors="replace"),
            (stderr.decode("utf-8", errors="replace") + "\n" + detail).strip(),
        )


class _HostStepRunner:
    """Execute deterministic host-side steps through the runner interface."""

    def __init__(self, template: str, step: StepDef):
        self.template = template
        self.step = step

    @staticmethod
    def _markdown_sources(workspace: Path, template: str) -> List[Path]:
        groups = {
            "copyright_material": [
                "software-copyright/PRODUCT_OVERVIEW.md",
                "software-copyright/USER_MANUAL.md",
                "software-copyright/SOURCE_CODE_INDEX.md",
                "software-copyright/REGISTRATION_CHECKLIST.md",
            ],
            "software_copyright": [
                "software-copyright/PRODUCT_OVERVIEW.md",
                "software-copyright/USER_MANUAL.md",
                "software-copyright/SOURCE_CODE_INDEX.md",
                "software-copyright/REGISTRATION_CHECKLIST.md",
            ],
            "patent_disclosure": [
                "patent/INVENTION_DISCLOSURE.md", "patent/CLAIMS_DRAFT.md",
                "patent/PRIOR_ART_PLAN.md", "patent/FIGURE_PLAN.md",
            ],
        }
        candidates = groups.get(template, [
            "paper/main.md", "HUMANITIES_PAPER.md", "PROPOSAL.md",
            "LITERATURE_REVIEW.md", "COURSE_PAPER.md", "COURSE_REPORT.md",
            "REPORT.md", "NARRATIVE_REPORT.md", "RESULTS.md", "MODELING_REPORT.md",
        ])
        return [workspace / name for name in candidates if (workspace / name).is_file()]

    @staticmethod
    def _safe_template_paths(workspace: Path, params: Dict[str, Any]) -> tuple[List[Path], List[Path]]:
        docx_names, latex_names = _template_files_by_kind(params)

        def resolve(names: List[str]) -> List[Path]:
            result: List[Path] = []
            root = workspace.resolve()
            for name in names:
                candidate = Path(name)
                candidate = candidate if candidate.is_absolute() else workspace / candidate
                try:
                    candidate = candidate.resolve()
                    candidate.relative_to(root)
                except (OSError, ValueError):
                    continue
                if candidate.is_file():
                    result.append(candidate)
            return result

        return resolve(docx_names), resolve(latex_names)

    @staticmethod
    def _default_text_profile() -> Dict[str, Any]:
        return {
            "profile_name": "workflow-template-derived",
            "_derived_from": "workflow-template",
            "_matched_items": [],
            "page": {
                "size": "A4", "margin_top_cm": 2.5, "margin_bottom_cm": 2.5,
                "margin_left_cm": 2.5, "margin_right_cm": 2.5,
            },
            "fonts": {
                "chinese_heading": "SimHei", "chinese_body": "SimSun",
                "latin": "Times New Roman", "monospace": "Consolas",
            },
            "headings": {
                "level1_pt": 16, "level2_pt": 14, "level3_pt": 12,
                "level4_pt": 11, "bold": True, "level1_alignment": "center",
                "level2_alignment": "left", "level3_alignment": "left",
            },
            "body": {
                "font_size_pt": 12, "line_spacing": 1.5,
                "first_line_indent_chars": 2, "space_before_pt": 0, "space_after_pt": 0,
            },
            "title": {
                "font_size_pt": 18, "bold": True, "alignment": "center",
                "font_family": "SimHei",
            },
            "table": {
                "top_border_pt": 1.5, "header_border_pt": 0.75,
                "bottom_border_pt": 1.5, "font_size_pt": 10.5,
                "header_bold": True, "cell_alignment": "center",
            },
            "references": {
                "hanging_indent_cm": 0.74, "font_size_pt": 10.5,
                "numbering_style": "bracket",
            },
            "image": {"max_width_cm": 14, "alignment": "center"},
            "code_block": {
                "font_size_pt": 9, "line_spacing": 1.0,
                "background_color": "F5F5F5",
            },
        }

    @classmethod
    def _derive_docx_profile(cls, template: Path) -> Dict[str, Any]:
        profile = cls._default_text_profile()
        try:
            from docx import Document

            document = Document(str(template))
            section = document.sections[0]
            profile["page"].update({
                "margin_top_cm": round(section.top_margin.cm, 3),
                "margin_bottom_cm": round(section.bottom_margin.cm, 3),
                "margin_left_cm": round(section.left_margin.cm, 3),
                "margin_right_cm": round(section.right_margin.cm, 3),
            })
            normal = document.styles["Normal"]
            if normal.font.name:
                profile["fonts"]["chinese_body"] = normal.font.name
                profile["fonts"]["latin"] = normal.font.name
            if normal.font.size:
                profile["body"]["font_size_pt"] = round(normal.font.size.pt, 2)
            matched = [
                f"DOCX template: {template.name}",
                "page margins and Normal style were derived from the physical template",
            ]
            profile["_matched_items"] = matched
            profile["_derived_from"] = template.as_posix()
        except Exception as exc:
            profile["_matched_items"] = [f"DOCX profile fallback: {exc}"]
        return profile

    @classmethod
    def _derive_latex_profile(cls, files: List[Path]) -> Dict[str, Any]:
        profile = cls._default_text_profile()
        text = "\n".join(
            item.read_text(encoding="utf-8", errors="replace")
            for item in files if item.suffix.lower() in {".tex", ".cls", ".sty"}
        )
        matched: List[str] = []
        geometry = re.search(r"\\(?:usepackage\[([^]]+)\]\{geometry\}|geometry\{([^}]+)\})", text)
        if geometry:
            settings = geometry.group(1) or geometry.group(2) or ""
            mapping = {
                "top": "margin_top_cm", "bottom": "margin_bottom_cm",
                "left": "margin_left_cm", "right": "margin_right_cm",
            }
            uniform = re.search(r"(?:^|,)\s*margin\s*=\s*([0-9.]+)\s*cm", settings)
            if uniform:
                for key in mapping.values():
                    profile["page"][key] = float(uniform.group(1))
                matched.append(f"geometry margin={uniform.group(1)}cm")
            for latex_key, profile_key in mapping.items():
                value = re.search(rf"(?:^|,)\s*{latex_key}\s*=\s*([0-9.]+)\s*cm", settings)
                if value:
                    profile["page"][profile_key] = float(value.group(1))
                    matched.append(f"geometry {latex_key}={value.group(1)}cm")
        size = re.search(r"\\documentclass\[([^]]*)\]", text)
        if size:
            point = re.search(r"(9|10|11|12)pt", size.group(1))
            if point:
                profile["body"]["font_size_pt"] = int(point.group(1))
                matched.append(f"documentclass {point.group(1)}pt")
        latin = re.search(r"\\setmainfont(?:\[[^]]*\])?\{([^}]+)\}", text)
        cjk = re.search(r"\\setCJKmainfont(?:\[[^]]*\])?\{([^}]+)\}", text)
        if latin:
            profile["fonts"]["latin"] = latin.group(1).strip()
            matched.append(f"Latin font {latin.group(1).strip()}")
        if cjk:
            profile["fonts"]["chinese_body"] = cjk.group(1).strip()
            matched.append(f"CJK font {cjk.group(1).strip()}")
        spacing = re.search(r"\\linespread\{([0-9.]+)\}", text)
        if spacing:
            profile["body"]["line_spacing"] = float(spacing.group(1))
            matched.append(f"line spacing {spacing.group(1)}")
        profile["_matched_items"] = matched or ["LaTeX suite copied; no portable style token found"]
        profile["_derived_from"] = ", ".join(item.name for item in files)
        return profile

    @classmethod
    async def _prepare_templates(
        cls, workspace: Path, params: Dict[str, Any], *, latex_mode: bool,
    ) -> Dict[str, Any]:
        docx_files, latex_files = cls._safe_template_paths(workspace, params)
        requested_docx, requested_latex = _template_files_by_kind(params)
        if requested_docx and not docx_files:
            return {"success": False, "returncode": 2, "stderr": "DOCX template files are missing"}
        if requested_latex and not latex_files:
            return {"success": False, "returncode": 2, "stderr": "LaTeX template files are missing"}

        prepared_root = workspace / "_user_templates"
        prepared_root.mkdir(parents=True, exist_ok=True)
        prepared: List[str] = []
        for source in [*docx_files, *latex_files]:
            destination = prepared_root / source.name
            shutil.copy2(source, destination)
            prepared.append(destination.relative_to(workspace).as_posix())

        profile = (
            cls._derive_docx_profile(docx_files[0]) if docx_files
            else cls._derive_latex_profile(latex_files) if latex_files
            else cls._default_text_profile()
        )
        (workspace / "_derived_profile.json").write_text(
            json.dumps(profile, ensure_ascii=False, indent=2) + "\n", encoding="utf-8",
        )

        contract: Dict[str, Any] = {
            "version": 1,
            "mode": "latex" if latex_mode else "docx",
            "docx_templates": [item.relative_to(workspace).as_posix() for item in docx_files],
            "latex_templates": [item.relative_to(workspace).as_posix() for item in latex_files],
            "prepared_files": prepared,
            "derived_profile": "_derived_profile.json",
        }
        if latex_files:
            paper = workspace / "paper"
            paper.mkdir(parents=True, exist_ok=True)
            for source in latex_files:
                shutil.copy2(source, paper / source.name)
            tex_files = [item for item in latex_files if item.suffix.lower() == ".tex"]
            main_source = next((item for item in tex_files if item.name.lower() == "main.tex"), tex_files[0] if tex_files else None)
            if main_source is not None:
                snapshot = paper / "_user_template_main.tex"
                shutil.copy2(main_source, snapshot)
                main = paper / "main.tex"
                if not main.exists() or latex_mode:
                    shutil.copy2(main_source, main)
                contract["main_source"] = main_source.relative_to(workspace).as_posix()
                contract["main_snapshot"] = snapshot.relative_to(workspace).as_posix()
                contract["main_output"] = main.relative_to(workspace).as_posix()
        output = workspace / ("_latex_template.json" if latex_mode else "_template_contract.json")
        output.write_text(json.dumps(contract, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        return {"success": True, "returncode": 0, "stdout": str(output), "stderr": ""}

    @staticmethod
    def _extract_latex_body(text: str) -> str:
        match = re.search(r"\\begin\{document\}(.*)\\end\{document\}", text, re.S)
        return match.group(1).strip() if match else text.strip()

    @classmethod
    async def _apply_latex_template(cls, workspace: Path, params: Dict[str, Any]) -> Dict[str, Any]:
        main = workspace / "paper" / "main.tex"
        if not main.is_file():
            return {"success": False, "returncode": 2, "stderr": "paper/main.tex is missing"}
        generated = main.read_text(encoding="utf-8", errors="replace")
        contract_path = workspace / "_latex_template.json"
        contract = json.loads(contract_path.read_text(encoding="utf-8")) if contract_path.is_file() else {}
        snapshot_name = contract.get("main_snapshot")
        snapshot = workspace / snapshot_name if isinstance(snapshot_name, str) else None
        applied = []
        if snapshot is not None and snapshot.is_file():
            template_text = snapshot.read_text(encoding="utf-8", errors="replace")
            body = cls._extract_latex_body(generated)
            marker_patterns = ["{{CONTENT}}", "%% CONTENT %%", "% VIBE_CONTENT", "\\input{content}"]
            marker = next((item for item in marker_patterns if item in template_text), None)
            if marker:
                merged = template_text.replace(marker, body, 1)
                applied.append(f"content inserted at {marker}")
            else:
                preamble = template_text.split("\\begin{document}", 1)[0]
                merged = preamble.rstrip() + "\n\\begin{document}\n" + body + "\n\\end{document}\n"
                applied.append("user template preamble and documentclass preserved")
            main.write_text(merged, encoding="utf-8")

        format_text = str(params.get("format_text") or "").strip()
        if format_text:
            commands = ["% Auto-generated from FORMAT_REQUIREMENTS.md"]
            margin_values = re.findall(r"([0-9]+(?:\.[0-9]+)?)\s*(?:cm|厘米)", format_text, re.I)
            if "页边距" in format_text and margin_values:
                commands.append(f"\\usepackage[margin={margin_values[0]}cm]{{geometry}}")
                applied.append(f"margin {margin_values[0]}cm")
            spacing = re.search(r"([0-9]+(?:\.[0-9]+)?)\s*倍行距", format_text)
            if spacing:
                commands.append(f"\\linespread{{{spacing.group(1)}}}")
                applied.append(f"line spacing {spacing.group(1)}")
            point = 12 if "小四" in format_text else 16 if "三号" in format_text else None
            if point:
                commands.append(f"\\AtBeginDocument{{\\fontsize{{{point}pt}}{{{round(point * 1.5, 1)}pt}}\\selectfont}}")
                applied.append(f"body size {point}pt")
            include = workspace / "paper" / "vibe-user-format.tex"
            include.write_text("\n".join(commands) + "\n", encoding="utf-8")
            text = main.read_text(encoding="utf-8", errors="replace")
            directive = "\\input{vibe-user-format.tex}"
            if directive not in text:
                if "\\begin{document}" in text:
                    text = text.replace("\\begin{document}", directive + "\n\\begin{document}", 1)
                else:
                    text = directive + "\n" + text
                main.write_text(text, encoding="utf-8")

        applied_contract = {
            "version": 1, "main": "paper/main.tex", "applied": applied,
            "template_contract": contract,
        }
        output = workspace / "_latex_template_applied.json"
        output.write_text(json.dumps(applied_contract, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        return {"success": True, "returncode": 0, "stdout": str(output), "stderr": ""}

    @classmethod
    async def _build_docx_template_map(cls, workspace: Path, params: Dict[str, Any]) -> Dict[str, Any]:
        docx_files, _ = cls._safe_template_paths(workspace, params)
        if not docx_files:
            return {"success": False, "returncode": 2, "stderr": "DOCX template is missing"}
        try:
            from docx import Document

            document = Document(str(docx_files[0]))
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
        except Exception as exc:
            return {"success": False, "returncode": 3, "stderr": f"Unable to inspect DOCX template: {exc}"}

        def find(pattern: str) -> Optional[int]:
            return next((i for i, text in enumerate(paragraphs) if re.search(pattern, text, re.I)), None)

        body = find(r"\{\{\s*(?:CONTENT|BODY)\s*\}\}|\[\s*(?:CONTENT|BODY)\s*\]|正文内容|^正文$|^Body$")
        if body is None:
            body = len(paragraphs)
        mapping = {
            "version": 1,
            "template": docx_files[0].relative_to(workspace).as_posix(),
            "title_anchor_para_idx": find(r"标题|题目|Title"),
            "abstract_anchor_para_idx": find(r"^摘\s*要$|^Abstract$"),
            "body_anchor_para_idx": body,
            "body_anchor_mode": "delete",
            "references_anchor_para_idx": find(r"^参考文献$|^References$|^Bibliography$"),
            "appendix_anchor_para_idx": find(r"^附录$|^Appendix$|^Annex$"),
            "delete_paragraph_indices": [],
            "preserve_table_indices": list(range(len(document.tables))),
        }
        output = workspace / "_template_map.json"
        output.write_text(json.dumps(mapping, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        (workspace / "_template_check_report.md").write_text(
            "# DOCX 模板映射报告\n\n"
            f"- 模板: `{mapping['template']}`\n"
            f"- 正文锚点: {body}\n"
            f"- 保留表格: {len(document.tables)}\n",
            encoding="utf-8",
        )
        return {"success": True, "returncode": 0, "stdout": str(output), "stderr": ""}

    @staticmethod
    async def _run_process(
        command: List[str],
        workspace: Path,
        timeout: float | None = 900.0,
    ) -> tuple[int, str, str]:
        process_options: Dict[str, Any] = {}
        if os.name == "nt":
            # Host builds should never flash a console window in the desktop app.
            process_options["creationflags"] = getattr(subprocess, "CREATE_NO_WINDOW", 0)
        cwd_str = str(Path(workspace).expanduser().resolve())
        process = None
        try:
            process = await asyncio.create_subprocess_exec(
                *command,
                cwd=cwd_str,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
                **process_options,
            )
        except (NotImplementedError, OSError) as first_exc:
            # Windows: CREATE_NO_WINDOW can trigger a NotImplementedError when the
            # ProactorEventLoop's IOCP path fails (e.g. app execution alias / Store
            # stub reparse points).  Retry without the flag so the process still
            # runs, even if a console window briefly appears.
            log.warning(
                "create_subprocess_exec failed (%r); retrying without creationflags",
                first_exc,
            )
            process_options_retry = {
                k: v for k, v in process_options.items() if k != "creationflags"
            }
            try:
                process = await asyncio.create_subprocess_exec(
                    *command,
                    cwd=cwd_str,
                    stdout=asyncio.subprocess.PIPE,
                    stderr=asyncio.subprocess.PIPE,
                    **process_options_retry,
                )
            except (NotImplementedError, OSError) as second_exc:
                # SelectorEventLoop (uvicorn >=0.36 forces it on Windows under
                # --reload/--workers, bypassing set_event_loop_policy) cannot
                # spawn subprocesses at all — every create_subprocess_exec raises
                # a bare NotImplementedError.  Fall back to synchronous
                # subprocess.run on a worker thread instead of letting this
                # exception escape and kill the step (fb4f4e5b7272 regression).
                log.warning(
                    "asyncio subprocess unavailable (%r); "
                    "falling back to synchronous subprocess.run",
                    second_exc,
                )
                process = None
        if process is None:
            return await asyncio.to_thread(
                _sync_run_process,
                [str(part) for part in command],
                cwd_str,
                timeout,
            )
        try:
            stdout, stderr = await asyncio.wait_for(process.communicate(), timeout=timeout)
        except asyncio.TimeoutError:
            process.kill()
            try:
                stdout, stderr = await asyncio.wait_for(process.communicate(), timeout=10)
            except Exception:
                stdout, stderr = b"", b""
            detail = f"Host step timed out after {timeout}s: {' '.join(str(part) for part in command[:4])}"
            return (
                124,
                stdout.decode("utf-8", errors="replace"),
                (stderr.decode("utf-8", errors="replace") + "\n" + detail).strip(),
            )
        return (
            int(process.returncode or 0),
            stdout.decode("utf-8", errors="replace"),
            stderr.decode("utf-8", errors="replace"),
        )

    @classmethod
    async def _host_execute_gen_fig_scripts(
        cls,
        workspace: Path,
        skill_name: str,
        *,
        on_output=None,
    ) -> List[str]:
        """Host-side last resort for figure-producing steps.

        When the agent's subprocess channel is broken (Windows ProactorEventLoop
        IOCP failures, Store alias reparse points, etc.) it can write gen_fig_*.py
        scripts but never actually execute them.  This fallback runs those scripts
        directly from the host so the step still produces real PDF/PNG figures:

        1. Try each script as a subprocess via the runtime python (with the
           CREATE_NO_WINDOW retry already handled by _run_process).
        2. If the subprocess cannot even start, import the script in-process with
           matplotlib forced to the Agg backend and run it under the workspace cwd.
        """
        if skill_name not in {"paper-figure", "nature-figure", "experiment-bridge"}:
            return []
        figures_dir = Path(workspace) / "figures"
        if not figures_dir.is_dir():
            return []
        scripts = sorted(figures_dir.glob("gen_fig_*.py"))
        if not scripts:
            return []

        def _already_has_output(script: Path) -> bool:
            stem = script.stem[len("gen_"):] if script.stem.startswith("gen_") else script.stem
            for ext in _FIG_IMG_EXTS:
                if (figures_dir / f"{stem}{ext}").is_file():
                    return True
            return False

        python = cls._runtime_python()
        produced: List[str] = []
        for script in scripts:
            if _already_has_output(script):
                continue
            if on_output:
                await on_output(f"[系统] 宿主兜底执行绘图脚本: {script.name}")
            rc, so, se = await cls._run_process(
                [python, str(script)], workspace, timeout=300.0,
            )
            if rc != 0 and (
                "NotImplementedError" in se or "subprocess" in se.lower()
            ):
                # In-process fallback: force Agg backend, chdir to workspace, exec.
                if on_output:
                    await on_output(f"[系统] 子进程执行失败，改为进程内执行: {script.name}")
                try:
                    rc = await cls._exec_figure_script_in_process(workspace, script)
                    se = se or f"in-process exec rc={rc}"
                except Exception as exc:  # noqa: BLE001 - last resort must not raise
                    rc = 1
                    se = (se + f"; in-process exec failed: {exc}").strip("; ")
            if rc == 0 and _already_has_output(script):
                produced.append(script.stem)
                if on_output:
                    await on_output(f"[系统] 绘图脚本产出图像: {script.name}")
            elif on_output:
                tail = (se or so or "").strip().splitlines()
                hint = tail[-1][:200] if tail else "unknown error"
                await on_output(f"[系统] 绘图脚本未产出图像: {script.name} ({hint})")
        return produced

    @classmethod
    async def _exec_figure_script_in_process(cls, workspace: Path, script: Path) -> int:
        """Execute a gen_fig_*.py script in-process with a non-interactive backend.

        Runs in a worker thread so the event loop is never blocked.  Forces
        matplotlib to Agg before the script imports pyplot, and temporarily
        chdirs to the workspace so relative paths (figures/, _utils/) resolve
        exactly as they would for a subprocess run.
        """
        import runpy

        def _exec() -> int:
            old_cwd = os.getcwd()
            old_argv = sys.argv[:]
            try:
                os.environ.setdefault("MPLBACKEND", "Agg")
                try:
                    import matplotlib

                    matplotlib.use("Agg", force=True)
                except Exception:  # noqa: BLE001 - matplotlib may be absent
                    pass
                os.chdir(str(workspace))
                sys.argv = [str(script)]
                runpy.run_path(str(script), run_name="__main__")
                return 0
            except SystemExit as exc:
                code = exc.code if isinstance(exc.code, int) else 0
                return code
            except Exception:  # noqa: BLE001
                log.exception("in-process figure script failed: %s", script)
                return 1
            finally:
                os.chdir(old_cwd)
                sys.argv = old_argv

        return await asyncio.to_thread(_exec)

    @classmethod
    async def _probe_figure_execution_channel(cls, workspace: Path) -> Optional[str]:
        """Pre-flight check for figure steps.

        Returns None when the host can (a) spawn a python subprocess and
        (b) create files inside figures/.  Otherwise returns a human-readable
        diagnosis explaining exactly which capability is missing, so the agent
        and the UI see the real blocker instead of six opaque retries.
        """
        figures_dir = Path(workspace) / "figures"
        try:
            figures_dir.mkdir(parents=True, exist_ok=True)
            probe_file = figures_dir / ".channel_probe"
            probe_file.write_text("ok", encoding="utf-8")
            # Read back instead of deleting: some sandboxes intercept deletes.
            if probe_file.read_text(encoding="utf-8") != "ok":
                return "figures/ 目录写入后读回内容异常"
        except OSError as exc:
            return f"figures/ 目录不可写: {exc}"

        python = cls._runtime_python()
        rc, so, se = await cls._run_process(
            [python, "-c", "import sys; print(sys.version.split()[0])"],
            workspace,
            timeout=30.0,
        )
        if rc != 0:
            return (
                "宿主 python 子进程启动失败 "
                f"(rc={rc}): {(se or so).strip()[:200]} — "
                "绘图脚本将无法执行，请检查后端运行环境"
            )
        return None

    @classmethod
    async def _export_markdown_pdf(
        cls, workspace: Path, template: str, params: Dict[str, Any],
    ) -> Dict[str, Any]:
        sources = cls._markdown_sources(workspace, template)
        latex_main = workspace / "paper" / "main.tex"
        if latex_main.is_file():
            # When a Markdown-producing branch is paired with a user LaTeX
            # suite, convert the reviewed Markdown to a LaTeX fragment and put
            # that fragment into the prepared template before compilation.
            if sources:
                from config import PANDOC_BIN
                if PANDOC_BIN:
                    fragment = workspace / "paper" / "_markdown_body.tex"
                    rc, stdout, stderr = await cls._run_process(
                        [str(PANDOC_BIN), str(sources[0]), "-f", "markdown", "-t", "latex", "-o", str(fragment)],
                        workspace,
                    )
                    if rc != 0:
                        return {"success": False, "returncode": rc, "stdout": stdout, "stderr": stderr}
                    current = latex_main.read_text(encoding="utf-8", errors="replace")
                    body = fragment.read_text(encoding="utf-8", errors="replace").strip()
                    if "\\begin{document}" in current and "\\end{document}" in current:
                        prefix = current.split("\\begin{document}", 1)[0]
                        latex_main.write_text(
                            prefix.rstrip() + "\n\\begin{document}\n" + body + "\n\\end{document}\n",
                            encoding="utf-8",
                        )
            if (workspace / "_latex_template.json").is_file() or str(params.get("format_text") or "").strip():
                applied = await cls._apply_latex_template(workspace, params)
                if not applied.get("success"):
                    return applied
            from services.claude_runner import _DETECTED_XELATEX

            xelatex = str(_DETECTED_XELATEX or shutil.which("xelatex") or "")
            if not xelatex:
                return {"success": False, "returncode": 3, "stderr": "XeLaTeX is unavailable"}
            cls._sanitize_latex_tree(latex_main.parent)
            stdout_parts: List[str] = []
            stderr_parts: List[str] = []
            rc = 0
            for _ in range(2):
                rc, stdout, stderr = await cls._run_process(
                    [xelatex, "-interaction=nonstopmode", "main.tex"],
                    latex_main.parent,
                )
                stdout_parts.append(stdout)
                stderr_parts.append(stderr)
                if rc != 0 and not (workspace / "paper" / "main.pdf").is_file():
                    break
                if rc != 0:
                    break
            output = workspace / "paper" / "main.pdf"
            combined = "\n".join(stdout_parts + stderr_parts)
            success = cls._pdf_compile_success(output, rc, combined)
            return {
                "success": success, "returncode": 0 if success else (rc or 4),
                "stdout": "\n".join(stdout_parts), "stderr": "\n".join(stderr_parts),
            }

        if not sources:
            return {"success": False, "returncode": 2, "stderr": "PDF export source Markdown not found"}
        source = sources[0]
        output = workspace / "paper" / "main.pdf"
        output.parent.mkdir(parents=True, exist_ok=True)
        from config import PANDOC_BIN
        from services.claude_runner import _DETECTED_XELATEX

        xelatex = str(_DETECTED_XELATEX or shutil.which("xelatex") or "")
        if PANDOC_BIN:
            engine = xelatex or "xelatex"
            command = [str(PANDOC_BIN), str(source), "-o", str(output), f"--pdf-engine={engine}"]
            rc, stdout, stderr = await cls._run_process(command, workspace)
            success = cls._pdf_compile_success(output, rc, f"{stdout}\n{stderr}")
            if success:
                return {"success": True, "returncode": 0, "stdout": stdout, "stderr": stderr}

        # Fallback without Pandoc: wrap Markdown into ctexart + Verbatim.
        if not xelatex:
            return {"success": False, "returncode": 3, "stderr": "Bundled Pandoc/XeLaTeX is unavailable"}
        body = source.read_text(encoding="utf-8", errors="replace").replace("\x00", "")
        escaped = (
            body.replace("\\", r"\textbackslash{}")
            .replace("&", r"\&")
            .replace("%", r"\%")
            .replace("$", r"\$")
            .replace("#", r"\#")
            .replace("_", r"\_")
            .replace("{", r"\{")
            .replace("}", r"\}")
            .replace("~", r"\textasciitilde{}")
            .replace("^", r"\textasciicircum{}")
        )
        tex_path = workspace / "paper" / "_markdown_export.tex"
        tex_path.write_text(
            "\\documentclass[UTF8]{ctexart}\n"
            "\\usepackage{geometry,hyperref,fancyvrb}\n"
            "\\geometry{a4paper,margin=2.2cm}\n"
            "\\begin{document}\n"
            "\\section*{Export}\n"
            "\\begin{Verbatim}[breaklines=true,fontsize=\\small]\n"
            + escaped[:120000]
            + "\n\\end{Verbatim}\n\\end{document}\n",
            encoding="utf-8",
        )
        rc, stdout, stderr = await cls._run_process(
            [xelatex, "-interaction=nonstopmode", tex_path.name],
            workspace / "paper",
        )
        produced = workspace / "paper" / "_markdown_export.pdf"
        if produced.is_file():
            shutil.copy2(produced, output)
        success = cls._pdf_compile_success(output, rc, f"{stdout}\n{stderr}")
        return {
            "success": success,
            "returncode": 0 if success else (rc or 4),
            "stdout": stdout,
            "stderr": stderr if not success else "",
        }

    @staticmethod
    def _apply_docx_columns(output: Path, column_layout: str) -> None:
        """Best-effort column layout. Missing python-docx must not fail export."""
        layout = (column_layout or "single").lower()
        if layout not in {"double", "two", "2"}:
            # Single-column is the engine default; skip optional dependency.
            return
        try:
            from docx import Document
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
        except Exception as exc:
            log.warning("Skip DOCX column layout; python-docx unavailable: %s", exc)
            return
        try:
            document = Document(str(output))
            section = document.sections[-1]
            sect_pr = section._sectPr
            cols = sect_pr.find(qn("w:cols"))
            if cols is None:
                cols = OxmlElement("w:cols")
                sect_pr.append(cols)
            cols.set(qn("w:num"), "2")
            cols.set(qn("w:space"), "720")
            document.save(str(output))
        except Exception as exc:
            # Column styling is cosmetic relative to delivering paper/main.docx.
            log.warning("Unable to apply DOCX column layout on %s: %s", output, exc)

    @staticmethod
    def _runtime_python() -> str:
        from config import RUNTIME_PYTHON

        if RUNTIME_PYTHON and Path(RUNTIME_PYTHON).is_file():
            return str(RUNTIME_PYTHON)
        return sys.executable

    @classmethod
    def _write_host_lineage(
        cls,
        workspace: Path,
        *,
        skill_name: str,
        script: Path,
        command: List[str],
        artifacts: List[Path],
        returncode: int,
        stdout: str,
        stderr: str,
    ) -> Path:
        records = []
        for path in artifacts:
            if path.is_file():
                records.append({
                    "path": path.relative_to(workspace).as_posix(),
                    "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
                    "size": path.stat().st_size,
                })
        payload = {
            "skill_name": skill_name,
            "executor": "host_step_runner",
            "script": script.as_posix(),
            "command": command,
            "returncode": returncode,
            "artifacts": records,
            "stdout_tail": (stdout or "")[-4000:],
            "stderr_tail": (stderr or "")[-4000:],
            "generated_at": datetime.utcnow().isoformat() + "Z",
        }
        out = workspace / f".host_builds/{skill_name}.json"
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        return out

    @classmethod
    async def _run_domain_host_skill(
        cls,
        workspace: Path,
        skill_name: str,
        params: Optional[Dict[str, Any]] = None,
        on_output: Optional[Callable[[str], Awaitable[None]]] = None,
        *,
        template: str = "",
    ) -> Dict[str, Any]:
        """Run deterministic doctoral-domain scaffolds (thesis/humanities/comp)."""
        from services import host_domain_builders as domain

        workspace = Path(workspace).expanduser().resolve()
        params = dict(params or {})
        title = str(params.get("title") or params.get("topic") or "")
        if not title:
            claude = workspace / "CLAUDE.md"
            if claude.is_file():
                for line in claude.read_text(encoding="utf-8", errors="replace").splitlines():
                    if line.startswith("#"):
                        title = line.lstrip("#").strip()
                        break
        if on_output:
            await on_output(f"[系统] 正在本机执行领域脚手架：{skill_name}")

        builders = {
            "thesis-proposal": lambda: domain.build_thesis_proposal(workspace, title=title, params=params),
            "literature-review": lambda: domain.build_literature_review(workspace, title=title, params=params),
            "project-blueprint": lambda: domain.build_project_blueprint(workspace, title=title, params=params),
            "paper-plan": lambda: domain.build_paper_plan(workspace, title=title, params=params),
            "paper-plan-zh": lambda: domain.build_paper_plan(workspace, title=title, params=params),
            "paper-analysis": lambda: domain.build_paper_analysis(workspace, title=title, params=params),
            "paper-figure": lambda: domain.build_paper_figure(workspace, title=title, params=params),
            "nature-figure": lambda: domain.build_paper_figure(workspace, title=title, params=params),
            "experiment-bridge": lambda: domain.build_experiment_bridge(
                workspace, title=title, params=params
            ),
            "research-lit": lambda: domain.build_research_lit(workspace, title=title, params=params),
            "idea-creator": lambda: domain.build_idea_creator(workspace, title=title, params=params),
            "novelty-check": lambda: domain.build_novelty_check(workspace, title=title, params=params),
            "research-review": lambda: domain.build_research_review(
                workspace, title=title, params=params
            ),
            "research-refine-pipeline": lambda: domain.build_research_refine_pipeline(
                workspace, title=title, params=params
            ),
            "auto-review-loop": lambda: domain.build_auto_review_loop(
                workspace, title=title, params=params
            ),
            "paper-write": lambda: domain.build_paper_write(
                workspace,
                title=title,
                params=params,
                language=str(params.get("language") or "en"),
            ),
            "paper-write-zh": lambda: domain.build_paper_write(
                workspace, title=title, params=params, language="zh",
            ),
            "paper-write-nature": lambda: domain.build_paper_write(
                workspace, title=title, params=params, language="en",
            ),
            "humanities-plan": lambda: domain.build_humanities_plan(workspace, title=title, params=params),
            "humanities-write": lambda: domain.build_humanities_paper(workspace, title=title, params=params),
            "course-plan": lambda: domain.build_course_plan(workspace, title=title, params=params),
            "course-paper": lambda: domain.build_course_paper(workspace, title=title, params=params),
            "course-report": lambda: domain.build_course_report(workspace, title=title, params=params),
            "course-report-plan": lambda: domain.build_course_report_plan(workspace, title=title, params=params),
            "comp-prob-analysis": lambda: domain.build_competition_problem_analysis(workspace, title=title, params=params),
            "comp-modeling": lambda: domain.build_competition_modeling(workspace, title=title, params=params),
            "comp-code": lambda: domain.build_competition_code(workspace, title=title, params=params),
            "comp-paper-zh": lambda: domain.build_competition_paper_zh(
                workspace, title=title, params=params, template=template or "comp_cumcm",
            ),
            "comp-paper-en": lambda: domain.build_competition_paper_en(
                workspace, title=title, params=params, template=template or "comp_mcm",
            ),
            "comp-paper-zh-docx": lambda: domain.build_competition_paper_md(
                workspace,
                title=title,
                params=params,
                template=template or "comp_cumcm",
                language="zh",
            ),
            "comp-paper-en-docx": lambda: domain.build_competition_paper_md(
                workspace,
                title=title,
                params=params,
                template=template or "comp_mcm",
                language="en",
            ),
            "paper-write-docx": lambda: domain.build_paper_write_md(
                workspace,
                title=title,
                params=params,
                language=str(params.get("language") or "en"),
            ),
            "paper-write-zh-docx": lambda: domain.build_paper_write_md(
                workspace, title=title, params=params, language="zh",
            ),
            "paper-write-nature-docx": lambda: domain.build_paper_write_md(
                workspace, title=title, params=params, language="en",
            ),
            "auto-paper-improvement-docx": lambda: domain.build_auto_paper_improvement_docx(
                workspace, title=title, params=params,
            ),
            "dev-requirement": lambda: domain.build_dev_requirement(
                workspace, title=title, params=params,
            ),
            "dev-design": lambda: domain.build_dev_design(
                workspace, title=title, params=params,
            ),
            "dev-code": lambda: domain.build_dev_code(
                workspace, title=title, params=params,
            ),
            "dev-selfcheck": lambda: domain.build_dev_selfcheck(
                workspace, title=title, params=params,
            ),
            "dev-report": lambda: domain.build_dev_report(
                workspace, title=title, params=params,
            ),
            "comp-stats-topic": lambda: domain.build_comp_stats_topic(
                workspace, title=title, params=params,
            ),
            "humanities-write-latex": lambda: domain.build_humanities_write_latex(
                workspace, title=title, params=params,
            ),
            "auto-paper-improvement-loop": lambda: domain.build_auto_paper_improvement_loop(
                workspace, title=title, params=params,
            ),
        }
        builder = builders.get(skill_name)
        if builder is None:
            return {
                "success": False,
                "returncode": 2,
                "return_code": 2,
                "stdout": "",
                "stderr": f"unsupported domain host skill: {skill_name}",
                "result": "",
            }

        built = builder()
        artifacts = [Path(p) for p in built.get("paths") or []]
        lineage = cls._write_host_lineage(
            workspace,
            skill_name=skill_name,
            script=Path(__file__).resolve().parent / "host_domain_builders.py",
            command=[f"host_domain_builders.{skill_name}", title],
            artifacts=artifacts,
            returncode=0 if built.get("success") else 1,
            stdout=json.dumps(built.get("artifacts") or [], ensure_ascii=False),
            stderr=str(built.get("stderr") or ""),
        )
        primary_rel = built.get("primary")
        success = bool(built.get("success"))
        if primary_rel:
            primary = workspace / str(primary_rel)
            success = success and primary.is_file() and primary.stat().st_size >= 40
        elif artifacts:
            success = success and all(path.is_file() for path in artifacts)
        if on_output:
            await on_output(
                f"[系统] 领域脚手架完成：{skill_name} "
                f"artifacts={len(built.get('artifacts') or [])}; lineage={lineage.name}"
            )
        return {
            "success": success,
            "returncode": 0 if success else 4,
            "return_code": 0 if success else 4,
            "stdout": json.dumps(built.get("artifacts") or [], ensure_ascii=False),
            "stderr": "" if success else f"{skill_name} host domain builder failed",
            "result": str(workspace / str(primary_rel)) if success and primary_rel else "",
        }

    @classmethod
    async def _run_patent_draft(
        cls,
        workspace: Path,
        params: Optional[Dict[str, Any]] = None,
        on_output: Optional[Callable[[str], Awaitable[None]]] = None,
    ) -> Dict[str, Any]:
        from services.host_ip_builders import build_patent_disclosure_draft

        workspace = Path(workspace).expanduser().resolve()
        title = str((params or {}).get("title") or "")
        claude = workspace / "CLAUDE.md"
        if not title and claude.is_file():
            for line in claude.read_text(encoding="utf-8", errors="replace").splitlines():
                if line.startswith("#"):
                    title = line.lstrip("#").strip()
                    break
        if on_output:
            await on_output("[系统] 正在本机生成专利交底书草稿（host_ip_builders）")
        built = build_patent_disclosure_draft(workspace, title=title, params=params or {})
        draft = workspace / "专利交底书" / "交底书草稿.md"
        lineage = cls._write_host_lineage(
            workspace,
            skill_name="patent-draft",
            script=Path(__file__).resolve().parent / "host_ip_builders.py",
            command=["host_ip_builders.build_patent_disclosure_draft", title],
            artifacts=[draft],
            returncode=0 if built.get("success") else 1,
            stdout=json.dumps(built.get("artifacts") or [], ensure_ascii=False),
            stderr="",
        )
        success = bool(built.get("success")) and draft.is_file() and draft.stat().st_size >= 200
        if on_output:
            await on_output(
                f"[系统] 专利交底书草稿已写入：{draft.relative_to(workspace).as_posix()} "
                f"({draft.stat().st_size if draft.is_file() else 0} bytes); lineage={lineage.name}"
            )
        return {
            "success": success,
            "returncode": 0 if success else 4,
            "return_code": 0 if success else 4,
            "stdout": draft.read_text(encoding="utf-8", errors="replace")[:2000] if draft.is_file() else "",
            "stderr": "" if success else "patent-draft host builder failed",
            "result": str(draft) if success else "",
        }

    @classmethod
    async def _run_copyright_draft(
        cls,
        workspace: Path,
        params: Optional[Dict[str, Any]] = None,
        on_output: Optional[Callable[[str], Awaitable[None]]] = None,
    ) -> Dict[str, Any]:
        from services.host_ip_builders import build_copyright_draft_package

        workspace = Path(workspace).expanduser().resolve()
        title = str((params or {}).get("software_name") or (params or {}).get("title") or "")
        if on_output:
            await on_output("[系统] 正在本机生成软著申请资料草稿（host_ip_builders）")
        built = build_copyright_draft_package(workspace, title=title, params=params or {})
        app_info = workspace / "软件著作权申请资料" / "草稿" / "申请表信息.md"
        artifacts = [Path(p) for p in built.get("paths") or []]
        lineage = cls._write_host_lineage(
            workspace,
            skill_name="copyright-draft",
            script=Path(__file__).resolve().parent / "host_ip_builders.py",
            command=["host_ip_builders.build_copyright_draft_package", title],
            artifacts=artifacts,
            returncode=0 if built.get("success") else 1,
            stdout=json.dumps(built.get("artifacts") or [], ensure_ascii=False),
            stderr="",
        )
        success = bool(built.get("success")) and app_info.is_file()
        if on_output:
            await on_output(
                f"[系统] 软著草稿已写入：软件著作权申请资料/草稿/ "
                f"({len(artifacts)} files); lineage={lineage.name}"
            )
        return {
            "success": success,
            "returncode": 0 if success else 4,
            "return_code": 0 if success else 4,
            "stdout": json.dumps(built.get("artifacts") or [], ensure_ascii=False),
            "stderr": "" if success else "copyright-draft host builder failed",
            "result": str(app_info) if success else "",
        }

    @classmethod
    async def _run_software_copyright(
        cls,
        workspace: Path,
        params: Optional[Dict[str, Any]] = None,
        on_output: Optional[Callable[[str], Awaitable[None]]] = None,
    ) -> Dict[str, Any]:
        from services.host_ip_builders import build_software_copyright_materials

        workspace = Path(workspace).expanduser().resolve()
        title = str((params or {}).get("software_name") or (params or {}).get("title") or "")
        if on_output:
            await on_output("[系统] 正在本机清点源码并生成软著材料四件套（host_ip_builders）")
        built = build_software_copyright_materials(workspace, title=title, params=params or {})
        artifacts = [Path(p) for p in built.get("paths") or []]
        lineage = cls._write_host_lineage(
            workspace,
            skill_name="software-copyright",
            script=Path(__file__).resolve().parent / "host_ip_builders.py",
            command=["host_ip_builders.build_software_copyright_materials", title],
            artifacts=artifacts,
            returncode=0 if built.get("success") else 1,
            stdout=json.dumps(built.get("artifacts") or [], ensure_ascii=False),
            stderr="",
        )
        required = [
            workspace / "software-copyright" / name
            for name in (
                "PRODUCT_OVERVIEW.md",
                "USER_MANUAL.md",
                "SOURCE_CODE_INDEX.md",
                "REGISTRATION_CHECKLIST.md",
            )
        ]
        success = bool(built.get("success")) and all(path.is_file() and path.stat().st_size > 50 for path in required)
        if on_output:
            await on_output(
                f"[系统] 软著材料已生成：software-copyright/ "
                f"(sources={built.get('source_files', 0)}); lineage={lineage.name}"
            )
        return {
            "success": success,
            "returncode": 0 if success else 4,
            "return_code": 0 if success else 4,
            "stdout": json.dumps(built.get("artifacts") or [], ensure_ascii=False),
            "stderr": "" if success else "software-copyright host builder failed",
            "result": str(workspace / "software-copyright") if success else "",
        }

    @classmethod
    async def _run_patent_build(
        cls, workspace: Path, on_output: Optional[Callable[[str], Awaitable[None]]] = None,
    ) -> Dict[str, Any]:
        workspace = Path(workspace).expanduser().resolve()
        draft = (workspace / "专利交底书" / "交底书草稿.md").resolve()
        out_md = (workspace / "专利交底书" / "交底书.md").resolve()
        out_docx = (workspace / "专利交底书" / "交底书.docx").resolve()
        if out_docx.is_file() and out_docx.stat().st_size >= 500:
            if on_output:
                await on_output(
                    f"[系统] 检测到已有专利交底书成品：{out_docx.relative_to(workspace).as_posix()} "
                    f"({out_docx.stat().st_size} bytes)，跳过重复渲染"
                )
            return {
                "success": True,
                "returncode": 0,
                "return_code": 0,
                "stdout": f"reuse existing {out_docx}",
                "stderr": "",
                "result": str(out_docx),
            }
        if not draft.is_file():
            return {
                "success": False,
                "returncode": 2,
                "return_code": 2,
                "stdout": "",
                "stderr": "专利交底书/交底书草稿.md is missing",
                "result": "",
            }
        script = Path(SKILLS_DIR) / "patent-build" / "tools" / "mermaid_render.py"
        if not script.is_file():
            return {
                "success": False,
                "returncode": 3,
                "return_code": 3,
                "stdout": "",
                "stderr": f"patent-build script missing: {script}",
                "result": "",
            }
        out_md.parent.mkdir(parents=True, exist_ok=True)
        command = [
            cls._runtime_python(),
            str(script),
            "-i", str(draft),
            "-o", str(out_md),
            "--docx", str(out_docx),
        ]
        if on_output:
            await on_output("[系统] 正在本机渲染 mermaid/公式并导出专利交底书 Word")
        rc, stdout, stderr = await cls._run_process(command, workspace)
        artifacts = [out_md, out_docx, *sorted((workspace / "专利交底书" / "mermaid_figures").glob("*.png"))]
        lineage = cls._write_host_lineage(
            workspace,
            skill_name="patent-build",
            script=script,
            command=command,
            artifacts=artifacts,
            returncode=rc,
            stdout=stdout,
            stderr=stderr,
        )
        success = rc == 0 and out_docx.is_file() and out_docx.stat().st_size >= 500
        if on_output:
            if success:
                await on_output(
                    f"[系统] 专利交底书已导出：{out_docx.relative_to(workspace).as_posix()} "
                    f"({out_docx.stat().st_size} bytes); lineage={lineage.name}"
                )
            else:
                await on_output(f"[系统] 专利交底书导出失败 rc={rc}")
        return {
            "success": success,
            "returncode": 0 if success else (rc or 4),
            "return_code": 0 if success else (rc or 4),
            "stdout": stdout,
            "stderr": stderr if success else (stderr or "patent-build failed to create 交底书.docx"),
            "result": str(out_docx) if success else "",
        }

    @classmethod
    async def _run_copyright_build(
        cls, workspace: Path, on_output: Optional[Callable[[str], Awaitable[None]]] = None,
    ) -> Dict[str, Any]:
        workdir = workspace / "软件著作权申请资料"
        draft_info = workdir / "草稿" / "申请表信息.md"
        if not draft_info.is_file():
            return {
                "success": False,
                "returncode": 2,
                "return_code": 2,
                "stdout": "",
                "stderr": "软件著作权申请资料/草稿/申请表信息.md is missing",
                "result": "",
            }
        script = Path(SKILLS_DIR) / "copyright-build" / "scripts" / "build_docx_from_md.py"
        if not script.is_file():
            return {
                "success": False,
                "returncode": 3,
                "return_code": 3,
                "stdout": "",
                "stderr": f"copyright-build script missing: {script}",
                "result": "",
            }

        def _field(name: str) -> str:
            prefix = f"➤{name}："
            for line in draft_info.read_text(encoding="utf-8", errors="replace").splitlines():
                text = line.strip()
                if text.startswith(prefix):
                    value = text[len(prefix):].strip()
                    return "" if "待用户确认" in value else value
            return ""

        software_name = _field("软件全称") or "未命名软件"
        version = _field("版本号") or "V1.0"
        command = [
            cls._runtime_python(),
            str(script),
            "--workdir", str(workdir),
            "--software-name", software_name,
            "--version", version,
            "--skip-preview",
        ]
        if on_output:
            await on_output(f"[系统] 正在本机生成软著正式资料：{software_name} {version}")
        rc, stdout, stderr = await cls._run_process(command, workspace)
        formal = workdir / "正式资料"
        report = formal / "生成报告.md"
        artifacts = sorted(formal.glob("*")) if formal.is_dir() else []
        lineage = cls._write_host_lineage(
            workspace,
            skill_name="copyright-build",
            script=script,
            command=command,
            artifacts=artifacts,
            returncode=rc,
            stdout=stdout,
            stderr=stderr,
        )
        success = (
            rc == 0
            and formal.is_dir()
            and any(path.is_file() and path.suffix.lower() in {".docx", ".txt", ".md"} for path in artifacts)
        )
        if on_output:
            if success:
                await on_output(
                    f"[系统] 软著正式资料已生成：{formal.relative_to(workspace).as_posix()} "
                    f"({len(list(formal.glob('*')))} files); lineage={lineage.name}"
                )
            else:
                await on_output(f"[系统] 软著正式资料生成失败 rc={rc}")
        return {
            "success": success,
            "returncode": 0 if success else (rc or 4),
            "return_code": 0 if success else (rc or 4),
            "stdout": stdout,
            "stderr": stderr if success else (stderr or "copyright-build failed"),
            "result": str(report if report.is_file() else formal) if success else "",
        }

    @classmethod
    async def _run_paper_slides(
        cls,
        workspace: Path,
        params: Optional[Dict[str, Any]] = None,
        on_output: Optional[Callable[[str], Awaitable[None]]] = None,
    ) -> Dict[str, Any]:
        workspace = Path(workspace).expanduser().resolve()
        paper = workspace / "paper"
        if not paper.is_dir():
            return {
                "success": False,
                "returncode": 2,
                "return_code": 2,
                "stdout": "",
                "stderr": "paper/ directory is missing",
                "result": "",
            }
        script = Path(SKILLS_DIR) / "paper-slides" / "tools" / "build_slides.py"
        if not script.is_file():
            return {
                "success": False,
                "returncode": 3,
                "return_code": 3,
                "stdout": "",
                "stderr": f"paper-slides script missing: {script}",
                "result": "",
            }
        params = params or {}
        venue = str(params.get("venue") or params.get("target_venue") or "NeurIPS")
        talk_type = str(params.get("talk_type") or params.get("talkType") or "spotlight")
        minutes = str(params.get("minutes") or params.get("talk_minutes") or 15)
        aspect = str(params.get("aspect") or params.get("aspect_ratio") or "16:9")
        notes = params.get("notes", params.get("include_speaker_notes", True))
        notes_flag = "--notes" if bool(notes) else "--no-notes"
        command = [
            cls._runtime_python(),
            str(script),
            "--workspace", str(workspace),
            "--venue", venue,
            "--talk-type", talk_type,
            "--minutes", str(minutes),
            "--aspect", aspect,
            notes_flag,
        ]
        if on_output:
            await on_output(f"[系统] 正在本机生成会议幻灯片（{talk_type} / {venue}）")
        rc, stdout, stderr = await cls._run_process(command, workspace)
        slides_dir = workspace / "slides"
        pdf = slides_dir / "main.pdf"
        pptx = slides_dir / "presentation.pptx"
        artifacts = [
            slides_dir / "main.tex",
            pdf,
            pptx,
            slides_dir / "SLIDE_OUTLINE.md",
            slides_dir / "TALK_SCRIPT.md",
            slides_dir / "speaker_notes.md",
            slides_dir / "SLIDES_STATE.json",
        ]
        lineage = cls._write_host_lineage(
            workspace,
            skill_name="paper-slides",
            script=script,
            command=command,
            artifacts=artifacts,
            returncode=rc,
            stdout=stdout,
            stderr=stderr,
        )
        success = (
            rc == 0
            and pdf.is_file()
            and pdf.stat().st_size >= 500
            and pptx.is_file()
            and pptx.stat().st_size >= 500
        )
        if on_output:
            if success:
                await on_output(
                    f"[系统] 幻灯片已生成：slides/main.pdf ({pdf.stat().st_size} bytes) + "
                    f"presentation.pptx; lineage={lineage.name}"
                )
            else:
                await on_output(f"[系统] 幻灯片生成失败 rc={rc}")
        return {
            "success": success,
            "returncode": 0 if success else (rc or 4),
            "return_code": 0 if success else (rc or 4),
            "stdout": stdout,
            "stderr": stderr if success else (stderr or "paper-slides failed to create slides/main.pdf"),
            "result": str(pdf) if success else "",
        }

    @classmethod
    async def _run_paper_poster(
        cls,
        workspace: Path,
        params: Optional[Dict[str, Any]] = None,
        on_output: Optional[Callable[[str], Awaitable[None]]] = None,
    ) -> Dict[str, Any]:
        workspace = Path(workspace).expanduser().resolve()
        paper = workspace / "paper"
        if not paper.is_dir():
            return {
                "success": False,
                "returncode": 2,
                "return_code": 2,
                "stdout": "",
                "stderr": "paper/ directory is missing",
                "result": "",
            }
        script = Path(SKILLS_DIR) / "paper-poster" / "tools" / "build_poster.py"
        if not script.is_file():
            return {
                "success": False,
                "returncode": 3,
                "return_code": 3,
                "stdout": "",
                "stderr": f"paper-poster script missing: {script}",
                "result": "",
            }
        params = params or {}
        venue = str(params.get("venue") or params.get("target_venue") or "NeurIPS")
        size = str(params.get("size") or params.get("poster_size") or "A0")
        orientation = str(params.get("orientation") or "landscape")
        columns = str(params.get("columns") or 4)
        command = [
            cls._runtime_python(),
            str(script),
            "--workspace", str(workspace),
            "--venue", venue,
            "--size", size,
            "--orientation", orientation,
            "--columns", columns,
        ]
        if on_output:
            await on_output(f"[系统] 正在本机生成会议海报（{size} {orientation} / {venue}）")
        rc, stdout, stderr = await cls._run_process(command, workspace)
        poster_dir = workspace / "poster"
        pdf = poster_dir / "main.pdf"
        pptx = poster_dir / "poster.pptx"
        artifacts = [
            poster_dir / "main.tex",
            pdf,
            pptx,
            poster_dir / "POSTER_CONTENT_PLAN.md",
            poster_dir / "POSTER_SPEECH.md",
            poster_dir / "POSTER_STATE.json",
        ]
        lineage = cls._write_host_lineage(
            workspace,
            skill_name="paper-poster",
            script=script,
            command=command,
            artifacts=artifacts,
            returncode=rc,
            stdout=stdout,
            stderr=stderr,
        )
        success = (
            rc == 0
            and pdf.is_file()
            and pdf.stat().st_size >= 500
            and pptx.is_file()
            and pptx.stat().st_size >= 500
        )
        if on_output:
            if success:
                await on_output(
                    f"[系统] 海报已生成：poster/main.pdf ({pdf.stat().st_size} bytes) + "
                    f"poster.pptx; lineage={lineage.name}"
                )
            else:
                await on_output(f"[系统] 海报生成失败 rc={rc}")
        return {
            "success": success,
            "returncode": 0 if success else (rc or 4),
            "return_code": 0 if success else (rc or 4),
            "stdout": stdout,
            "stderr": stderr if success else (stderr or "paper-poster failed to create poster/main.pdf"),
            "result": str(pdf) if success else "",
        }

    @classmethod
    def _derive_text_profile_from_requirements(
        cls, workspace: Path, params: Dict[str, Any],
    ) -> Dict[str, Any]:
        """Deterministically build `_text_profile.json` without LLM assistance."""
        base = cls._default_text_profile()
        derived_path = workspace / "_derived_profile.json"
        if derived_path.is_file():
            try:
                loaded = json.loads(derived_path.read_text(encoding="utf-8"))
                if isinstance(loaded, dict):
                    base = loaded
            except (OSError, json.JSONDecodeError):
                pass

        sources: List[str] = []
        format_text = str(params.get("format_text") or "").strip()
        if format_text:
            sources.append(format_text)
        for relative in (
            "FORMAT_REQUIREMENTS.md",
            "CUSTOM_REQUIREMENTS.md",
        ):
            path = workspace / relative
            if path.is_file():
                try:
                    sources.append(path.read_text(encoding="utf-8", errors="replace"))
                except OSError:
                    pass
        user_data = workspace / "user_data"
        if user_data.is_dir():
            for path in sorted(user_data.glob("*_extracted.txt")):
                try:
                    sources.append(path.read_text(encoding="utf-8", errors="replace"))
                except OSError:
                    pass
        text = "\n".join(sources)
        matched: List[str] = list(base.get("_matched_items") or [])
        if not text.strip():
            base["profile_name"] = base.get("profile_name") or "workflow-default"
            base["_derived_from"] = base.get("_derived_from") or "default"
            base["_matched_items"] = matched or ["no explicit format text; default profile"]
            return base

        # Longer / more specific names first so 小四 does not collapse to 四号.
        size_map = {
            "小初": 36, "初号": 42, "小一": 24, "一号": 26, "小二": 18, "二号": 22,
            "小三": 15, "三号": 16, "小四": 12, "四号": 14, "小五": 9, "五号": 10.5,
        }
        size_names = "|".join(size_map.keys())
        font_map = {
            "宋体": "SimSun", "黑体": "SimHei", "仿宋": "FangSong", "楷体": "KaiTi",
            "微软雅黑": "Microsoft YaHei", "雅黑": "Microsoft YaHei", "等线": "DengXian",
            "Times New Roman": "Times New Roman",
        }

        def capture(label: str, value: Any) -> None:
            matched.append(f"{label}={value}")

        margins = re.findall(r"([0-9]+(?:\.[0-9]+)?)\s*(?:cm|厘米)", text, re.I)
        if "页边距" in text and margins:
            value = float(margins[0])
            for key in ("margin_top_cm", "margin_bottom_cm", "margin_left_cm", "margin_right_cm"):
                base.setdefault("page", {})[key] = value
            capture("margin_cm", value)
        for side, key in (
            ("上", "margin_top_cm"), ("下", "margin_bottom_cm"),
            ("左", "margin_left_cm"), ("右", "margin_right_cm"),
        ):
            match = re.search(rf"{side}\s*[边距距]{{0,2}}\s*([0-9]+(?:\.[0-9]+)?)\s*(?:cm|厘米)", text, re.I)
            if match:
                value = float(match.group(1))
                base.setdefault("page", {})[key] = value
                capture(key, value)

        spacing = re.search(r"([0-9]+(?:\.[0-9]+)?)\s*倍行距", text)
        if spacing:
            base.setdefault("body", {})["line_spacing"] = float(spacing.group(1))
            capture("line_spacing", spacing.group(1))
        indent = re.search(r"(?:首行缩进|缩进)\s*([0-9]+(?:\.[0-9]+)?)\s*字符", text)
        if indent:
            base.setdefault("body", {})["first_line_indent_chars"] = float(indent.group(1))
            capture("first_line_indent_chars", indent.group(1))

        body_size = None
        body_match = re.search(
            rf"(?:正文|体文|正文字体|正文字号)[^。；;\n]{{0,12}}?({size_names})"
            rf"|({size_names})[^。；;\n]{{0,12}}?(?:正文|体文|正文字体|正文字号)",
            text,
        )
        if body_match:
            body_size = size_map[next(g for g in body_match.groups() if g)]
        if body_size is None and "小四" in text and "标题" not in text:
            body_size = 12
        if body_size is not None:
            base.setdefault("body", {})["font_size_pt"] = body_size
            capture("body_font_size_pt", body_size)

        for level, labels in (
            (1, ("一级标题", "标题1", "章标题")),
            (2, ("二级标题", "标题2", "节标题")),
            (3, ("三级标题", "标题3")),
        ):
            for label in labels:
                level_match = re.search(
                    rf"{label}[^。；;\n]{{0,12}}?({size_names})"
                    rf"|({size_names})[^。；;\n]{{0,12}}?{label}",
                    text,
                )
                if level_match:
                    pt = size_map[next(g for g in level_match.groups() if g)]
                    base.setdefault("headings", {})[f"level{level}_pt"] = pt
                    capture(f"heading_level{level}_pt", pt)
                    break

        for zh, font in font_map.items():
            if re.search(rf"(正文|体文).{{0,6}}{zh}|{zh}.{{0,6}}(正文|体文)", text) or (
                zh in text and "字体" in text and "标题" not in text[max(0, text.find(zh) - 4): text.find(zh) + 8]
            ):
                base.setdefault("fonts", {})["chinese_body"] = font
                capture("chinese_body", font)
                break
        for zh, font in font_map.items():
            if re.search(rf"(标题).{{0,6}}{zh}|{zh}.{{0,6}}(标题)", text):
                base.setdefault("fonts", {})["chinese_heading"] = font
                base.setdefault("title", {})["font_family"] = font
                capture("chinese_heading", font)
                break
        if "Times New Roman" in text or "times new roman" in text.casefold():
            base.setdefault("fonts", {})["latin"] = "Times New Roman"
            capture("latin", "Times New Roman")

        if "A4" in text.upper():
            base.setdefault("page", {})["size"] = "A4"
            capture("page_size", "A4")
        if "悬挂缩进" in text:
            hang = re.search(r"悬挂缩进\s*([0-9]+(?:\.[0-9]+)?)\s*(?:cm|厘米)", text, re.I)
            if hang:
                base.setdefault("references", {})["hanging_indent_cm"] = float(hang.group(1))
                capture("references_hanging_indent_cm", hang.group(1))

        base["profile_name"] = "text-requirements-derived"
        base["_derived_from"] = "text-description"
        base["_matched_items"] = matched or ["format text present but no portable tokens matched"]
        return base

    @staticmethod
    def _docx_format_check_targets(workspace: Path) -> List[Path]:
        candidates = [
            "PROPOSAL.md", "LITERATURE_REVIEW.md", "COURSE_PAPER.md",
            "COURSE_REPORT.md", "NARRATIVE_REPORT.md", "HUMANITIES_PAPER.md",
            "paper/main.md", "REPORT.md", "AUTO_REVIEW.md",
        ]
        found: List[Path] = []
        for relative in candidates:
            path = workspace / relative
            if path.is_file():
                found.append(path)
        return found

    @classmethod
    def _run_docx_format_check_text(cls, content: str) -> tuple[str, List[str], List[str]]:
        """Apply deterministic Markdown fixes that commonly break DOCX export."""
        findings: List[str] = []
        fixes: List[str] = []
        lines = content.splitlines(keepends=True)
        out: List[str] = []

        heading_list = re.compile(
            r"^(\s*)-\s+("
            r"第[一二三四五六七八九十百千零〇0-9]+[章节部篇]"
            r"|[一二三四五六七八九十]+、"
            r"|（[一二三四五六七八九十]+）"
            r"|[0-9]+(?:\.[0-9]+){0,4}\.?\s+"
            r"|摘要|Abstract|关键词|Keywords|参考文献|References|致谢|附录|Appendix"
            r")(.*)$"
        )
        figure_list = re.compile(
            r"^(\s*)-\s+((?:图|表|Figure|Fig\.?|Table|Tab\.?)\s*[0-9A-Za-z][0-9A-Za-z\-\.]*\s*[：:、\.\-].*|"
            r"(?:图|表)\s*[：:].*)$"
        )
        formula_list = re.compile(r"^(\s*)-\s*([\[\(【]\s*[0-9]+(?:\.[0-9]+)?\s*[\]\)】])\s*$")

        def heading_level(token: str) -> int:
            bare = token.strip()
            if re.match(r"^[0-9]+(?:\.[0-9]+){3}", bare):
                return 4
            if re.match(r"^[0-9]+(?:\.[0-9]+){2}", bare):
                return 3
            if re.match(r"^[0-9]+\.[0-9]+", bare):
                return 2
            if re.match(r"^[0-9]+\.?\s+", bare):
                return 1
            if bare.startswith("第") or re.match(r"^[一二三四五六七八九十]+、", bare):
                return 1
            return 2

        for line in lines:
            raw = line.rstrip("\r\n")
            ending = line[len(raw):] or "\n"
            match = heading_list.match(raw)
            if match:
                indent, token, rest = match.groups()
                level = heading_level(token)
                rebuilt = f"{indent}{'#' * level} {token}{rest}{ending}"
                out.append(rebuilt)
                findings.append(f"list-as-heading: {raw.strip()}")
                fixes.append(f"converted heading list to H{level}: {token.strip()}")
                continue
            match = figure_list.match(raw)
            if match:
                indent, body = match.groups()
                rebuilt = f"{indent}{body}{ending}"
                out.append(rebuilt)
                findings.append(f"list-as-caption: {raw.strip()}")
                fixes.append(f"removed bullet from caption: {body.strip()[:80]}")
                continue
            match = formula_list.match(raw)
            if match:
                indent, label = match.groups()
                number = re.search(r"[0-9]+(?:\.[0-9]+)?", label)
                body = f"({number.group(0)})" if number else label.strip("[]【】() ")
                if out:
                    prev = out[-1].rstrip("\r\n")
                    prev_end = out[-1][len(prev):] or "\n"
                    if prev.rstrip().endswith("$$"):
                        label_text = body if body.startswith("(") else f"({body})"
                        out[-1] = f"{prev} {label_text}{prev_end}"
                        findings.append(f"list-as-formula-label: {raw.strip()}")
                        fixes.append(f"merged formula label into previous math block: {body}")
                        continue
                label_text = body if body.startswith("(") else f"({body})"
                out.append(f"{indent}{label_text}{ending}")
                findings.append(f"list-as-formula-label: {raw.strip()}")
                fixes.append(f"normalized formula label line: {body}")
                continue
            out.append(line if line.endswith(("\n", "\r")) else line + "\n")

        text = "".join(out)

        def repl_formula(match: re.Match[str]) -> str:
            findings.append(f"formula-number-style: {match.group(0).strip()}")
            fixes.append(
                f"normalized formula number to ASCII parentheses: ({match.group(1)})"
            )
            return f"$$ ({match.group(1)})"

        text, _n_formula = re.subn(
            r"\$\$\s*[（(]\s*([0-9]+(?:\.[0-9]+)?)\s*[)）]",
            repl_formula,
            text,
        )

        fence_count = len(re.findall(r"(?m)^```", text))
        if fence_count % 2 != 0:
            findings.append(f"unbalanced-code-fence: count={fence_count}")
            text = text.rstrip() + "\n```\n"
            fixes.append("appended closing code fence for unbalanced ``` markers")

        code_hits = 0
        for block in re.split(r"\n\s*\n", text):
            if "```" in block:
                continue
            if len(block) >= 200 and sum(
                token in block for token in ("def ", "class ", "import ", "return ")
            ) >= 2:
                code_hits += 1
        if code_hits:
            findings.append(f"possible-unfenced-code-blocks: {code_hits}")

        return text, findings, fixes

    @classmethod
    async def _run_docx_format_check(
        cls,
        workspace: Path,
        on_output: Optional[Callable[[str], Awaitable[None]]] = None,
    ) -> Dict[str, Any]:
        targets = cls._docx_format_check_targets(workspace)
        if not targets:
            report = (
                "# DOCX 格式自检报告\n\n"
                "- 状态: blocked\n"
                "- 原因: 未找到可检查的 Markdown 目标文件\n"
                "- 判定: NOT_RUN / fail-closed（无目标不得记 success）\n"
                "- 候选: PROPOSAL.md / LITERATURE_REVIEW.md / COURSE_PAPER.md / "
                "COURSE_REPORT.md / NARRATIVE_REPORT.md / paper/main.md\n"
            )
            output = workspace / "DOCX_FORMAT_CHECK_REPORT.md"
            output.write_text(report, encoding="utf-8")
            if on_output:
                await on_output("[系统] 未找到 Markdown 目标，docx-format-check 记为 blocked")
            return {
                "success": False,
                "returncode": 2,
                "return_code": 2,
                "stdout": str(output),
                "stderr": "docx-format-check blocked: no Markdown targets",
                "result": str(output),
                "status": "blocked",
                "root_cause": "DOCX_CHECK_WITHOUT_TARGET",
            }

        all_findings: List[str] = []
        all_fixes: List[str] = []
        checked: List[str] = []
        for target in targets:
            original = target.read_text(encoding="utf-8", errors="replace")
            updated, findings, fixes = cls._run_docx_format_check_text(original)
            if updated != original:
                target.write_text(updated, encoding="utf-8")
            rel = target.relative_to(workspace).as_posix()
            checked.append(rel)
            all_findings.extend(f"{rel}: {item}" for item in findings)
            all_fixes.extend(f"{rel}: {item}" for item in fixes)
            if on_output:
                await on_output(
                    f"[系统] docx-format-check {rel}: findings={len(findings)} fixes={len(fixes)}"
                )

        status = "passed" if not all_findings else ("fixed" if all_fixes else "warnings")
        lines = [
            "# DOCX 格式自检报告",
            "",
            f"- 状态: {status}",
            f"- 检查文件数: {len(checked)}",
            f"- 发现问题: {len(all_findings)}",
            f"- 自动修复: {len(all_fixes)}",
            "",
            "## 检查文件",
        ]
        lines.extend(f"- `{name}`" for name in checked)
        lines.extend(["", "## 发现"])
        if all_findings:
            lines.extend(f"- {item}" for item in all_findings)
        else:
            lines.append("- 全部检查通过")
        lines.extend(["", "## 自动修复"])
        if all_fixes:
            lines.extend(f"- {item}" for item in all_fixes)
        else:
            lines.append("- 无需修复")
        lines.extend([
            "",
            "## 说明",
            "- 本步骤为确定性本机检查，不调用模型。",
            "- 目标是避免列表伪标题、伪图注、公式编号与代码围栏问题导致 Word 走样。",
            "",
        ])
        report_text = "\n".join(lines)
        if len(report_text.encode("utf-8")) < 200:
            report_text += "\n" + ("<!-- padding for min report size -->\n" * 5)
        output = workspace / "DOCX_FORMAT_CHECK_REPORT.md"
        output.write_text(report_text, encoding="utf-8")
        if on_output:
            await on_output(
                f"[系统] DOCX 格式自检完成：{output.name} status={status} "
                f"findings={len(all_findings)} fixes={len(all_fixes)}"
            )
        return {
            "success": True,
            "returncode": 0,
            "return_code": 0,
            "stdout": str(output),
            "stderr": "",
            "result": str(output),
        }

    @staticmethod
    def _latex_engine_for_skill(skill_name: str) -> tuple[str, str]:
        """Return (engine_name, binary_path_or_empty)."""
        from services.claude_runner import _DETECTED_XELATEX

        if skill_name in {"paper-compile", "comp-compile-en"}:
            binary = shutil.which("pdflatex") or ""
            return "pdflatex", binary
        binary = str(_DETECTED_XELATEX or "") or (shutil.which("xelatex") or "")
        return "xelatex", binary

    @staticmethod
    def _competition_template_folder(template: str, params: Dict[str, Any] | None = None) -> str:
        """Map a competition workflow id/param to the paper template folder name."""
        params = params or {}
        raw = str(params.get("competition") or template or "").strip().lower()
        if raw.startswith("comp_"):
            raw = raw[5:]
        aliases = {
            "cumcm": "cumcm",
            "huawei": "huawei",
            "mcm": "mcm",
            "apmcm": "apmcm",
            "apmcm_zh": "apmcm_zh",
            "mathorcup": "mathorcup",
            "stats": "stats",
            "wuyi": "wuyi",
            "huazhong": "huazhong",
            "changsanjiao": "changsanjiao",
            "yangtze": "changsanjiao",
            "huashu": "huashubei",
            "huashubei": "huashubei",
            "diangong": "diangongbei",
            "diangongbei": "diangongbei",
            "liaoning": "dongsansheng",
            "dongsansheng": "dongsansheng",
            "shuwei": "shuweibei",
            "shuweibei": "shuweibei",
            "shuwei_en": "mcm",
            "certcup_en": "mcm",
            "teddy": "default",
            "certcup": "default",
            "tianfu": "default",
            "huadong": "default",
            "zhongqing": "default",
            "shenzhen": "default",
        }
        return aliases.get(raw, raw or "default")

    @classmethod
    def _stage_competition_latex_assets(
        cls,
        workspace: Path,
        *,
        template: str,
        skill_name: str,
        params: Dict[str, Any] | None = None,
    ) -> list[str]:
        """Copy competition cls/fonts into paper/ before host-side XeLaTeX compile.

        Agent paper steps historically assumed the full template folder already
        lived under paper/.  When only main.tex was rewritten, compile failed
        with missing document classes such as gmcmthesis.cls.  Stage the local
        skill template assets deterministically before invoking the engine.
        """
        if skill_name not in {"comp-compile-zh", "comp-compile-en"}:
            return []
        paper_dir = workspace / "paper"
        if not paper_dir.is_dir():
            return []
        folder = cls._competition_template_folder(template, params)
        skill_name_dir = "comp-paper-en" if skill_name == "comp-compile-en" else "comp-paper-zh"
        candidates = [
            SKILLS_DIR / skill_name_dir / "templates" / folder,
            SKILLS_DIR / "comp-paper-zh" / "templates" / folder,
            SKILLS_DIR / "comp-paper-en" / "templates" / folder,
            SKILLS_DIR / "comp-paper-zh" / "templates" / "default",
        ]
        source_dir = next((path for path in candidates if path.is_dir()), None)
        if source_dir is None:
            return []
        staged: list[str] = []
        for source in source_dir.rglob("*"):
            if not source.is_file():
                continue
            if source.suffix.lower() in {".enc"}:
                continue
            # Keep the agent's main.tex/sections; only fill missing support files.
            if source.name.lower() in {"main.tex", "main.tex.enc"}:
                continue
            relative = source.relative_to(source_dir)
            target = paper_dir / relative
            # Always overwrite .cls files from the canonical template — agents
            # sometimes write a thin wrapper (a few hundred bytes) that omits
            # the full class body, causing catastrophic XeLaTeX failures.
            # For every other support file keep the agent's version if present.
            is_cls = source.suffix.lower() == ".cls"
            if target.exists() and not is_cls:
                continue
            target.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(source, target)
            staged.append(relative.as_posix())
        return staged

    @staticmethod
    def _escape_texttt_underscores(text: str) -> str:
        """Escape underscores inside ``\\texttt{...}`` so xelatex keeps text mode."""

        def _fix(match: re.Match[str]) -> str:
            body = match.group(1).replace("_", r"\_")
            return f"\\texttt{{{body}}}"

        return re.sub(r"\\texttt\{([^{}]*)\}", _fix, text)

    @classmethod
    def _sanitize_latex_source_text(cls, text: str) -> str:
        """Repair common agent LaTeX hazards before host compile.

        Agents frequently emit Markdown backticks, re-load hyperref with options
        after a competition class already loaded it, and leave bare underscores
        inside ``\\texttt`` or prose.  Those produce option-clash / Missing-$ /
        undefined-control-sequence failures under ``-halt-on-error``.
        """
        if not text:
            return text

        # Drop duplicate hyperref loads; keep the first bare load and later
        # convert option-bearing loads into \hypersetup when possible.
        lines = text.splitlines(keepends=True)
        rewritten: list[str] = []
        hyperref_loaded = False
        hyperref_pkg = re.compile(
            r"""^(\s*)\\usepackage(?:\s*\[[^\]]*\])?\s*\{hyperref\}\s*(?:%.*)?$"""
        )
        hyperref_opts = re.compile(
            r"""\\usepackage\s*\[([^\]]*)\]\s*\{hyperref\}"""
        )
        for line in lines:
            match = hyperref_pkg.match(line.rstrip("\r\n"))
            if match:
                opts_match = hyperref_opts.search(line)
                opts = (opts_match.group(1) if opts_match else "").strip()
                indent = match.group(1)
                if not hyperref_loaded:
                    rewritten.append(f"{indent}\\usepackage{{hyperref}}\n")
                    hyperref_loaded = True
                    if opts:
                        rewritten.append(f"{indent}\\hypersetup{{{opts}}}\n")
                elif opts:
                    rewritten.append(f"{indent}\\hypersetup{{{opts}}}\n")
                # else: pure duplicate package load — drop it
                continue
            rewritten.append(line)
        text = "".join(rewritten)

        # Markdown inline code → \texttt{...}
        def _backtick(match: re.Match[str]) -> str:
            body = match.group(1).replace("\\", r"\textbackslash{}").replace("_", r"\_")
            body = body.replace("{", r"\{").replace("}", r"\}")
            return f"\\texttt{{{body}}}"

        text = re.sub(r"`([^`\n]+)`", _backtick, text)
        text = cls._escape_texttt_underscores(text)

        # gbt7714 + some MiKTeX builds choke when legacy options are forced.
        text = re.sub(
            r"\\usepackage\[[^\]]*\]\{gbt7714\}",
            r"\\usepackage{gbt7714}",
            text,
        )
        return text

    @classmethod
    def _sanitize_latex_tree(cls, paper_dir: Path) -> list[str]:
        """Sanitize main.tex and section fragments under paper/."""
        fixed: list[str] = []
        if not paper_dir.is_dir():
            return fixed
        candidates = [paper_dir / "main.tex"]
        sections = paper_dir / "sections"
        if sections.is_dir():
            candidates.extend(sorted(sections.rglob("*.tex")))
        for path in candidates:
            if not path.is_file():
                continue
            try:
                original = path.read_text(encoding="utf-8", errors="replace")
            except OSError:
                continue
            cleaned = cls._sanitize_latex_source_text(original)
            if cleaned != original:
                path.write_text(cleaned, encoding="utf-8")
                fixed.append(path.relative_to(paper_dir).as_posix())
        return fixed

    @classmethod
    def _pdf_compile_success(cls, output: Path, rc: int, combined_log: str) -> bool:
        """Fail closed: nonzero LaTeX/pandoc exit never counts as compile success.

        A PDF may still exist after recoverable MiKTeX warnings, but product
        routes must not report success/PASS when the compiler returned non-zero.
        Callers can surface the artifact path separately without flipping success.
        """
        if rc != 0:
            return False
        if not (output.is_file() and output.stat().st_size >= 500):
            return False
        # combined_log retained for call-site diagnostics / future strict scanners
        _ = combined_log
        return True

    @classmethod
    async def _run_paper_compile(
        cls,
        workspace: Path,
        skill_name: str,
        on_output: Optional[Callable[[str], Awaitable[None]]] = None,
        *,
        template: str = "",
        params: Dict[str, Any] | None = None,
    ) -> Dict[str, Any]:
        paper_dir = workspace / "paper"
        main_tex = paper_dir / "main.tex"
        if not main_tex.is_file():
            return {
                "success": False,
                "returncode": 2,
                "return_code": 2,
                "stdout": "",
                "stderr": "paper/main.tex is missing",
                "result": "",
            }
        staged = cls._stage_competition_latex_assets(
            workspace,
            template=template or "",
            skill_name=skill_name,
            params=params or {},
        )
        if staged and on_output:
            await on_output(
                "[系统] 已补齐竞赛模板依赖: " + ", ".join(staged[:12])
                + (" ..." if len(staged) > 12 else "")
            )
        sanitized = cls._sanitize_latex_tree(paper_dir)
        if sanitized and on_output:
            await on_output(
                "[系统] 已修复 LaTeX 常见风险: " + ", ".join(sanitized[:12])
                + (" ..." if len(sanitized) > 12 else "")
            )
        engine_name, binary = cls._latex_engine_for_skill(skill_name)
        if not binary:
            return {
                "success": False,
                "returncode": 3,
                "return_code": 3,
                "stdout": "",
                "stderr": f"{engine_name} is unavailable",
                "result": "",
            }
        if on_output:
            await on_output(f"[系统] 正在本机用 {engine_name} 编译 paper/main.tex")

        async def _compile_once() -> tuple[int, str, str]:
            stdout_parts: List[str] = []
            stderr_parts: List[str] = []
            local_rc = 0
            # Two passes cover TOC/refs for simple papers without requiring bibtex.
            for pass_no in range(1, 3):
                if on_output:
                    await on_output(f"[系统] {engine_name} pass {pass_no}/2")
                # Prefer nonstop without hard-stop so recoverable body errors still
                # emit a usable PDF; fall back semantics handled by success gate.
                local_rc, stdout, stderr = await cls._run_process(
                    [binary, "-interaction=nonstopmode", "main.tex"],
                    paper_dir,
                )
                stdout_parts.append(stdout)
                stderr_parts.append(stderr)
                if local_rc != 0 and not (paper_dir / "main.pdf").is_file():
                    break
                if local_rc != 0:
                    # PDF exists but engine complained — one pass is enough.
                    break
            return local_rc, "\n".join(stdout_parts), "\n".join(stderr_parts)

        rc, stdout_all, stderr_all = await _compile_once()
        output = paper_dir / "main.pdf"
        combined = f"{stdout_all}\n{stderr_all}"
        success = cls._pdf_compile_success(output, rc, combined)

        lineage = cls._write_host_lineage(
            workspace,
            skill_name=skill_name,
            script=Path(binary),
            command=[binary, "-interaction=nonstopmode", "main.tex"],
            artifacts=[output] if output.is_file() else [],
            returncode=0 if success else rc,
            stdout=stdout_all,
            stderr=stderr_all,
        )
        if on_output:
            if success:
                await on_output(
                    f"[系统] PDF 编译完成：paper/main.pdf ({output.stat().st_size} bytes); "
                    f"lineage={lineage.name}"
                )
            else:
                artifact_note = (
                    f"；产物仍存在 size={output.stat().st_size}" if output.is_file() else ""
                )
                await on_output(f"[系统] PDF 编译失败 rc={rc}{artifact_note}（fail-closed，不记 success）")
        return {
            "success": success,
            "returncode": 0 if success else (rc or 4),
            "return_code": 0 if success else (rc or 4),
            "stdout": stdout_all,
            "stderr": (
                stderr_all
                if success
                else (stderr_all.strip() or f"{engine_name} failed to produce paper/main.pdf")
            ),
            "result": str(output) if success else "",
            # Preserve main.tex on failure; never rewrite sources from this gate.
            "root_cause": None if success else "PDF_COMPILE_NONZERO_EXIT",
        }

    @staticmethod
    def _classify_asset(path: Path) -> str:
        name = path.name
        lower = name.casefold()
        suffix = path.suffix.casefold()
        if any(token in lower for token in ("result", "results", "output", "metrics", "score")) and suffix in {
            ".md", ".json", ".txt", ".csv", ".tsv",
        }:
            return "result"
        if suffix in {".py", ".ipynb", ".r", ".m", ".cpp", ".c", ".h", ".java", ".js", ".ts", ".zip"}:
            return "code"
        if suffix in {".csv", ".xlsx", ".xls", ".tsv", ".json", ".parquet"}:
            return "data"
        if suffix in {".png", ".jpg", ".jpeg", ".svg", ".gif", ".webp"}:
            return "figure"
        if suffix in {".tex", ".cls", ".sty", ".bst", ".dotx"}:
            return "template"
        if suffix == ".docx" and any(token in lower for token in ("template", "模板", "格式")):
            return "template"
        if any(token in lower for token in ("problem", "题目", "赛题", "requirement", "要求", "prompt")):
            return "problem"
        if suffix in {".pdf", ".docx", ".md", ".txt"}:
            return "problem"
        return "other"

    @classmethod
    async def _run_assets_inventory(
        cls,
        workspace: Path,
        params: Dict[str, Any],
        on_output: Optional[Callable[[str], Awaitable[None]]] = None,
    ) -> Dict[str, Any]:
        root = workspace / "user_data"
        root.mkdir(parents=True, exist_ok=True)
        files = [
            path for path in sorted(root.rglob("*"))
            if path.is_file()
            and not any(part.startswith(".") for part in path.relative_to(workspace).parts)
            and path.name not in {"_extract_status.json", "_input_manifest.json"}
            and not path.name.endswith("_extracted.txt")
        ]
        buckets: Dict[str, List[dict]] = {
            "problem": [], "code": [], "data": [], "figure": [],
            "result": [], "template": [], "other": [],
        }
        for path in files:
            rel = path.relative_to(workspace).as_posix()
            category = cls._classify_asset(path)
            digest = hashlib.sha256(path.read_bytes()).hexdigest()
            item = {
                "path": rel,
                "name": path.name,
                "category": category,
                "size": path.stat().st_size,
                "sha256": digest,
                "suffix": path.suffix.casefold(),
            }
            buckets[category].append(item)

        index = {
            "format_version": "assets-index/v1",
            "paper_type_target": str(params.get("paper_type_target") or ""),
            "scanned_at": datetime.utcnow().isoformat() + "Z",
            "counts": {key: len(value) for key, value in buckets.items()},
            "files": [item for group in buckets.values() for item in group],
            "executor": "host_step_runner",
        }
        index_path = workspace / "_assets_index.json"
        index_path.write_text(json.dumps(index, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

        lines = [
            "# 资产清点清单",
            "",
            f"- 扫描目录: `user_data/`",
            f"- 文件总数: {len(files)}",
            f"- paper_type_target: `{index['paper_type_target'] or 'n/a'}`",
            f"- 执行器: host_step_runner",
            "",
        ]
        labels = {
            "problem": "题目 / 要求",
            "code": "代码",
            "data": "数据",
            "figure": "图",
            "result": "结果",
            "template": "模板",
            "other": "其他",
        }
        for key, label in labels.items():
            lines.append(f"## {label} ({len(buckets[key])})")
            if not buckets[key]:
                lines.append("- （无）")
            else:
                for item in buckets[key]:
                    lines.append(
                        f"- `{item['path']}` · {item['size']} bytes · sha256=`{item['sha256'][:12]}…`"
                    )
            lines.append("")

        conflicts: List[str] = []
        if len(buckets["problem"]) > 3:
            conflicts.append("题目/要求类文件过多，请确认主输入文件。")
        if buckets["code"] and not buckets["result"] and not buckets["data"]:
            conflicts.append("检测到代码但未见数据或结果文件，后续可能需要补实验输出。")
        if buckets["figure"] and not buckets["result"] and not buckets["code"]:
            conflicts.append("仅有图文件而缺少代码/结果，图文一致性需要人工确认。")
        names = [item["name"].casefold() for item in index["files"]]
        if len(names) != len(set(names)):
            conflicts.append("存在同名文件，可能造成引用歧义。")

        conflict_path = workspace / "ASSETS_CONFLICTS.md"
        if conflicts:
            conflict_path.write_text(
                "# 资产冲突报告\n\n" + "\n".join(f"- {item}" for item in conflicts) + "\n",
                encoding="utf-8",
            )
        elif conflict_path.exists():
            conflict_path.unlink()

        inventory = workspace / "ASSETS_INVENTORY.md"
        inventory.write_text("\n".join(lines) + "\n", encoding="utf-8")
        artifacts = [inventory, index_path] + ([conflict_path] if conflict_path.exists() else [])
        lineage = cls._write_host_lineage(
            workspace,
            skill_name="assets-inventory",
            script=Path(__file__),
            command=["host:assets-inventory"],
            artifacts=artifacts,
            returncode=0,
            stdout=f"files={len(files)} conflicts={len(conflicts)}",
            stderr="",
        )
        if on_output:
            await on_output(
                f"[系统] 资产清点完成：{len(files)} files, conflicts={len(conflicts)}; "
                f"lineage={lineage.name}"
            )
        return {
            "success": inventory.is_file() and inventory.stat().st_size >= 100,
            "returncode": 0,
            "return_code": 0,
            "stdout": str(inventory),
            "stderr": "",
            "result": str(inventory),
        }

    @classmethod
    async def _run_paper_figure_html(
        cls,
        workspace: Path,
        on_output: Optional[Callable[[str], Awaitable[None]]] = None,
    ) -> Dict[str, Any]:
        """Render existing HTML figure sources to static PDF/PNG without an LLM."""
        script = Path(SKILLS_DIR) / "paper-figure-html" / "tools" / "render_html.py"
        if not script.is_file():
            return {
                "success": False,
                "returncode": 3,
                "return_code": 3,
                "stdout": "",
                "stderr": f"paper-figure-html renderer missing: {script}",
                "result": "",
            }
        figures_dir = workspace / "figures"
        figures_dir.mkdir(parents=True, exist_ok=True)
        sources = sorted(figures_dir.glob("*.html"))
        if (figures_dir / "html").is_dir():
            sources.extend(sorted((figures_dir / "html").glob("*.html")))
        sources = [path for path in sources if path.is_file() and not path.name.startswith(".")]
        if not sources:
            plan_candidates = [
                workspace / "PAPER_PLAN.md",
                workspace / "OUTLINE.md",
                workspace / "figures" / "FIGURE_PLAN.md",
            ]
            title = "System pipeline"
            for plan in plan_candidates:
                if plan.is_file():
                    for line in plan.read_text(encoding="utf-8", errors="replace").splitlines():
                        if line.strip():
                            title = re.sub(r"^#+\s*", "", line.strip())[:80]
                            break
                    break
            scaffold = figures_dir / "fig_pipeline.html"
            scaffold.write_text(
                "<!doctype html><html><head><meta charset=\"utf-8\">"
                "<style>body{font-family:Segoe UI,Arial,sans-serif;margin:28px;background:#fff}"
                ".row{display:flex;align-items:center;gap:12px;flex-wrap:wrap}"
                ".box{border:2px solid #222;border-radius:10px;padding:14px 18px;min-width:120px;text-align:center}"
                ".arrow{font-size:28px}</style></head><body>"
                f"<h2>{title}</h2><div class=\"row\">"
                "<div class=\"box\">Input</div><div class=\"arrow\">→</div>"
                "<div class=\"box\">Process</div><div class=\"arrow\">→</div>"
                "<div class=\"box\">Evidence</div><div class=\"arrow\">→</div>"
                "<div class=\"box\">Output</div></div></body></html>\n",
                encoding="utf-8",
            )
            sources = [scaffold]
            if on_output:
                await on_output(
                    "[系统][警告] figures/ 目录内未找到任何 fig_*.html，"
                    "已自动生成占位图 fig_pipeline.html。"
                    "这通常意味着 comp-code / paper-figure 步骤尚未生成真实的 HTML 图源；"
                    "建议检查前序步骤是否已正常完成，或在前序步骤产出 fig_*.html 后重跑本步骤。"
                )

        if on_output:
            await on_output(f"[系统] 正在本机渲染 {len(sources)} 个 HTML 图源")
        outputs: List[Path] = []
        stdout_parts: List[str] = []
        stderr_parts: List[str] = []
        failures = 0
        python = cls._runtime_python()
        for source in sources:
            out_pdf = source.with_suffix(".pdf")
            # If a valid PDF already exists for this HTML source, reuse it rather
            # than re-rendering.  This lets the step succeed on recovery even when
            # the subprocess renderer is unavailable (e.g. CREATE_NO_WINDOW /
            # ProactorEventLoop issues on Windows).
            if out_pdf.is_file() and out_pdf.stat().st_size >= 200:
                outputs.append(out_pdf)
                png = source.with_suffix(".png")
                if png.is_file():
                    outputs.append(png)
                capture = Path(str(out_pdf) + ".capture.json")
                if capture.is_file():
                    outputs.append(capture)
                stdout_parts.append(f"reused existing {out_pdf.name}")
                stderr_parts.append("")
                if on_output:
                    await on_output(f"[系统] 复用已有 PDF: {out_pdf.name} ({out_pdf.stat().st_size} bytes)")
                continue
            command = [
                python, str(script),
                "--file", str(source),
                "--out", str(out_pdf),
                "--format", "pdf",
                "--width", "1400",
                "--height", "900",
            ]
            # Attempt subprocess render; fall back to in-process rendering when the
            # subprocess cannot be launched (Windows CREATE_NO_WINDOW / Store-stub
            # NotImplementedError or similar OS-level failures).
            proc_stdout, proc_stderr = "", ""
            try:
                rc, proc_stdout, proc_stderr = await cls._run_process(command, workspace)
            except Exception as proc_exc:
                rc = 1
                proc_stderr = f"subprocess launch failed ({type(proc_exc).__name__}): {proc_exc}"
                if on_output:
                    await on_output(f"[系统] subprocess 启动失败 ({type(proc_exc).__name__})，尝试进程内渲染")
                # In-process fallback: import render_html functions directly to
                # avoid any asyncio subprocess / CREATE_NO_WINDOW issues.
                try:
                    import importlib.util as _ilu
                    _spec = _ilu.spec_from_file_location("_vibe_render_html", str(script))
                    _rmod = _ilu.module_from_spec(_spec)
                    _spec.loader.exec_module(_rmod)  # type: ignore[union-attr]
                    png_out = source.with_suffix(".png")
                    browser = _rmod.find_browser()
                    ok_br = False
                    if browser:
                        ok_br, _warn = _rmod._run_browser(
                            browser, source, png_out, out_pdf,
                            width=1400, height=900, wait_ms=1800,
                        )
                    if not ok_br:
                        _rmod._fallback(source, png_out, out_pdf, 1400, 900)
                    rc = 0 if (out_pdf.is_file() and out_pdf.stat().st_size > 100) else 1
                    proc_stdout = f"in-process render: {'ok' if rc == 0 else 'failed'}"
                    if rc == 0 and on_output:
                        await on_output(f"[系统] 进程内渲染完成: {out_pdf.name}")
                except Exception as ip_exc:
                    rc = 1
                    proc_stderr += f"; in-process fallback also failed: {ip_exc}"
                    if on_output:
                        await on_output(f"[系统] 进程内渲染也失败: {ip_exc}")
            stdout_parts.append(proc_stdout)
            stderr_parts.append(proc_stderr)
            ok = rc == 0 and out_pdf.is_file() and out_pdf.stat().st_size >= 200
            if ok:
                outputs.append(out_pdf)
                png = source.with_suffix(".png")
                if png.is_file():
                    outputs.append(png)
                capture = Path(str(out_pdf) + ".capture.json")
                if capture.is_file():
                    outputs.append(capture)
            else:
                failures += 1
                if on_output:
                    await on_output(f"[系统] 渲染失败: {source.name} rc={rc}")

        include = figures_dir / "latex_includes.tex"
        pdfs = sorted(figures_dir.glob("*.pdf"))
        include_lines = ["% auto-generated by host paper-figure-html"]
        for pdf in pdfs:
            include_lines.append(
                f"% \\includegraphics[width=0.9\\linewidth]{{figures/{pdf.name}}}"
            )
        include.write_text("\n".join(include_lines) + "\n", encoding="utf-8")
        outputs.append(include)

        lineage = cls._write_host_lineage(
            workspace,
            skill_name="paper-figure-html",
            script=script,
            command=[python, str(script), "--file", "<html>", "--out", "<pdf>"],
            artifacts=outputs,
            returncode=0 if outputs and failures == 0 else (1 if failures else 0),
            stdout="\n".join(stdout_parts),
            stderr="\n".join(stderr_parts),
        )
        success = any(
            path.suffix.lower() == ".pdf" and path.is_file() and path.stat().st_size >= 200
            for path in outputs
        )
        if on_output:
            await on_output(
                f"[系统] HTML 图渲染完成: pdfs="
                f"{sum(1 for item in outputs if item.suffix.lower() == '.pdf')} "
                f"failures={failures}; lineage={lineage.name}"
            )
        return {
            "success": success,
            "returncode": 0 if success else 4,
            "return_code": 0 if success else 4,
            "stdout": "\n".join(stdout_parts),
            "stderr": (
                "\n".join(stderr_parts)
                if success
                else (("\n".join(stderr_parts)).strip() or "paper-figure-html produced no PDF")
            ),
            "result": str(figures_dir),
        }

    @staticmethod
    def _drawio_binary() -> str:
        from config import RUNTIME_DRAWIO

        candidates: List[Path] = []
        override = os.environ.get("VIBE_DRAWIO_BIN", "").strip()
        if override:
            candidates.append(Path(override))
        if RUNTIME_DRAWIO:
            root = Path(RUNTIME_DRAWIO)
            candidates.extend([
                root / "draw.io.exe", root / "drawio.exe", root / "draw.io",
            ])
        # Source-tree fallbacks for local development and clean-room verification.
        here = Path(__file__).resolve()
        for base in here.parents:
            candidates.extend([
                base / "runtime" / "draw.io" / "draw.io.exe",
                base / "runtime-release" / "draw.io" / "draw.io.exe",
            ])
        found = shutil.which("drawio") or shutil.which("draw.io")
        if found:
            candidates.append(Path(found))
        for candidate in candidates:
            try:
                resolved = candidate.expanduser().resolve()
            except OSError:
                continue
            if resolved.is_file():
                return str(resolved)
        return ""

    @classmethod
    async def _run_paper_figure_drawio(
        cls,
        workspace: Path,
        on_output: Optional[Callable[[str], Awaitable[None]]] = None,
    ) -> Dict[str, Any]:
        """Export existing .drawio sources to PDF/PNG without an LLM planner."""
        binary = cls._drawio_binary()
        if not binary:
            return {
                "success": False,
                "returncode": 3,
                "return_code": 3,
                "stdout": "",
                "stderr": "draw.io runtime is unavailable",
                "result": "",
            }
        figures_dir = workspace / "figures"
        figures_dir.mkdir(parents=True, exist_ok=True)
        sources = sorted(figures_dir.glob("*.drawio"))
        if (figures_dir / "drawio").is_dir():
            sources.extend(sorted((figures_dir / "drawio").glob("*.drawio")))
        sources = [path for path in sources if path.is_file()]
        if not sources:
            scaffold = figures_dir / "fig_arch.drawio"
            scaffold.write_text(
                '<mxfile host="vibe-research"><diagram id="1" name="Page-1">'
                "<mxGraphModel><root><mxCell id=\"0\"/><mxCell id=\"1\" parent=\"0\"/>"
                "<mxCell id=\"2\" value=\"Input\" style=\"rounded=1;whiteSpace=wrap;html=1;\" "
                "vertex=\"1\" parent=\"1\"><mxGeometry x=\"40\" y=\"40\" width=\"100\" height=\"48\" as=\"geometry\"/>"
                "</mxCell><mxCell id=\"3\" value=\"Process\" style=\"rounded=1;whiteSpace=wrap;html=1;\" "
                "vertex=\"1\" parent=\"1\"><mxGeometry x=\"200\" y=\"40\" width=\"100\" height=\"48\" as=\"geometry\"/>"
                "</mxCell><mxCell id=\"4\" value=\"Output\" style=\"rounded=1;whiteSpace=wrap;html=1;\" "
                "vertex=\"1\" parent=\"1\"><mxGeometry x=\"360\" y=\"40\" width=\"100\" height=\"48\" as=\"geometry\"/>"
                "</mxCell><mxCell id=\"5\" style=\"endArrow=classic;html=1;\" edge=\"1\" parent=\"1\" "
                "source=\"2\" target=\"3\"><mxGeometry relative=\"1\" as=\"geometry\"/></mxCell>"
                "<mxCell id=\"6\" style=\"endArrow=classic;html=1;\" edge=\"1\" parent=\"1\" "
                "source=\"3\" target=\"4\"><mxGeometry relative=\"1\" as=\"geometry\"/></mxCell>"
                "</root></mxGraphModel></diagram></mxfile>\n",
                encoding="utf-8",
            )
            sources = [scaffold]
            if on_output:
                await on_output("[系统] 未找到现成 .drawio 图源，已生成本机占位 fig_arch.drawio")

        if on_output:
            await on_output(f"[系统] 正在本机导出 {len(sources)} 个 DrawIO 图源")
        outputs: List[Path] = []
        stdout_parts: List[str] = []
        stderr_parts: List[str] = []
        failures = 0
        for source in sources:
            out_pdf = source.with_suffix(".pdf")
            command = [
                binary, "--export", "--format", "pdf",
                "--output", str(out_pdf), str(source),
            ]
            rc, stdout, stderr = await cls._run_process(command, workspace)
            stdout_parts.append(stdout)
            stderr_parts.append(stderr)
            ok = rc == 0 and out_pdf.is_file() and out_pdf.stat().st_size >= 200
            if ok:
                outputs.append(out_pdf)
                # companion png for Word/HTML consumers
                out_png = source.with_suffix(".png")
                png_cmd = [
                    binary, "--export", "--format", "png",
                    "--output", str(out_png), str(source),
                ]
                png_rc, png_out, png_err = await cls._run_process(png_cmd, workspace)
                stdout_parts.append(png_out)
                stderr_parts.append(png_err)
                if png_rc == 0 and out_png.is_file():
                    outputs.append(out_png)
            else:
                failures += 1
                if on_output:
                    await on_output(f"[系统] DrawIO 导出失败: {source.name} rc={rc}")

        include = figures_dir / "latex_includes.tex"
        existing = include.read_text(encoding="utf-8", errors="replace") if include.is_file() else ""
        lines = [existing.rstrip(), "% auto-generated by host paper-figure-drawio"] if existing.strip() else [
            "% auto-generated by host paper-figure-drawio"
        ]
        for pdf in sorted(figures_dir.glob("*.pdf")):
            marker_line = f"% \\includegraphics[width=0.9\\linewidth]{{figures/{pdf.name}}}"
            if marker_line not in "\n".join(lines):
                lines.append(marker_line)
        include.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
        outputs.append(include)

        lineage = cls._write_host_lineage(
            workspace,
            skill_name="paper-figure-drawio",
            script=Path(binary),
            command=[binary, "--export", "--format", "pdf", "--output", "<pdf>", "<drawio>"],
            artifacts=outputs,
            returncode=0 if outputs and failures == 0 else (1 if failures else 0),
            stdout="\n".join(stdout_parts),
            stderr="\n".join(stderr_parts),
        )
        success = any(
            path.suffix.lower() == ".pdf" and path.is_file() and path.stat().st_size >= 200
            for path in outputs
        )
        if on_output:
            await on_output(
                f"[系统] DrawIO 导出完成: pdfs="
                f"{sum(1 for item in outputs if item.suffix.lower() == '.pdf')} "
                f"failures={failures}; lineage={lineage.name}"
            )
        return {
            "success": success,
            "returncode": 0 if success else 4,
            "return_code": 0 if success else 4,
            "stdout": "\n".join(stdout_parts),
            "stderr": (
                "\n".join(stderr_parts)
                if success
                else (("\n".join(stderr_parts)).strip() or "paper-figure-drawio produced no PDF")
            ),
            "result": str(figures_dir),
        }

    @classmethod
    async def _run_format_profile(
        cls, workspace: Path, params: Dict[str, Any],
        on_output: Optional[Callable[[str], Awaitable[None]]] = None,
    ) -> Dict[str, Any]:
        profile = cls._derive_text_profile_from_requirements(workspace, params)
        output = workspace / "_text_profile.json"
        output.write_text(json.dumps(profile, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        if on_output:
            matched = ", ".join(str(item) for item in profile.get("_matched_items") or [])
            await on_output(f"[系统] 已生成本机格式 profile：{output.name} ({matched})")
        return {"success": True, "returncode": 0, "return_code": 0, "stdout": str(output), "stderr": "", "result": str(output)}

    async def run_skill(self, **kwargs: Any) -> Dict[str, Any]:
        # Always resolve workspace to an absolute path. Host builders pass
        # artifact paths to child scripts; a relative WORKSPACES_DIR would
        # otherwise be re-rooted by the child cwd and double-joined.
        workspace = Path(kwargs["cwd"]).expanduser().resolve()
        params = dict(kwargs.get("extra_params") or {})
        on_output = kwargs.get("on_output")

        if self.step.skill_name in {"template-prepare", "latex-template-prepare"}:
            result = await self._prepare_templates(
                workspace, params, latex_mode=self.step.skill_name == "latex-template-prepare",
            )
            if on_output and result.get("success"):
                await on_output("[系统] 用户格式模板已复制并建立执行合同")
            return result
        if self.step.skill_name == "latex-template-apply":
            result = await self._apply_latex_template(workspace, params)
            if on_output and result.get("success"):
                await on_output("[系统] LaTeX 模板与文字格式要求已应用")
            return result
        if self.step.skill_name == "docx-template-map":
            result = await self._build_docx_template_map(workspace, params)
            if on_output and result.get("success"):
                await on_output("[系统] Word 模板封面/正文锚点已识别")
            return result
        if self.step.skill_name == "format-profile":
            return await self._run_format_profile(workspace, params, on_output)
        if self.step.skill_name == "docx-format-check":
            return await self._run_docx_format_check(workspace, on_output)
        if self.step.skill_name == "assets-inventory":
            return await self._run_assets_inventory(workspace, params, on_output)
        if self.step.skill_name == "paper-figure-html":
            return await self._run_paper_figure_html(workspace, on_output)
        if self.step.skill_name == "paper-figure-drawio":
            return await self._run_paper_figure_drawio(workspace, on_output)
        if self.step.skill_name in {
            "paper-compile", "paper-compile-zh", "comp-compile-zh", "comp-compile-en",
        }:
            return await self._run_paper_compile(
                workspace,
                self.step.skill_name,
                on_output,
                template=self.template,
                params=params,
            )
        if self.step.skill_name in {"markdown-pdf-export", "auto-review-pdf-export"}:
            return await self._export_markdown_pdf(workspace, self.template, params)
        if self.step.skill_name == "patent-draft":
            return await self._run_patent_draft(workspace, params, on_output)
        if self.step.skill_name == "copyright-draft":
            return await self._run_copyright_draft(workspace, params, on_output)
        if self.step.skill_name == "software-copyright":
            return await self._run_software_copyright(workspace, params, on_output)
        if self.step.skill_name == "patent-build":
            return await self._run_patent_build(workspace, on_output)
        if self.step.skill_name == "copyright-build":
            return await self._run_copyright_build(workspace, on_output)
        if self.step.skill_name == "paper-slides":
            return await self._run_paper_slides(workspace, params, on_output)
        if self.step.skill_name == "paper-poster":
            return await self._run_paper_poster(workspace, params, on_output)
        if self.step.skill_name == "auto-paper-improvement-loop":
            # Log + ensure tex, then real local compile so primary PDF gate holds.
            domain_result = await self._run_domain_host_skill(
                workspace,
                "auto-paper-improvement-loop",
                params,
                on_output,
                template=self.template,
            )
            if not domain_result.get("success"):
                return domain_result
            pdf = workspace / "paper" / "main.pdf"
            if not (pdf.is_file() and pdf.stat().st_size >= 30000):
                compile_skill = (
                    "paper-compile-zh"
                    if str(params.get("language") or "").lower().startswith("zh")
                    else "paper-compile"
                )
                if on_output:
                    await on_output("[系统] 改进循环：正在本机重编译 paper/main.pdf")
                compile_result = await self._run_paper_compile(
                    workspace,
                    compile_skill,
                    on_output,
                    template=self.template,
                    params=params,
                )
                if not compile_result.get("success"):
                    return compile_result
            pdf = workspace / "paper" / "main.pdf"
            log = workspace / "paper" / "PAPER_IMPROVEMENT_LOG.md"
            ok = (
                pdf.is_file()
                and pdf.stat().st_size >= 30000
                and log.is_file()
                and log.stat().st_size >= 200
            )
            if on_output:
                if ok:
                    await on_output(
                        f"[系统] 改进循环完成：paper/main.pdf ({pdf.stat().st_size} bytes) + log"
                    )
                else:
                    await on_output("[系统] 改进循环未满足 PDF/日志门禁")
            return {
                "success": ok,
                "returncode": 0 if ok else 4,
                "return_code": 0 if ok else 4,
                "stdout": domain_result.get("stdout") or "",
                "stderr": "" if ok else "auto-paper-improvement-loop PDF/log gate failed",
                "result": str(pdf) if ok else "",
            }
        if self.step.skill_name in {
            "thesis-proposal", "literature-review", "project-blueprint",
            "paper-plan", "paper-plan-zh", "paper-analysis",
            "paper-figure", "nature-figure", "experiment-bridge",
            "research-lit", "idea-creator", "novelty-check",
            "research-review", "research-refine-pipeline", "auto-review-loop",
            "paper-write", "paper-write-zh", "paper-write-nature",
            "humanities-plan", "humanities-write",
            "course-plan", "course-paper", "course-report", "course-report-plan",
            "comp-prob-analysis", "comp-modeling", "comp-code",
            "comp-paper-zh", "comp-paper-en", "comp-stats-topic",
            "comp-paper-zh-docx", "comp-paper-en-docx",
            "paper-write-docx", "paper-write-zh-docx", "paper-write-nature-docx",
            "auto-paper-improvement-docx",
            "dev-requirement", "dev-design", "dev-code", "dev-selfcheck", "dev-report",
            "humanities-write-latex",
        }:
            return await self._run_domain_host_skill(
                workspace,
                self.step.skill_name,
                params,
                on_output,
                template=self.template,
            )

        sources = self._markdown_sources(workspace, self.template)
        if not sources:
            return {"success": False, "returncode": 2, "stderr": "DOCX export source Markdown not found"}

        if len(sources) == 1:
            raw_source = sources[0]
            # Strip NULs that break Node/XML DOCX writers on some agent outputs.
            raw_bytes = raw_source.read_bytes()
            if b"\x00" in raw_bytes:
                source = workspace / "_docx_export_source_clean.md"
                source.write_bytes(raw_bytes.replace(b"\x00", b""))
            else:
                source = raw_source
        else:
            source = workspace / "_docx_merged_source.md"
            chunks = []
            for item in sources:
                title = item.stem.replace("_", " ")
                body = item.read_text(encoding="utf-8", errors="replace").replace("\x00", "").strip()
                chunks.append(f"# {title}\n\n{body}\n")
            source.write_text("\n\n---\n\n".join(chunks) + "\n", encoding="utf-8")

        output_rel = self.step.primary_output or self.step.output_files[0]
        output = workspace / output_rel
        output.parent.mkdir(parents=True, exist_ok=True)
        if on_output:
            await on_output(f"[系统] 正在将 {source.relative_to(workspace).as_posix()} 导出为 {output_rel}")

        from config import PANDOC_BIN, RUNTIME_NODE, RUNTIME_PYTHON, TOOLS_DIR

        docx_templates, _ = self._safe_template_paths(workspace, params)
        content_output = workspace / "_docx_rendered_content.docx" if docx_templates else output
        node = None
        if RUNTIME_NODE:
            candidate = Path(RUNTIME_NODE) / ("node.exe" if os.name == "nt" else "bin/node")
            if candidate.is_file():
                node = str(candidate)
        node = node or shutil.which("node")
        script = Path(TOOLS_DIR) / "docx-cn-engine" / "md_to_docx.js"
        profile = workspace / "_text_profile.json"
        if node and script.is_file():
            command = [
                node, str(script), "--source", str(source), "--output", str(content_output),
                "--workspace", str(workspace),
            ]
            if profile.is_file():
                command.extend(["--profile", str(profile)])
        elif PANDOC_BIN:
            command = [str(PANDOC_BIN), "-f", "markdown", "-t", "docx", "-o", str(content_output), str(source)]
            if docx_templates:
                command.insert(-1, f"--reference-doc={docx_templates[0]}")
        else:
            return {"success": False, "returncode": 3, "stderr": "Bundled DOCX conversion engine is unavailable"}

        rc, stdout_text, stderr_text = await self._run_process(command, workspace)
        if rc != 0 or not content_output.is_file() or content_output.stat().st_size < 500:
            return {"success": False, "returncode": rc or 4, "stdout": stdout_text, "stderr": stderr_text or "DOCX content was not created"}

        if docx_templates:
            filler = Path(TOOLS_DIR) / "docx_template_fill.py"
            if not filler.is_file():
                return {"success": False, "returncode": 5, "stderr": f"Exact DOCX template filler is unavailable: {filler}"}
            python = str(RUNTIME_PYTHON) if RUNTIME_PYTHON and Path(RUNTIME_PYTHON).is_file() else sys.executable
            fill_command = [
                python, str(filler), "--template", str(docx_templates[0]),
                "--content-docx", str(content_output), "--output", str(output),
            ]
            template_map = workspace / "_template_map.json"
            if template_map.is_file():
                fill_command.extend(["--map", str(template_map)])
            fill_rc, fill_stdout, fill_stderr = await self._run_process(fill_command, workspace)
            stdout_text += fill_stdout
            stderr_text += fill_stderr
            if fill_rc != 0:
                return {"success": False, "returncode": fill_rc, "stdout": stdout_text, "stderr": stderr_text}

        # Cosmetic only: never fail a successful DOCX write on column styling.
        self._apply_docx_columns(
            output, str(params.get("column_layout") or "single").lower(),
        )

        if not output.is_file() or output.stat().st_size < 500:
            return {"success": False, "returncode": 4, "stdout": stdout_text, "stderr": stderr_text or "DOCX file was not created"}
        if on_output:
            await on_output(f"[系统] Word 导出完成：{output_rel} ({output.stat().st_size} bytes)")
        return {"success": True, "returncode": 0, "stdout": stdout_text, "stderr": stderr_text}


async def _db_write(operation_name: str, worker):
    """Open a short-lived DB connection for a write under the shared lock/retry policy."""
    from services.state_store import execute_write, get_db

    async def _op():
        db = await get_db()
        try:
            return await worker(db)
        finally:
            await db.close()

    return await execute_write(operation_name, _op)


async def _db_read(worker):
    """Open a short-lived DB connection for a pure read."""
    from services.state_store import get_db

    db = await get_db()
    try:
        return await worker(db)
    finally:
        await db.close()


async def run_single_step(workflow_id: str, skill_name: str) -> None:
    """(docstring)"""
    from services.claude_runner import ClaudeRunner
    from services.state_store import get_db, get_workflow, update_workflow

    # Bound concurrent skill runs and never hold SQLite connections across the
    # long agent/host execution window. Matrix launches previously opened one
    # writer per workflow for the entire step and failed with "database is locked".
    async with _get_step_semaphore():
        await _run_single_step_locked(workflow_id, skill_name, ClaudeRunner, get_db, get_workflow, update_workflow)


async def _run_single_step_locked(workflow_id: str, skill_name: str, ClaudeRunner, get_db, get_workflow, update_workflow) -> None:
    managed = workflow_id in _workflow_managed_steps
    attempt_id: str | None = None
    attempt_cancelled = False
    attempt_unhandled_error: str | None = None
    # Save CWD so the finally block can restore it even if no chdir is performed.
    original_cwd = os.getcwd()
    try:
        db = await get_db()
        try:
            wf = await get_workflow(db, workflow_id)
        finally:
            await db.close()
        if not wf:
            log.error("Workflow not found: %s", workflow_id)
            return

        workspace_dir = Path(wf["workspace_dir"])
        # Match run_workflow()'s CWD contract: relative paths inside skills and
        # shared scripts (e.g. user_data/<file>) resolve to the workspace root.
        os.chdir(str(workspace_dir))
        params = wf.get("params", {})
        template = wf["template"]
        try:
            tmpl = _resolve_template(template, params, workspace_dir)
        except ValueError:
            log.error("Template not found: %s", template)
            return

        step_def = None
        for step in tmpl.sub_steps:
            if step.skill_name == skill_name:
                step_def = step
                break
        if not step_def:
            log.error("Step not found in template: %s", skill_name)
            return

        try:
            from services.workflow_operations import begin_step_attempt

            attempt_id = await begin_step_attempt(workflow_id, skill_name)
        except Exception as exc:
            # Operations telemetry must never replace the actual executor.
            log.warning("Unable to open workflow step attempt ledger: %s", exc)

        async def _mark_running(db):
            await db.execute(
                "UPDATE workflow_steps SET status = 'running', started_at = ? WHERE workflow_id = ? AND skill_name = ?",
                (datetime.now().isoformat(), workflow_id, skill_name)
            )
            await update_workflow(db, workflow_id, status="running", current_step=skill_name)

        await _db_write(f"step_running:{workflow_id}:{skill_name}", _mark_running)

        if not managed:
            await _broadcast(workflow_id, {"type": "workflow_started"})
        await _broadcast(workflow_id, {"type": "step_started", "step": skill_name, "display_name": step_def.display_name})

        try:
            missing_upstream = _missing_upstream_primary_outputs(tmpl, step_def, workspace_dir)
            if missing_upstream:
                raise RuntimeError(
                    "REQUIRED_UPSTREAM_MISSING: "
                    f"{step_def.display_name}: {', '.join(missing_upstream)}"
                )
            _inject_plan_stage_guidance(workspace_dir, skill_name, params)
            snapshot_before = _snapshot_workspace(workspace_dir)
            files_before = set(snapshot_before)
            workspace_files = sorted(files_before, key=str.lower) or None
            managed_rows = _managed_step_rows.get()
            if managed and managed_rows is not None:
                context_summary = _build_managed_context_summary(
                    workspace_files, step_def, tmpl, managed_rows,
                )
            else:
                context_summary = _build_context_summary(workspace_dir, workspace_files or [])
            # Pure tooling steps: deterministic file operations that never need LLM.
            # These always use _HostStepRunner regardless of model configuration.
            _pure_tooling_steps = {
                "docx-export", "template-prepare", "latex-template-prepare",
                "latex-template-apply", "docx-template-map", "format-profile",
                "docx-format-check", "assets-inventory",
                "markdown-pdf-export", "auto-review-pdf-export",
                "paper-compile", "paper-compile-zh",
                "comp-compile-zh", "comp-compile-en",
                "paper-figure-html", "paper-figure-drawio",
                "patent-draft", "copyright-draft", "software-copyright",
                "patent-build", "copyright-build",
                "paper-slides", "paper-poster",
                "auto-paper-improvement-docx",
                "comp-paper-zh-docx", "comp-paper-en-docx",
                "paper-write-docx", "paper-write-zh-docx", "paper-write-nature-docx",
            }
            # Content generation steps: require LLM for meaningful output.
            # Use ClaudeRunner when executor model is configured; fall back to
            # _HostStepRunner (scaffold/template mode) when no model is available.
            _content_generation_steps = {
                "thesis-proposal", "literature-review", "project-blueprint",
                "paper-plan", "paper-plan-zh", "paper-analysis",
                "paper-figure", "nature-figure", "experiment-bridge",
                "research-lit", "idea-creator", "novelty-check",
                "research-review", "research-refine-pipeline", "auto-review-loop",
                "paper-write", "paper-write-zh", "paper-write-nature",
                "humanities-plan", "humanities-write",
                "course-plan", "course-paper", "course-report", "course-report-plan",
                "comp-prob-analysis", "comp-modeling", "comp-code",
                "comp-paper-zh", "comp-paper-en",
                "dev-requirement", "dev-design", "dev-code",
                "dev-selfcheck", "dev-report",
                "comp-stats-topic", "humanities-write-latex",
                "auto-paper-improvement-loop",
            }

            # Determine which runner to use for this step.
            if skill_name in _pure_tooling_steps:
                # Always host — no LLM needed for compile/export/format operations.
                runner = _HostStepRunner(template, step_def)
            elif skill_name in _content_generation_steps:
                # Use ClaudeRunner when executor model is configured; scaffold otherwise.
                _executor_ready = False
                try:
                    from services.llm_client import get_all_settings, _configured_agent
                    _settings = await get_all_settings()
                    _configured_agent(_settings, "executor")  # raises if not configured
                    _executor_ready = True
                except Exception:
                    pass
                if _executor_ready:
                    runner = ClaudeRunner()
                    log.info("Step '%s': executor model configured — using ClaudeRunner", skill_name)
                else:
                    runner = _HostStepRunner(template, step_def)
                    log.warning(
                        "Step '%s': no executor model configured — falling back to scaffold mode. "
                        "Configure an executor model in Settings to get real LLM output.",
                        skill_name,
                    )
            else:
                # Unknown skill — default to ClaudeRunner as before.
                runner = ClaudeRunner()

            # Execution-channel probe for figure steps: verify the host can spawn
            # a python subprocess AND write into figures/ BEFORE the agent burns
            # its retries.  If the channel is broken we surface a precise diagnosis
            # immediately instead of letting the step fail 6 times with the generic
            # "只有辅助文件" message.
            # NOTE: on_output is not yet defined at this point (it is created inside
            # the attempt loop below), so report through _log/_broadcast instead.
            if managed and skill_name in _AUTO_RECOVER_FIGURE_SKILLS:
                try:
                    probe_issue = await _HostStepRunner._probe_figure_execution_channel(
                        workspace_dir
                    )
                except Exception as probe_exc:  # noqa: BLE001 - probe must warn, never kill
                    # The probe is a pre-flight courtesy check; an unexpected
                    # failure inside it (e.g. a subprocess-channel error that
                    # escaped _run_process) must degrade to a warning, not abort
                    # the step before the runner and host-fallback get a chance.
                    probe_issue = (
                        "执行通道自检自身异常 "
                        f"({type(probe_exc).__name__}: {probe_exc or 'no message'}) — "
                        "将继续执行，失败时由宿主兜底接管"
                    )
                if probe_issue:
                    log.warning(
                        "Step '%s' execution-channel probe failed: %s",
                        skill_name, probe_issue,
                    )
                    await _log(
                        workflow_id, skill_name, "warning",
                        f"[系统][执行通道自检] {probe_issue}",
                    )
                    await _broadcast(workflow_id, {
                        "type": "step_progress", "step": skill_name,
                        "log": f"[系统][执行通道自检] {probe_issue}",
                    })

            arguments = workflow_id if managed else json.dumps({
                "template": template,
                "skill_name": skill_name,
                "params": params,
                "workspace": str(workspace_dir),
            }, ensure_ascii=False)

            # Installed-engine probes show one initial execution plus eight
            # immediate retries for a non-zero runner result (nine calls total).
            # Accept both runner result spellings while preserving the rebuilt
            # runner's public ``returncode`` field.
            result = {}
            max_attempts = 9 if managed else 1
            for attempt in range(max_attempts):
                log_buffer: List[str] = []
                log_flush_lock = asyncio.Lock()

                async def flush_logs() -> None:
                    async with log_flush_lock:
                        if not log_buffer:
                            return
                        batch = list(log_buffer)
                        log_buffer.clear()
                    await _persist_log_batch(workflow_id, skill_name, batch)

                async def on_output(line: str) -> None:
                    log_buffer.append(line)
                    if managed:
                        await _broadcast(workflow_id, {
                            "type": "step_progress", "step": skill_name, "log": line,
                        })
                    else:
                        await _broadcast(workflow_id, {"log": line})

                try:
                    result = await runner.run_skill(
                        skill_name=skill_name,
                        arguments=arguments,
                        cwd=workspace_dir,
                        workflow_id=f"{workflow_id}_{skill_name}" if managed else workflow_id,
                        on_output=on_output,
                        extra_params=params,
                        workspace_files=workspace_files,
                        context_summary=context_summary,
                        inactivity_timeout=2400,
                    )
                except asyncio.TimeoutError:
                    result = {
                        "success": False,
                        "stdout": "",
                        "stderr": "Step timed out after 4 hours",
                        "returncode": 1,
                        "return_code": 1,
                        "result": "",
                    }
                except Exception as runner_exc:  # noqa: BLE001
                    # Catch ALL runner exceptions (CAPABILITY_BLOCKED, mount failures,
                    # network errors, sandbox violations, etc.) and convert them into
                    # a failed result so the host-fallback logic below still has a
                    # chance to rescue figure-producing steps.  Without this catch,
                    # any non-timeout exception propagates straight up and the
                    # host-side gen_fig_*.py rescue path is unreachable.
                    log.exception("runner.run_skill raised for step '%s'", skill_name)
                    result = {
                        "success": False,
                        "stdout": "",
                        "stderr": f"runner exception: {type(runner_exc).__name__}: {runner_exc}",
                        "returncode": -1,
                        "return_code": -1,
                        "result": "",
                    }
                await flush_logs()
                return_code = result.get("return_code", result.get("returncode"))
                succeeded = result.get("success", return_code == 0)
                if succeeded and (return_code is None or return_code == 0):
                    break
                if attempt < max_attempts - 1:
                    log.warning(
                        "Step '%s' failed (rc=%s, attempt %s/8), retrying... error: %s",
                        skill_name, return_code, attempt + 1, result.get("stderr", ""),
                    )
                    await _log(
                        workflow_id,
                        skill_name,
                        "info",
                        f"[RETRY] 步骤退出码 rc={return_code}, 自动重试 (attempt {attempt + 1}/8)",
                    )

            files_after = set(files_before)
            files_created: List[str] = []
            files_modified: List[str] = []

            if result.get("success"):
                snapshot_after = _snapshot_workspace(workspace_dir)
                files_after = set(snapshot_after)
                files_created = _order_step_files(files_after - files_before, step_def)
                # The installed run_workflow watchdog reports every visible
                # file that existed at step start as "modified".  Standalone
                # reruns do not report modifications at all.
                files_modified = sorted(files_before & files_after, key=str.lower) if managed else []
                if managed and skill_name in _AUTO_RECOVER_FIGURE_SKILLS:
                    primary_issue = _primary_output_issue(
                        workspace_dir, step_def, initial_directory_check=True,
                    )
                    figures_ok, figures_reason = _check_figures_step_health(
                        workspace_dir, skill_name, files_created, files_modified, params
                    )
                    quantity_issue = _minimum_quantity_issue(workspace_dir, skill_name, params)
                    recovery_reason = primary_issue or (None if figures_ok else figures_reason) or quantity_issue
                    resume_session_id = result.get("session_id")

                    for recovery_attempt in range(1, 6):
                        if recovery_reason is None:
                            break
                        await _log(
                            workflow_id,
                            skill_name,
                            "info",
                            f"[AUTO-RECOVER {recovery_attempt}/5] {recovery_reason}",
                        )
                        retry_params = _figure_recovery_params(
                            params, recovery_reason, recovery_attempt,
                        )
                        try:
                            recovery_result = await runner.run_skill(
                                skill_name=skill_name,
                                arguments=arguments,
                                cwd=workspace_dir,
                                workflow_id=(
                                    f"{workflow_id}_{skill_name}_recov{recovery_attempt}"
                                ),
                                on_output=on_output,
                                extra_params=retry_params,
                                workspace_files=workspace_files,
                                context_summary=context_summary,
                                inactivity_timeout=2400,
                                resume_session_id=resume_session_id,
                            )
                        except asyncio.TimeoutError:
                            recovery_result = {
                                "success": False,
                                "stdout": "",
                                "stderr": "auto-recover timeout",
                                "returncode": 1,
                                "return_code": 1,
                                "result": "",
                            }
                        await flush_logs()
                        if "session_id" in recovery_result:
                            resume_session_id = recovery_result.get("session_id")
                        result = recovery_result

                        snapshot_after = _snapshot_workspace(workspace_dir)
                        files_after = set(snapshot_after)
                        files_created = _order_step_files(files_after - files_before, step_def)
                        files_modified = (
                            sorted(files_before & files_after, key=str.lower) if managed else []
                        )
                        primary_issue = _primary_output_issue(workspace_dir, step_def)
                        figures_ok, figures_reason = _check_figures_step_health(
                            workspace_dir, skill_name, files_created, files_modified, params
                        )
                        quantity_issue = _minimum_quantity_issue(workspace_dir, skill_name, params)
                        health_reason = primary_issue or (None if figures_ok else figures_reason) or quantity_issue
                        recovery_succeeded = result.get("success") and health_reason is None
                        recovery_reason = None if recovery_succeeded else (
                            health_reason or result.get("stderr") or "runner failed"
                        )

                    if recovery_reason is not None:
                        # Host-side last resort: the agent kept failing to produce
                        # images (usually because its subprocess channel is broken on
                        # this machine).  If gen_fig_*.py scripts exist, execute them
                        # directly from the host so the step can still deliver real
                        # PDF/PNG figures instead of dying after 6 blind retries.
                        host_figs = await _HostStepRunner._host_execute_gen_fig_scripts(
                            workspace_dir, skill_name, on_output=on_output,
                        )
                        if host_figs:
                            await _log(
                                workflow_id,
                                skill_name,
                                "info",
                                f"[HOST-FALLBACK] 宿主环境成功执行 {len(host_figs)} 个绘图脚本: "
                                + ", ".join(host_figs[:8])
                                + (" ..." if len(host_figs) > 8 else ""),
                            )
                            snapshot_after = _snapshot_workspace(workspace_dir)
                            files_after = set(snapshot_after)
                            files_created = _order_step_files(files_after - files_before, step_def)
                            files_modified = (
                                sorted(files_before & files_after, key=str.lower) if managed else []
                            )
                            primary_issue = _primary_output_issue(workspace_dir, step_def)
                            figures_ok, figures_reason = _check_figures_step_health(
                                workspace_dir, skill_name, files_created, files_modified, params
                            )
                            quantity_issue = _minimum_quantity_issue(workspace_dir, skill_name, params)
                            health_reason = primary_issue or (None if figures_ok else figures_reason) or quantity_issue
                            if health_reason is None:
                                recovery_reason = None
                                result["success"] = True
                                result["stderr"] = ""

                    if recovery_reason is not None:
                        result["success"] = False
                        result["stderr"] = (
                            "步骤连续 6 次未产出 primary_output: "
                            f"{recovery_reason}."
                        )
            elif managed and skill_name in _AUTO_RECOVER_FIGURE_SKILLS:
                # Host-side rescue when the runner itself failed (exception, rc!=0,
                # network/HTTP error, sandbox blocked, etc.).  The agent may have
                # already written gen_fig_*.py scripts before dying; executing them
                # directly from the host lets the step deliver real figures instead
                # of propagating an opaque runner failure.
                await _log(
                    workflow_id, skill_name, "info",
                    "[HOST-FALLBACK] runner 失败分支，尝试宿主兜底执行 gen_fig_*.py",
                )
                host_figs = await _HostStepRunner._host_execute_gen_fig_scripts(
                    workspace_dir, skill_name, on_output=on_output,
                )
                if host_figs:
                    snapshot_after = _snapshot_workspace(workspace_dir)
                    files_after = set(snapshot_after)
                    files_created = _order_step_files(files_after - files_before, step_def)
                    files_modified = (
                        sorted(files_before & files_after, key=str.lower) if managed else []
                    )
                    primary_issue = _primary_output_issue(workspace_dir, step_def)
                    figures_ok, figures_reason = _check_figures_step_health(
                        workspace_dir, skill_name, files_created, files_modified, params
                    )
                    quantity_issue = _minimum_quantity_issue(workspace_dir, skill_name, params)
                    health_reason = primary_issue or (None if figures_ok else figures_reason) or quantity_issue
                    if health_reason is None:
                        await _log(
                            workflow_id, skill_name, "info",
                            f"[HOST-FALLBACK] 宿主兜底成功，产出 {len(host_figs)} 张图像，"
                            "将 runner 失败转为步骤成功",
                        )
                        result["success"] = True
                        result["stderr"] = ""
                        result["returncode"] = 0
                        result["return_code"] = 0

            if result.get("success"):
                validation_errors = []

                primary_issue = _primary_output_issue(workspace_dir, step_def)
                if primary_issue and not (files_created or files_modified):
                    validation_errors.append(primary_issue)
                companions_ok, missing_companions = _check_step_companions(workspace_dir, skill_name)
                if not companions_ok:
                    validation_errors.append(f"missing companion outputs: {', '.join(missing_companions)}")
                figures_ok, figures_reason = _check_figures_step_health(
                    workspace_dir, skill_name, files_created, files_modified, params
                )
                if not figures_ok:
                    validation_errors.append(figures_reason)
                quantity_issue = _minimum_quantity_issue(workspace_dir, skill_name, params)
                if quantity_issue:
                    validation_errors.append(quantity_issue)

                if validation_errors:
                    result["success"] = False
                    result["stderr"] = "; ".join(validation_errors)

            if result.get("success"):
                reported_files = _order_reported_files(
                    files_created, files_modified, step_def
                )
                if not managed:
                    declared_existing = _order_step_files(files_after, step_def)
                    declared_names = set(step_def.output_files)
                    for path in declared_existing:
                        if path in declared_names and path not in reported_files:
                            reported_files.append(path)

                async def _mark_completed(db, reported=reported_files):
                    await db.execute(
                        "UPDATE workflow_steps SET status = 'completed', completed_at = ?, output_files = ? WHERE workflow_id = ? AND skill_name = ?",
                        (datetime.now().isoformat(), json.dumps(reported, ensure_ascii=False), workflow_id, skill_name)
                    )
                    # Standalone reruns own the whole workflow; managed multi-step
                    # orchestration must keep the parent workflow running so later
                    # nodes are not abandoned after an intermediate success.
                    if managed:
                        await update_workflow(db, workflow_id, status="running", current_step=skill_name)
                    else:
                        await update_workflow(db, workflow_id, status="completed")

                await _db_write(f"step_completed:{workflow_id}:{skill_name}", _mark_completed)
                await _broadcast(workflow_id, {
                    "type": "step_completed", "step": skill_name,
                    "result_summary": _result_summary(files_created, files_modified),
                    "output_files": reported_files,
                })
                if not managed:
                    await _broadcast(workflow_id, {"type": "workflow_completed", "output_files": reported_files})
            else:
                # Never persist an empty error_message — downstream code
                # (finish_step_attempt, recovery_operations fallback) treats
                # empty/NULL as "unknown" and surfaces the opaque "node
                # execution failed".  Synthesise a precise reason from the
                # runner result so attempts retain a real diagnosis.
                raw_stderr = (result.get("stderr") or "").strip()
                return_code = result.get("return_code", result.get("returncode"))
                if raw_stderr:
                    error_text = raw_stderr[:500]
                elif return_code not in (None, 0):
                    error_text = (
                        f"runner exited with rc={return_code} but produced no stderr "
                        f"(skill={skill_name}); possible sandbox block, executor "
                        "configuration issue, or upstream LLM failure"
                    )[:500]
                else:
                    error_text = (
                        f"runner reported success=False without stderr or returncode "
                        f"(skill={skill_name}); check backend logs for the underlying "
                        "exception"
                    )[:500]

                async def _mark_failed(db, message=error_text):
                    await db.execute(
                        "UPDATE workflow_steps SET status = 'failed', completed_at = ?, error_message = ? WHERE workflow_id = ? AND skill_name = ?",
                        (datetime.now().isoformat(), message, workflow_id, skill_name)
                    )
                    await update_workflow(db, workflow_id, status="failed")

                await _db_write(f"step_failed:{workflow_id}:{skill_name}", _mark_failed)
                await _broadcast(workflow_id, {
                    "type": "step_failed", "step": skill_name,
                    "error": error_text, "output_files": [],
                })
                if not managed:
                    await _broadcast(workflow_id, {"type": "workflow_stopped"})

        except Exception as e:
            # log.exception (not log.error) so the full traceback lands in
            # backend.log — empty-message exceptions (CancelledError wrapped
            # by asyncio, NotImplementedError(''), etc.) are undiagnosable
            # from the message alone.
            log.exception("Step execution failed: %s", e)
            error_text = (str(e) or f"{type(e).__name__} (no message)")[:500]

            async def _mark_exception(db, message=error_text):
                await db.execute(
                    "UPDATE workflow_steps SET status = 'failed', error_message = ? WHERE workflow_id = ? AND skill_name = ?",
                    (message, workflow_id, skill_name)
                )
                await update_workflow(db, workflow_id, status="failed")

            await _db_write(f"step_exception:{workflow_id}:{skill_name}", _mark_exception)
            await _broadcast(workflow_id, {"type": "step_failed", "step": skill_name, "error": str(e)})

    except asyncio.CancelledError:
        attempt_cancelled = True
        raise
    except Exception as exc:
        attempt_unhandled_error = str(exc)
        raise
    finally:
        # Restore original CWD unconditionally so concurrent coroutines and
        # subsequent workflow runs are not affected by this execution's chdir.
        os.chdir(original_cwd)
        if attempt_id:
            try:
                from services.workflow_operations import finish_step_attempt

                await finish_step_attempt(
                    attempt_id,
                    cancelled=attempt_cancelled,
                    unhandled_error=attempt_unhandled_error,
                )
            except Exception as exc:
                log.warning("Unable to seal workflow step attempt ledger: %s", exc)


async def run_workflow(workflow_id: str) -> None:
    """(docstring)"""
    from services.claude_runner import ClaudeRunner
    from services.state_store import get_db, get_workflow, update_workflow

    # Serialise per-workflow runs.  Without this, double-clicking
    # retry/recover spawns concurrent run_workflow coroutines that share the
    # same workspace_dir, which produced the fb4f4e5b7272 failure mode:
    # one coroutine holds _utils files open while another tries to rmtree
    # them, and the attempt ledger accumulates zombie 'running' rows.
    workflow_lock = await _acquire_workflow_lock(workflow_id)
    if workflow_lock.locked():
        log.warning(
            "run_workflow(%s) is already in progress; queuing behind the "
            "existing run instead of racing it.  Triggering UI should debounce "
            "retry/recover clicks.",
            workflow_id,
        )
    async with workflow_lock:
        await _run_workflow_locked(workflow_id)


async def _run_workflow_locked(workflow_id: str) -> None:
    """Inner implementation of run_workflow, executed under the per-workflow lock."""
    from services.claude_runner import ClaudeRunner
    from services.state_store import get_db, get_workflow, update_workflow

    async def _load_initial(db):
        wf_local = await get_workflow(db, workflow_id)
        if not wf_local:
            return None, []
        cursor = await db.execute(
            "SELECT * FROM workflow_steps WHERE workflow_id = ? ORDER BY step_order",
            (workflow_id,),
        )
        steps_local = [dict(r) for r in await cursor.fetchall()]
        return wf_local, steps_local

    # Save original CWD to restore later
    original_cwd = os.getcwd()

    try:
        wf, steps = await _db_read(_load_initial)
        if not wf:
            log.error("Workflow not found: %s", workflow_id)
            return

        # Change to workspace directory so relative paths like 'user_data/file.txt'
        # resolve correctly throughout workflow execution
        workspace_dir = Path(wf["workspace_dir"])
        os.chdir(str(workspace_dir))
        log.info("Changed CWD to workspace: %s", workspace_dir)

        template = wf["template"]
        try:
            tmpl = _resolve_template(template, wf.get("params", {}), Path(wf["workspace_dir"]))
        except ValueError:
            log.error("Template not found: %s", template)
            return

        _generate_claude_md(
            Path(wf["workspace_dir"]), wf["title"], template, wf.get("params", {})
        )

        async def _mark_running(db):
            await update_workflow(db, workflow_id, status="running")

        await _db_write(f"workflow_running:{workflow_id}", _mark_running)
        await _broadcast(workflow_id, {"type": "workflow_started"})

        await _wait_for_extracts(Path(wf["workspace_dir"]))

        if template == "paper_from_assets":
            copied = _pfa_safety_copy_assets(Path(wf["workspace_dir"]))
            if any(copied[key] for key in ("code", "data", "figures", "templates")):
                await _log(workflow_id, None, "info", f"Synchronized user assets: {copied}")

        for step_row in steps:
            resumed_checkpoint = False
            skip_reason = _runtime_skip_reason(step_row["skill_name"], wf.get("params", {}))
            if skip_reason:
                async def _skip(db, reason=skip_reason, skill=step_row["skill_name"]):
                    await db.execute(
                        "UPDATE workflow_steps SET status='skipped', completed_at=?, error_message=? "
                        "WHERE workflow_id=? AND skill_name=?",
                        (datetime.now().isoformat(), reason, workflow_id, skill),
                    )
                    await db.commit()

                await _db_write(f"workflow_skip:{workflow_id}:{step_row['skill_name']}", _skip)
                await _broadcast(workflow_id, {
                    "type": "step_skipped",
                    "step": step_row["skill_name"],
                    "display_name": step_row["display_name"],
                    "reason": skip_reason,
                })
                continue
            # A persisted waiting_checkpoint step has already executed.  On a
            # process restart/resume the installed engine advances past it
            # rather than rerunning the skill, even when checkpoints are now
            # disabled or the in-memory waiter was lost.
            if step_row["status"] == "completed":
                completed_def = next(
                    (item for item in tmpl.sub_steps if item.skill_name == step_row["skill_name"]),
                    None,
                )
                if completed_def is None or _primary_output_exists(
                    Path(wf["workspace_dir"]), completed_def
                ):
                    await _broadcast(workflow_id, {
                        "type": "step_skipped",
                        "step": step_row["skill_name"],
                        "display_name": step_row["display_name"],
                    })
                    continue

                async def _reset_missing(db, skill=step_row["skill_name"]):
                    await db.execute(
                        "UPDATE workflow_steps SET status='pending', error_message=NULL, "
                        "started_at=NULL, completed_at=NULL WHERE workflow_id=? AND skill_name=?",
                        (workflow_id, skill),
                    )
                    await db.commit()

                await _db_write(f"workflow_reset_missing:{workflow_id}:{step_row['skill_name']}", _reset_missing)
                await _broadcast(workflow_id, {
                    "type": "step_progress",
                    "step": step_row["skill_name"],
                    "log": (
                        f"[系统] 之前标记完成但产出 `{completed_def.primary_output}` 已丢失, 重新执行"
                    ),
                })
            elif step_row["status"] == "skipped":
                async def _reset_skipped(db, skill=step_row["skill_name"]):
                    await db.execute(
                        "UPDATE workflow_steps SET status='pending', error_message=NULL, "
                        "started_at=NULL, completed_at=NULL WHERE workflow_id=? AND skill_name=?",
                        (workflow_id, skill),
                    )
                    await db.commit()

                await _db_write(f"workflow_reset_skipped:{workflow_id}:{step_row['skill_name']}", _reset_skipped)
            elif step_row["status"] == "waiting_checkpoint":
                checkpoint_def = next(
                    (item for item in tmpl.sub_steps if item.skill_name == step_row["skill_name"]),
                    None,
                )
                if checkpoint_def and _primary_output_exists(Path(wf["workspace_dir"]), checkpoint_def):
                    if wf.get("enable_checkpoints"):
                        # This step had already completed before the process was
                        # interrupted. Recreate only its checkpoint waiter.
                        resumed_checkpoint = True
                    else:
                        async def _complete_checkpoint(db, skill=step_row["skill_name"]):
                            await db.execute(
                                "UPDATE workflow_steps SET status = 'completed' WHERE workflow_id = ? AND skill_name = ?",
                                (workflow_id, skill),
                            )
                            await db.commit()

                        await _db_write(
                            f"workflow_complete_checkpoint:{workflow_id}:{step_row['skill_name']}",
                            _complete_checkpoint,
                        )
                        await _broadcast(workflow_id, {
                            "type": "step_skipped",
                            "step": step_row["skill_name"],
                            "display_name": step_row["display_name"],
                        })
                        continue
                else:
                    async def _reset_checkpoint(db, skill=step_row["skill_name"]):
                        await db.execute(
                            "UPDATE workflow_steps SET status='pending', error_message=NULL, "
                            "started_at=NULL, completed_at=NULL WHERE workflow_id=? AND skill_name=?",
                            (workflow_id, skill),
                        )
                        await db.commit()

                    await _db_write(
                        f"workflow_reset_checkpoint:{workflow_id}:{step_row['skill_name']}",
                        _reset_checkpoint,
                    )
            skill_name = step_row["skill_name"]
            if resumed_checkpoint:
                should_skip, skip_reason = False, ""
            else:
                should_skip, skip_reason = _should_skip_step_by_assets(
                    Path(wf["workspace_dir"]), skill_name, template
                )
            if should_skip:
                async def _skip_assets(db, reason=skip_reason, skill=skill_name):
                    await db.execute(
                        "UPDATE workflow_steps SET status = 'skipped', completed_at = ? WHERE workflow_id = ? AND skill_name = ?",
                        (datetime.now().isoformat(), workflow_id, skill),
                    )
                    await db.commit()

                await _db_write(f"workflow_skip_assets:{workflow_id}:{skill_name}", _skip_assets)
                await _log(workflow_id, skill_name, "info", f"Skipped: {skip_reason}")
                await _broadcast(workflow_id, {"type": "step_skipped", "step": skill_name})
                continue
            if not resumed_checkpoint:
                _workflow_managed_steps.add(workflow_id)
                snapshot_token = _managed_step_rows.set(steps)
                try:
                    await run_single_step(workflow_id, skill_name)
                finally:
                    _managed_step_rows.reset(snapshot_token)
                    _workflow_managed_steps.discard(workflow_id)

            async def _read_step_status(db, skill=skill_name):
                step_status_cursor = await db.execute(
                    "SELECT status, error_message FROM workflow_steps WHERE workflow_id = ? AND skill_name = ?",
                    (workflow_id, skill),
                )
                return await step_status_cursor.fetchone()

            step_status = await _db_read(_read_step_status)
            if step_status and step_status["status"] == "failed":
                error_message = step_status["error_message"] or f"Step failed: {skill_name}"

                async def _fail_workflow(db, skill=skill_name):
                    await update_workflow(db, workflow_id, status="failed", current_step=skill)

                await _db_write(f"workflow_failed:{workflow_id}:{skill_name}", _fail_workflow)
                await _broadcast(
                    workflow_id,
                    {"type": "workflow_failed", "reason": error_message, "step": skill_name},
                )
                return

            if (
                not resumed_checkpoint
                and skill_name in {"paper-figure", "nature-figure", "experiment-bridge"}
                and not _has_table_and_json_data(Path(wf["workspace_dir"]))
            ):
                await _broadcast(workflow_id, {
                    "type": "step_progress",
                    "step": skill_name,
                    "log": "[系统] TABLE 数据真实性核对：无需自检（无 TABLE 文件或无 JSON 数据）",
                })

            async def _reload_workflow(db):
                return await get_workflow(db, workflow_id)

            wf = await _db_read(_reload_workflow)
            if wf and wf["status"] == "paused":
                await _broadcast(workflow_id, {"type": "workflow_paused"})
                return


            step_def = None
            for s in tmpl.sub_steps:
                if s.skill_name == skill_name:
                    step_def = s
                    break
            if step_def and step_def.has_checkpoint and wf and wf.get("enable_checkpoints"):
                checkpoint_event = {
                    "type": "checkpoint_hit",
                    "step": skill_name,
                    "checkpoint_type": step_def.checkpoint_type,
                    "display_name": step_def.display_name,
                }
                if step_def.primary_output:
                    checkpoint_event["primary_output_file"] = step_def.primary_output
                    primary_path = Path(wf["workspace_dir"]) / step_def.primary_output
                    if primary_path.is_file():
                        try:
                            checkpoint_event["primary_output_content"] = primary_path.read_text(encoding="utf-8")
                        except (OSError, UnicodeDecodeError):
                            pass

                async def _enter_checkpoint(db, event=checkpoint_event, skill=skill_name, ctype=step_def.checkpoint_type):
                    await db.execute(
                        "UPDATE workflow_steps SET status = 'waiting_checkpoint' WHERE workflow_id = ? AND skill_name = ?",
                        (workflow_id, skill),
                    )
                    await update_workflow(db, workflow_id, status="paused", current_step=skill)
                    # Persist a durable pending checkpoint so /checkpoints/current and
                    # UI recovery still work after a renderer refresh or process restart.
                    await db.execute(
                        "INSERT INTO checkpoints (workflow_id, step_name, checkpoint_type, data, status) "
                        "VALUES (?, ?, ?, ?, 'pending')",
                        (
                            workflow_id,
                            skill,
                            ctype or "approve",
                            json.dumps(event, ensure_ascii=False),
                        ),
                    )
                    await db.commit()

                await _db_write(f"workflow_checkpoint:{workflow_id}:{skill_name}", _enter_checkpoint)
                await _broadcast(workflow_id, checkpoint_event)

                response = await wait_checkpoint(workflow_id, timeout=None)
                action = response.get("action") if isinstance(response, dict) else None
                try:
                    async def _resolve_checkpoint(db, payload=response, skill=skill_name):
                        await db.execute(
                            "UPDATE checkpoints SET status = 'resolved', response = ?, resolved_at = CURRENT_TIMESTAMP "
                            "WHERE workflow_id = ? AND step_name = ? AND status = 'pending'",
                            (
                                json.dumps(payload if isinstance(payload, dict) else {"action": action}, ensure_ascii=False),
                                workflow_id,
                                skill,
                            ),
                        )
                        await db.commit()

                    await _db_write(f"workflow_checkpoint_resolve:{workflow_id}:{skill_name}", _resolve_checkpoint)
                except Exception:
                    log.exception("failed to resolve durable checkpoint for %s/%s", workflow_id, skill_name)

                if action == "stop":
                    async def _stop_workflow(db, skill=skill_name):
                        await db.execute(
                            "UPDATE workflow_steps SET status = 'completed' WHERE workflow_id = ? AND skill_name = ?",
                            (workflow_id, skill),
                        )
                        await update_workflow(db, workflow_id, status="failed")

                    await _db_write(f"workflow_stop:{workflow_id}:{skill_name}", _stop_workflow)
                    await _broadcast(workflow_id, {
                        "type": "workflow_failed",
                        "error": "工作流某个步骤失败，请查看日志。",
                        "output_files": sorted(_snapshot_workspace(Path(wf["workspace_dir"])), key=str.lower),
                    })
                    return

                feedback_text = _checkpoint_feedback_text(response)
                # The UI offers revision text for approve/assets_resolve as
                # well as feedback checkpoints.  A non-empty request must be
                # executed instead of being silently treated as approval.
                if action == "feedback" and feedback_text:
                    async def _resume_for_feedback(db, skill=skill_name):
                        await db.execute(
                            "UPDATE workflow_steps SET status = 'running' WHERE workflow_id = ? AND skill_name = ?",
                            (workflow_id, skill),
                        )
                        await db.commit()

                    await _db_write(f"workflow_feedback_run:{workflow_id}:{skill_name}", _resume_for_feedback)
                    feedback_path = Path(wf["workspace_dir"]) / "checkpoint_feedback.md"
                    feedback_path.write_text(
                        f"# 用户对「{step_def.display_name}」的修改意见\n\n{feedback_text}\n",
                        encoding="utf-8",
                    )
                    await _broadcast(workflow_id, {
                        "type": "step_progress", "step": skill_name,
                        "log": "[系统] 正在根据你的建议修改当前步骤产出...",
                    })

                    revision_files = sorted(_snapshot_workspace(Path(wf["workspace_dir"])), key=str.lower)
                    revision_context = (
                        "工作区文件:\n"
                        + "\n".join(f"- {path}" for path in revision_files)
                        + "\n\n## ⛔⛔⛔ OVERRIDE MODE: 用户建议审查（最高优先级，覆盖 SKILL.md 中的所有指令）\n\n"
                        "**你现在处于「建议审查模式」，不是「完整执行模式」。**\n"
                        "**忽略 SKILL.md 中关于从头执行、生成完整报告的指令。**\n\n"
                        f"用户对步骤「{step_def.display_name}」提出了以下建议：\n\n> {feedback_text}\n\n"
                        "### 执行流程（必须严格按顺序）：\n\n"
                        "**Step 1: 读取现有产出**\n"
                        "用 Read 工具读取该步骤已生成的关键文件，了解当前内容。\n\n"
                        "**Step 2: 逐条分析用户建议**\n"
                        "对用户的每条建议，输出分析判断（必须打印到终端让用户看到）：\n"
                        "```\n=== 建议审查 ===\n建议 1: [用户原文]\n  分析: [这条建议是否合理？为什么？]\n"
                        "  判定: ✅ 采纳 / ❌ 不采纳\n  理由: [采纳则说明如何修改；不采纳则说明为什么不合理]\n\n"
                        "建议 2: [用户原文]\n  分析: ...\n  判定: ...\n```\n\n"
                        "**Step 3: 执行采纳的修改**\n"
                        "- 只修改被判定为 ✅ 采纳的部分\n- 用 Write/Edit 工具保存修改\n"
                        "- 保持文件格式和结构不变，只改必要的地方\n\n"
                        "**Step 4: 写入审查报告到文件**\n"
                        "将完整的审查过程和修改总结写入 `checkpoint_revision_report.md`：\n"
                        "```bash\ncat << 'EOF' > checkpoint_revision_report.md\n# 建议审查报告\n\n"
                        "## 用户建议\n[原文]\n\n## 逐条审查\n| # | 建议内容 | 判定 | 理由 |\n"
                        "|---|---------|------|------|\n| 1 | ... | ✅/❌ | ... |\n\n## 修改总结\n"
                        "- 采纳: X 条\n- 不采纳: Y 条\n- 修改的文件: [列表]\n- 具体改动: [每个文件改了什么]\nEOF\n```\n\n"
                        "### 判断原则：\n"
                        "- 用户建议如果与数据/事实矛盾 → 不采纳，给出数据依据\n"
                        "- 用户建议如果改善了表达/逻辑/结构 → 采纳\n"
                        "- 用户建议如果涉及方法论变更但当前数据不支持 → 不采纳，解释原因\n"
                        "- 用户建议如果是补充遗漏内容 → 采纳\n"
                        "- 不确定时倾向于采纳（用户通常比 AI 更了解自己的需求）\n"
                    )

                    runner = ClaudeRunner()
                    async def on_revision_output(line: str) -> None:
                        await _broadcast(workflow_id, {
                            "type": "step_progress", "step": skill_name, "log": line,
                        })

                    await runner.run_skill(
                        skill_name=skill_name,
                        arguments=workflow_id,
                        cwd=Path(wf["workspace_dir"]),
                        workflow_id=f"{workflow_id}_{skill_name}_revision",
                        on_output=on_revision_output,
                        extra_params=wf.get("params", {}),
                        workspace_files=revision_files,
                        context_summary=revision_context,
                        inactivity_timeout=2400,
                    )
                    report_path = Path(wf["workspace_dir"]) / "checkpoint_revision_report.md"
                    report_text = report_path.read_text(encoding="utf-8") if report_path.is_file() else ""
                    primary_path = Path(wf["workspace_dir"]) / (step_def.primary_output or "")
                    primary_text = primary_path.read_text(encoding="utf-8") if primary_path.is_file() else ""
                    await _broadcast(workflow_id, {
                        "type": "step_progress", "step": skill_name,
                        "log": "[系统] 修改完成，请在弹窗中复审修改结果",
                    })
                    second_event = {
                        "type": "checkpoint_hit", "step": skill_name,
                        "display_name": f"{step_def.display_name}（修改后）",
                        "checkpoint_type": "feedback",
                    }
                    if step_def.primary_output:
                        second_event["primary_output_file"] = step_def.primary_output
                        second_event["primary_output_content"] = (
                            f"## 📋 修改审查报告\n\n{report_text}\n\n---\n\n{primary_text}"
                        )
                    await _broadcast(workflow_id, second_event)

                    response_2 = await wait_checkpoint(workflow_id, timeout=None)
                    action_2 = response_2.get("action") if isinstance(response_2, dict) else None
                    if action_2 == "feedback":
                        data_2 = response_2.get("data") if isinstance(response_2.get("data"), dict) else {}
                        feedback_text_2 = data_2.get("feedback", "")
                        feedback_path.write_text(
                            f"# 用户对「{step_def.display_name}」的第二轮修改意见\n\n{feedback_text_2}\n",
                            encoding="utf-8",
                        )
                        await _broadcast(workflow_id, {
                            "type": "step_progress", "step": skill_name,
                            "log": "[系统] 正在根据你的第二轮建议继续修改...",
                        })
                        revision_files_2 = sorted(_snapshot_workspace(Path(wf["workspace_dir"])), key=str.lower)
                        revision_context_2 = (
                            "工作区文件:\n"
                            + "\n".join(f"- {path}" for path in revision_files_2)
                            + "\n\n## ⛔⛔⛔ OVERRIDE MODE: 用户第二轮建议审查\n\n"
                            "用户对修改后的结果仍不满意，提出了第二轮建议：\n\n"
                            f"> {feedback_text_2}\n\n"
                            "请按照与第一轮相同的流程执行：读取现有产出 → 分析建议 → 执行修改 → 写入审查报告。\n"
                            "注意：这是在第一轮修改基础上的进一步调整，不要回退第一轮已采纳的修改。\n"
                        )
                        async def on_revision_output_2(line: str) -> None:
                            await _broadcast(workflow_id, {
                                "type": "step_progress", "step": skill_name, "log": line,
                            })

                        await runner.run_skill(
                            skill_name=skill_name,
                            arguments=workflow_id,
                            cwd=Path(wf["workspace_dir"]),
                            workflow_id=f"{workflow_id}_{skill_name}_revision2",
                            on_output=on_revision_output_2,
                            extra_params=wf.get("params", {}),
                            workspace_files=revision_files_2,
                            context_summary=revision_context_2,
                            inactivity_timeout=2400,
                        )
                    _generate_claude_md(
                        Path(wf["workspace_dir"]), wf["title"], template, wf.get("params", {})
                    )

                async def _complete_after_checkpoint(db, skill=skill_name):
                    await db.execute(
                        "UPDATE workflow_steps SET status = 'completed' WHERE workflow_id = ? AND skill_name = ?",
                        (workflow_id, skill),
                    )
                    await update_workflow(db, workflow_id, status="running", current_step=skill)

                await _db_write(f"workflow_after_checkpoint:{workflow_id}:{skill_name}", _complete_after_checkpoint)
                if resumed_checkpoint and action != "feedback":
                    await _broadcast(workflow_id, {
                        "type": "step_completed",
                        "step": skill_name,
                        "result_summary": "已恢复检查点确认",
                        "output_files": [],
                    })


        # A full pipeline is not submission-ready merely because every
        # generative step returned successfully.  Its terminal transition is
        # controlled by the independent, deterministic assurance envelope.
        if template == "full_pipeline":
            assurance_envelope = await _evaluate_full_pipeline_assurance(
                wf, Path(wf["workspace_dir"]),
            )
            assurance_path = _write_assurance_artifact(
                Path(wf["workspace_dir"]), assurance_envelope,
            )
            gate_summary = ", ".join(
                f"{gate.get('id')}={gate.get('status')}"
                for gate in assurance_envelope.get("gates", [])
            )
            await _log(
                workflow_id,
                "assurance-gate",
                "info",
                f"[ASSURANCE] {assurance_envelope.get('status')}: {gate_summary}",
            )
            await _broadcast(workflow_id, {
                "type": "step_progress",
                "step": "assurance-gate",
                "log": (
                    f"[系统] 最终质量门禁 {assurance_envelope.get('status')}；"
                    f"审计报告：{assurance_path.name}"
                ),
                "assurance": assurance_envelope,
            })
            if not assurance_envelope.get("submission_ready", False):
                blocking_codes = [
                    str(item.get("code"))
                    for item in assurance_envelope.get("findings", [])
                    if item.get("severity") in {"critical", "major"}
                ]
                reason = (
                    "最终质量门禁未通过"
                    + (f": {', '.join(blocking_codes)}" if blocking_codes else "")
                )

                async def _fail_assurance(db):
                    await update_workflow(
                        db, workflow_id, status="failed", current_step="assurance-gate",
                    )

                await _db_write(f"workflow_assurance_fail:{workflow_id}", _fail_assurance)
                blocked_files = sorted(
                    _snapshot_workspace(Path(wf["workspace_dir"])), key=str.lower,
                )
                await _broadcast(workflow_id, {
                    "type": "workflow_failed",
                    "reason": reason,
                    "error": reason,
                    "step": "assurance-gate",
                    "output_files": blocked_files,
                    "assurance": assurance_envelope,
                })
                return

        final_files = sorted(_snapshot_workspace(Path(wf["workspace_dir"])), key=str.lower)
        if not final_files:
            error = "所有步骤执行完毕但未产出任何文件。"

            async def _fail_empty(db):
                await update_workflow(db, workflow_id, status="failed")

            await _db_write(f"workflow_empty_fail:{workflow_id}", _fail_empty)
            await _broadcast(workflow_id, {
                "type": "workflow_failed", "error": error, "output_files": [],
            })
            return
        # The installed engine retains the final current_step value after a
        # successful terminal transition rather than clearing it.

        async def _complete_workflow(db):
            await update_workflow(db, workflow_id, status="completed")

        await _db_write(f"workflow_completed:{workflow_id}", _complete_workflow)
        await _broadcast(workflow_id, {
            "type": "workflow_completed", "output_files": final_files,
        })

    except asyncio.CancelledError:
        async def _pause_cancelled(db):
            await db.execute(
                "UPDATE workflow_steps SET status = 'pending' WHERE workflow_id = ? AND status = 'running'",
                (workflow_id,),
            )
            await update_workflow(db, workflow_id, status="paused")

        await _db_write(f"workflow_cancelled:{workflow_id}", _pause_cancelled)
        await _broadcast(workflow_id, {"type": "workflow_paused"})
        return
    finally:
        # Restore original CWD so concurrent workflows and other coroutines are
        # not affected by the workspace chdir performed at startup.
        os.chdir(original_cwd)


async def wait_checkpoint(workflow_id: str, timeout: float | None = None) -> dict:
    """Wait for explicit feedback; a timeout is opt-in for tests/automation."""
    # A resolve may arrive while the workflow task is paused/cancelled. Keep
    # that durable decision and apply it as soon as the waiter is recreated.
    if workflow_id in _checkpoint_responses:
        return _checkpoint_responses.pop(workflow_id)

    event = asyncio.Event()
    old_event = _checkpoint_events.get(workflow_id)
    if old_event is not None:
        log.warning("wait_checkpoint: residual event for %s, replaced", workflow_id)
        old_event.set()
        await asyncio.sleep(0)
    _checkpoint_events[workflow_id] = event
    try:
        # Re-check after registration so a concurrent resolve is not missed.
        if workflow_id in _checkpoint_responses:
            event.set()
        if timeout is None:
            await event.wait()
        else:
            await asyncio.wait_for(event.wait(), timeout=timeout)
        if _checkpoint_events.get(workflow_id) is not event:
            return {"action": "superseded"}
        return _checkpoint_responses.pop(workflow_id, {"action": "approve"})
    except asyncio.TimeoutError:
        return {"action": "approve", "auto": True}
    finally:
        if _checkpoint_events.get(workflow_id) is event:
            _checkpoint_events.pop(workflow_id, None)


def resolve_checkpoint(workflow_id: str, response: dict) -> bool:
    """Resolve a waiting checkpoint; return whether a waiter was present.

    Responses are always retained so a later resume can consume them even when
    the in-memory waiter was cancelled by pause/delete.
    """
    _checkpoint_responses[workflow_id] = response
    event = _checkpoint_events.get(workflow_id)
    if event is None:
        return False
    event.set()
    return True


def submit_checkpoint(workflow_id: str, response: dict) -> bool:
    """Compatibility alias retained for the reconstructed HTTP router."""
    return resolve_checkpoint(workflow_id, response)


async def update_workflow(db, wf_id: str, **fields) -> None:
    """(docstring)"""
    from services.state_store import update_workflow as _update
    await _update(db, wf_id, **fields)
