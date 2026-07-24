"""Generate and inspect real LaTeX, PDF, DOCX, and image artifacts locally."""
from __future__ import annotations

import asyncio
import hashlib
import json
import os
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
BACKEND = ROOT / "backend"
sys.path.insert(0, str(BACKEND))


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def find_tool(name: str, bundled: Path) -> str:
    if bundled.exists():
        return str(bundled)
    value = shutil.which(name)
    if value:
        return value
    raise FileNotFoundError(name)


def run(cmd: list[str], cwd: Path) -> dict:
    result = subprocess.run(
        cmd,
        cwd=str(cwd),
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=120,
        creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
    )
    return {"cmd": cmd, "returncode": result.returncode, "stdout": result.stdout[-2000:], "stderr": result.stderr[-2000:]}


async def main(output_root: Path) -> dict:
    from config import RUNTIME_DRAWIO, RUNTIME_TEXLIVE, TOOLS_DIR
    from services.docx_tool_loader import get_markdown_to_docx

    output_root.mkdir(parents=True, exist_ok=True)
    paper = output_root / "paper"
    figures = output_root / "figures"
    paper.mkdir(exist_ok=True)
    figures.mkdir(exist_ok=True)

    drawio_xml = """<mxfile host="app.diagrams.net"><diagram name="Page-1"><mxGraphModel dx="800" dy="600" grid="1" page="1" pageWidth="827" pageHeight="1169"><root><mxCell id="0"/><mxCell id="1" parent="0"/><mxCell id="2" value="Input" style="rounded=1;whiteSpace=wrap;html=1;fillColor=#dae8fc;strokeColor=#6c8ebf;" vertex="1" parent="1"><mxGeometry x="80" y="100" width="120" height="60" as="geometry"/></mxCell><mxCell id="3" value="Validated Output" style="rounded=1;whiteSpace=wrap;html=1;fillColor=#d5e8d4;strokeColor=#82b366;" vertex="1" parent="1"><mxGeometry x="300" y="100" width="160" height="60" as="geometry"/></mxCell><mxCell id="4" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;endArrow=block;" edge="1" parent="1" source="2" target="3"><mxGeometry relative="1" as="geometry"/></mxCell></root></mxGraphModel></diagram></mxfile>"""
    drawio_source = figures / "pipeline.drawio"
    drawio_source.write_text(drawio_xml, encoding="utf-8")

    runtime_drawio = Path(RUNTIME_DRAWIO) if RUNTIME_DRAWIO else ROOT / "runtime" / "draw.io"
    if not runtime_drawio.exists():
        runtime_drawio = ROOT / "runtime" / "draw.io"
    drawio = find_tool("drawio", runtime_drawio / "draw.io.exe")
    image_png = figures / "pipeline.png"
    image_pdf = figures / "pipeline.pdf"
    drawio_png_log = run([drawio, "--export", "--format", "png", "--scale", "2", "--output", str(image_png), str(drawio_source)], output_root)
    drawio_pdf_log = run([drawio, "--export", "--format", "pdf", "--output", str(image_pdf), str(drawio_source)], output_root)

    tex = r"""\documentclass[11pt]{article}
\usepackage[margin=1in]{geometry}
\usepackage{graphicx}
\usepackage{hyperref}
\title{Vibe Research Document Pipeline Verification}
\author{Local Artifact Test}
\date{July 2026}
\begin{document}
\maketitle
\begin{abstract}
This document validates a real local artifact chain: LaTeX source, a compiled PDF, a Word document, and a generated diagram image.
\end{abstract}
\section{Pipeline}
The application invokes bundled or local document tools and preserves concrete outputs for inspection.
\begin{figure}[ht]
\centering
\includegraphics[width=0.75\linewidth]{../figures/pipeline.png}
\caption{Locally generated document pipeline diagram.}
\end{figure}
\section{Quality Gate}
The PDF must open, contain at least one page, embed the diagram, and contain extractable text. The DOCX must contain document XML, media, relationships, and non-empty paragraphs.
\end{document}
"""
    tex_path = paper / "main.tex"
    tex_path.write_text(tex, encoding="utf-8")

    runtime_tex = Path(RUNTIME_TEXLIVE) if RUNTIME_TEXLIVE else ROOT / "runtime" / "texlive"
    if not runtime_tex.exists():
        runtime_tex = ROOT / "runtime" / "texlive"
    xelatex = find_tool("xelatex", runtime_tex / "miktex" / "bin" / "x64" / "xelatex.exe")
    compile_runs = [
        run([xelatex, "-interaction=nonstopmode", "-halt-on-error", "main.tex"], paper),
        run([xelatex, "-interaction=nonstopmode", "-halt-on-error", "main.tex"], paper),
    ]
    pdf_path = paper / "main.pdf"

    markdown = """# Vibe Research 文档流水线实产物验证

## 摘要

本文档验证 Vibe Research 的本地文档流水线。验证对象包括 LaTeX 源文件、PDF、DOCX 与真实图片，并检查文件结构、可读文本、媒体嵌入和调用链日志。

## 调用链

1. draw.io CLI 从 `pipeline.drawio` 生成 PNG/PDF 图。
2. XeLaTeX 读取 `paper/main.tex` 并嵌入 PNG，生成 `paper/main.pdf`。
3. 产品 DOCX 引擎（Node `docx-cn-engine` 或 Python `tools/docx_export.py`）将本文 Markdown 转换为 Word。

![本地文档流水线](../figures/pipeline.png)

## 质量检查

| 产物 | 验证项 |
|---|---|
| PDF | 页数、文本、图片对象、文件大小 |
| DOCX | ZIP 结构、段落、媒体、关系文件 |
| image | PNG 魔数、尺寸、非空像素内容 |

结论：所有产物均在本地生成，不依赖外部网络服务。
"""
    md_path = paper / "main.md"
    md_path.write_text(markdown, encoding="utf-8")
    docx_path = paper / "main.docx"

    converter_label = None
    node_bin = shutil.which("node")
    node_script = ROOT / "tools" / "docx-cn-engine" / "md_to_docx.js"
    if node_bin and node_script.is_file():
        converter_label = f"node:{node_script}"
        node_log = run(
            [
                node_bin,
                str(node_script),
                "--source",
                str(md_path),
                "--output",
                str(docx_path),
                "--workspace",
                str(output_root),
            ],
            output_root,
        )
        if node_log["returncode"] != 0 or not docx_path.is_file():
            converter = get_markdown_to_docx()
            if converter is None:
                raise FileNotFoundError(f"Node DOCX failed and Python converter missing: {node_log}")
            converter_label = "python:tools/docx_export.py"
            await asyncio.to_thread(converter, md_path, docx_path, None, output_root, "python")
    else:
        converter = get_markdown_to_docx()
        if converter is None:
            raise FileNotFoundError("tools/docx_export.py")
        converter_label = "python:tools/docx_export.py"
        await asyncio.to_thread(converter, md_path, docx_path, None, output_root, "python")

    checks: dict = {}
    checks["image"] = {"exists": image_png.exists(), "size": image_png.stat().st_size if image_png.exists() else 0}
    try:
        from PIL import Image
        with Image.open(image_png) as image:
            checks["image"].update({"format": image.format, "width": image.width, "height": image.height})
    except Exception as exc:
        checks["image"]["inspect_error"] = str(exc)

    checks["pdf"] = {"exists": pdf_path.exists(), "size": pdf_path.stat().st_size if pdf_path.exists() else 0}
    if pdf_path.exists():
        try:
            import fitz
            document = fitz.open(pdf_path)
            checks["pdf"].update({
                "pages": document.page_count,
                "text_chars": sum(len(page.get_text()) for page in document),
                "image_objects": sum(len(page.get_images(full=True)) for page in document),
            })
            document.close()
        except Exception as exc:
            checks["pdf"]["inspect_error"] = str(exc)

    checks["docx"] = {"exists": docx_path.exists(), "size": docx_path.stat().st_size if docx_path.exists() else 0}
    if docx_path.exists():
        with zipfile.ZipFile(docx_path) as archive:
            names = archive.namelist()
            checks["docx"].update({
                "zip_valid": True,
                "has_document_xml": "word/document.xml" in names,
                "media_files": [name for name in names if name.startswith("word/media/")],
                "has_relationships": "word/_rels/document.xml.rels" in names,
            })
        try:
            from docx import Document
            document = Document(docx_path)
            checks["docx"].update({
                "paragraphs": len(document.paragraphs),
                "nonempty_paragraphs": sum(bool(p.text.strip()) for p in document.paragraphs),
                "inline_shapes": len(document.inline_shapes),
            })
        except Exception as exc:
            checks["docx"]["inspect_error"] = str(exc)

    assert checks["image"].get("format") == "PNG" and checks["image"].get("width", 0) > 100
    assert checks["pdf"].get("pages", 0) >= 1 and checks["pdf"].get("text_chars", 0) > 100
    assert checks["pdf"].get("image_objects", 0) >= 1
    assert checks["docx"].get("zip_valid") and checks["docx"].get("has_document_xml")
    assert checks["docx"].get("nonempty_paragraphs", 0) >= 5
    assert checks["docx"].get("inline_shapes", 0) >= 1

    artifacts = {}
    for path in [drawio_source, image_png, image_pdf, tex_path, pdf_path, md_path, docx_path]:
        artifacts[str(path.relative_to(output_root)).replace("\\", "/")] = {
            "size": path.stat().st_size,
            "sha256": sha256(path),
        }
    return {
        "output_root": str(output_root),
        "tools": {"drawio": drawio, "xelatex": xelatex, "docx_converter": converter_label},
        "calls": {"drawio_png": drawio_png_log, "drawio_pdf": drawio_pdf_log, "xelatex": compile_runs},
        "checks": checks,
        "artifacts": artifacts,
    }


if __name__ == "__main__":
    target = Path(sys.argv[1]) if len(sys.argv) > 1 else ROOT / "tests" / "document_artifacts"
    report = asyncio.run(main(target.resolve()))
    print(json.dumps(report, ensure_ascii=False, indent=2))
