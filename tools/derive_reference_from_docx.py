#!/usr/bin/env python3
"""从一份完整的 .docx 论文派生通用参考资料（不是模板，不是 profile）。

输入：一份完整的论文 .docx（如外部提供的优质参考论文）
输出：
  1. reference_paper_general.json — 结构化样式参考（字号/字体/对齐/颜色/边距/行距/首行缩进）
  2. reference_structure_general.md — 人类可读 + LLM 友好的「格式自检对照表」

派生时做"通用化"：
  - 字号、字体、颜色、对齐、边距、行距、首行缩进 → 保留（通用规范）
  - 章节大纲 → 抽取通用骨架（去掉具体题目/问题/内容），归一为 "标题/摘要/关键词/正文章节/参考文献/附录"
  - 题注格式 → 抽取范式（"图 X：xxx" / "表 X：xxx" 等）
  - 公式编号格式 → 抽取范式（(N) / \tag{} / X.Y 等）
  - Bold 使用模式 → 抽取通用规则（关键词标签加粗、章节标题加粗、首行术语加粗）

用法：
    python derive_reference_from_docx.py <input.docx> [--out-dir <dir>]
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter, defaultdict
from pathlib import Path
from typing import Any, Optional

try:
    from docx import Document
    from docx.shared import Cm, Pt
except ImportError:
    print("ERROR: 需要 python-docx，请先 pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ---------------------------------------------------------------------------
# 工具
# ---------------------------------------------------------------------------

_HEADING_LEVEL_RE = re.compile(r"^Heading\s*(\d+)$", re.IGNORECASE)


def _emu_to_cm(emu: int) -> float:
    """English Metric Unit → 厘米（python-docx 内部用 EMU 存边距）"""
    return round(emu / 360000, 2)


def _twips_to_pt(twips: int) -> float:
    return round(twips / 20, 1)


def _detect_heading_level(style_name: str) -> Optional[int]:
    if not style_name:
        return None
    m = _HEADING_LEVEL_RE.match(style_name.strip())
    if m:
        return int(m.group(1))
    return None


def _is_caption_text(text: str) -> bool:
    """启发式：是否像图/表注释（不依赖 style.caption）"""
    if not text or len(text) > 100:
        return False
    s = text.strip()
    # "图 1：xxx" / "图 1-2：xxx" / "Figure 1:..." / "表 1：xxx" / "Table 1:..."
    if re.match(
        r"^(?:图|表|Figure|Fig\.?|Table|Tab\.?)\s*[\dA-Za-z][\dA-Za-z\-\.]*\s*[：:、\.\-]",
        s,
    ):
        return True
    # 无编号的"图 xxx" / "表 xxx：xxx"
    if re.match(r"^(?:图|表)\s*[：:]\s*\S", s):
        return True
    return False


def _is_formula_number_pattern(text: str) -> Optional[str]:
    """识别公式编号风格。返回模式名或 None"""
    if not text:
        return None
    s = text.strip()
    # "$$ ... $$ (1)" — 同行尾追加
    if re.search(r"\)\s*$", s) and re.search(r"\$\$.+\$\$", s):
        return "trailing_paren"
    # \tag{1}
    if r"\tag{" in s:
        return "inline_tag"
    return None


# ---------------------------------------------------------------------------
# 派生
# ---------------------------------------------------------------------------

def derive_reference(docx_path: Path) -> dict[str, Any]:
    """从 docx 派生结构化参考资料"""
    doc = Document(str(docx_path))

    # ========== 1. 页面布局 ==========
    sec = doc.sections[0]
    page = {
        "size": "A4",  # python-docx 不直接给纸张名，按常识标
        "width_cm": _emu_to_cm(sec.page_width) if sec.page_width else None,
        "height_cm": _emu_to_cm(sec.page_height) if sec.page_height else None,
        "margin_top_cm": _emu_to_cm(sec.top_margin) if sec.top_margin else 2.5,
        "margin_bottom_cm": _emu_to_cm(sec.bottom_margin) if sec.bottom_margin else 2.5,
        "margin_left_cm": _emu_to_cm(sec.left_margin) if sec.left_margin else 2.5,
        "margin_right_cm": _emu_to_cm(sec.right_margin) if sec.right_margin else 2.5,
    }

    # ========== 2. 样式属性（Heading / Normal） ==========
    styles_info: dict[str, dict] = {}
    for st in doc.styles:
        try:
            sn = st.name
        except Exception:
            continue
        # 只关心 Heading 1-6 / Title / Normal / Caption
        if not (
            (sn and _detect_heading_level(sn))
            or sn in ("Title", "Normal", "Caption", "Subtitle")
        ):
            continue
        try:
            f = st.font
            info = {
                "font_name": f.name,
                "size_pt": f.size.pt if f.size else None,
                "bold": f.bold,
                "italic": f.italic,
                "color_rgb": str(f.color.rgb) if (f.color and f.color.rgb) else None,
            }
            # 段落属性
            try:
                pf = st.paragraph_format
                info["alignment"] = str(pf.alignment).split(".")[-1].lower() if pf.alignment is not None else None
                info["line_spacing"] = pf.line_spacing
                info["first_line_indent_pt"] = pf.first_line_indent.pt if pf.first_line_indent else None
                info["space_before_pt"] = pf.space_before.pt if pf.space_before else None
                info["space_after_pt"] = pf.space_after.pt if pf.space_after else None
            except Exception:
                pass
            styles_info[sn] = info
        except Exception:
            pass

    # ========== 3. 实测段落级 run 属性（fallback：很多模板的样式只是占位，真实属性在 run 上） ==========
    # 思路：扫描所有段落，按 style 分组，统计 run 的实际字号/字体/颜色/加粗的"众数"
    run_stats: dict[str, dict] = defaultdict(lambda: {
        "size_pt": Counter(),
        "font_name": Counter(),
        "color": Counter(),
        "bold_ratio": [],  # 该 style 下 bold run 的占比
    })
    body_first_line_indent = []
    body_line_spacing = []

    for p in doc.paragraphs:
        sn = p.style.name if p.style else "NoStyle"
        bold_count = 0
        run_count = 0
        for r in p.runs:
            run_count += 1
            if r.font.size:
                run_stats[sn]["size_pt"][r.font.size.pt] += 1
            if r.font.name:
                run_stats[sn]["font_name"][r.font.name] += 1
            if r.font.color and r.font.color.rgb:
                run_stats[sn]["color"][str(r.font.color.rgb)] += 1
            if r.bold:
                bold_count += 1
        if run_count > 0:
            run_stats[sn]["bold_ratio"].append(bold_count / run_count)
        # 收集 body 段落的首行缩进与行距
        if sn in ("Normal", "NoStyle") and p.text.strip():
            try:
                if p.paragraph_format.first_line_indent:
                    body_first_line_indent.append(p.paragraph_format.first_line_indent.pt)
                if p.paragraph_format.line_spacing:
                    body_line_spacing.append(p.paragraph_format.line_spacing)
            except Exception:
                pass

    measured: dict[str, dict] = {}
    for sn, st in run_stats.items():
        info = {}
        if st["size_pt"]:
            info["size_pt_mode"] = st["size_pt"].most_common(1)[0][0]
        if st["font_name"]:
            info["font_name_mode"] = st["font_name"].most_common(1)[0][0]
        if st["color"]:
            info["color_mode"] = st["color"].most_common(1)[0][0]
        if st["bold_ratio"]:
            info["bold_avg"] = round(sum(st["bold_ratio"]) / len(st["bold_ratio"]), 2)
        measured[sn] = info

    # ========== 4. 章节大纲（去内容、保骨架） ==========
    outline = []
    for p in doc.paragraphs:
        lvl = _detect_heading_level(p.style.name if p.style else "")
        if lvl is None:
            continue
        text = p.text.strip()
        if not text:
            continue
        outline.append({"level": lvl, "text": text[:100]})

    # 通用化大纲：把具体章节抽象成结构槽位
    generic_outline = _genericize_outline(outline)

    # ========== 5. 题注 / 公式编号范式 ==========
    captions_raw = []
    formula_styles = Counter()
    # 同时扫描 main.md 之类的内容时检测块公式中的编号风格
    full_text_for_formula = "\n".join(p.text for p in doc.paragraphs)
    # 检测 "$$ ... $$ (1)" / "\tag{}" 在整个文档里
    if re.search(r"\\tag\{", full_text_for_formula):
        formula_styles["inline_tag"] = len(re.findall(r"\\tag\{", full_text_for_formula))
    if re.search(r"\$\$.+\$\$\s*\(\d+\)", full_text_for_formula):
        formula_styles["trailing_paren"] = len(
            re.findall(r"\$\$.+?\$\$\s*\(\d+\)", full_text_for_formula)
        )
    # 同时检测 "公式 X" / "式 X" / "(X.Y)" 等编号风格
    if re.search(r"^\s*\(\d+(?:\.\d+)?\)\s*$", full_text_for_formula, re.MULTILINE):
        formula_styles["standalone_paren"] = len(
            re.findall(r"^\s*\(\d+(?:\.\d+)?\)\s*$", full_text_for_formula, re.MULTILINE)
        )

    for p in doc.paragraphs:
        text = p.text.strip()
        if _is_caption_text(text):
            if len(captions_raw) < 10:
                captions_raw.append(text[:100])
    # 通用化：脱敏题注，保留范式
    captions = [_generalize_caption(c) for c in captions_raw]
    # 去重
    captions = list(dict.fromkeys(captions))

    # ========== 6. 加粗使用模式 ==========
    bold_patterns = []
    for p in doc.paragraphs:
        for r in p.runs:
            if r.bold and r.text.strip() and len(r.text) < 30:
                t = r.text.strip()
                # 只收集有结构特征的（关键词标签 / 假设标签 / 章节序号），并通用化文本
                m_label = re.match(
                    r"^(关键词|Key\s*words|Keywords|假设|引理|定理|命题|证明|算法|步骤|备注|注)([一二三四五六七八九十\d]*)\s*[：:]?\s*$",
                    t,
                )
                if m_label:
                    label = m_label.group(1)
                    suffix = m_label.group(2)
                    bold_patterns.append(f"{label}{suffix}：（结构化标签）")
                    continue
                m_chap = re.match(r"^第([一二三四五六七八九十百千零〇\d]+)([章节部篇])\s*$", t)
                if m_chap:
                    bold_patterns.append(f"第{m_chap.group(1)}{m_chap.group(2)}（章节序号）")
                    continue
    bold_patterns = list(dict.fromkeys(bold_patterns))[:10]

    # ========== 7. 表格使用 ==========
    table_count = len(doc.tables)
    # 看第一张表的结构
    sample_table_info = None
    if doc.tables:
        t0 = doc.tables[0]
        sample_table_info = {
            "row_count": len(t0.rows),
            "col_count": len(t0.columns),
            "has_header_style": any("header" in (r.style.name.lower() if r.style else "") for r in t0.rows[:1] if hasattr(r, "style")),
        }

    return {
        "page": page,
        "styles": styles_info,
        "measured_runs": measured,
        "outline_raw": outline,
        "outline_generic": generic_outline,
        "body_indent_pt_avg": round(sum(body_first_line_indent) / len(body_first_line_indent), 1) if body_first_line_indent else None,
        "body_line_spacing_avg": round(sum(body_line_spacing) / len(body_line_spacing), 2) if body_line_spacing else None,
        "captions_samples": captions,
        "formula_number_styles": dict(formula_styles),
        "bold_patterns_samples": bold_patterns,
        "table_count": table_count,
        "sample_table": sample_table_info,
        "source_file": docx_path.name,
    }


def _genericize_outline(outline: list[dict]) -> list[dict]:
    """章节大纲通用化：按 Heading 层级映射到通用槽位

    规则：
      - 第一个 H1 视为「论文题目」（slot=title），不参与槽位匹配
      - 其它 Heading 按文本匹配通用槽位；匹配不到则 slot=section
    """
    SLOT_RULES = [
        (r"摘\s*要|Abstract", "abstract"),
        (r"关键词|Keywords", "keywords"),
        (r"^问题重述|^引言|^Introduction|^绪论|^一、|^1\s+(?:引言|背景|绪论)", "intro"),
        (r"^问题分析|^背景分析|^二、", "analysis"),
        (r"^.{0,4}假设|^.{0,4}符号说明|^术语|^Notation|^三、", "assumptions"),
        (r"^模型(?:建立|构建)|^方法|^Method|^理论|^四、", "methodology"),
        (r"^(?:模型)?求解|^计算结果|^实验结果|^数值结果|^Results|^五、", "experiments"),
        (r"^结论|^总结|^Conclusion", "conclusion"),
        (r"^参考文献|^References|^Bibliography", "references"),
        (r"^附录|^Appendix", "appendix"),
        (r"^致谢|^Acknowledg", "acknowledgement"),
    ]
    generic = []
    title_seen = False
    for item in outline:
        text = item["text"]
        # 第一个 H1 一律是论文题目
        if item["level"] == 1 and not title_seen:
            slot = "title"
            title_seen = True
            generic.append({"level": item["level"], "slot": slot, "original": text[:50]})
            continue
        # H2 级别才尝试匹配槽位（深层级一律 section）
        if item["level"] == 2:
            slot = "section"
            for pat, name in SLOT_RULES:
                if re.search(pat, text):
                    slot = name
                    break
        else:
            slot = "section"
        generic.append({"level": item["level"], "slot": slot, "original": text[:50]})
    return generic


def _generalize_caption(text: str) -> str:
    """题注脱敏：保留范式（'图 X：' 前缀），把具体内容替换成 ... """
    s = text.strip()
    # 保留前缀 + 编号 + 分隔符
    m = re.match(
        r"^((?:图|表|Figure|Fig\.?|Table|Tab\.?)\s*[\dA-Za-z][\dA-Za-z\-\.]*\s*[：:、\.\-])\s*(.+)$",
        s,
    )
    if m:
        return m.group(1) + " ……（具体说明文字）"
    m2 = re.match(r"^((?:图|表)\s*[：:])\s*(.+)$", s)
    if m2:
        return m2.group(1) + " ……（具体说明文字）"
    return s


# ---------------------------------------------------------------------------
# 输出生成
# ---------------------------------------------------------------------------

def build_general_profile(derived: dict) -> dict:
    """从派生数据中抽出通用 profile（JSON），剥离竞赛/学科特定的内容"""
    page = derived["page"]
    styles = derived.get("styles", {})
    measured = derived.get("measured_runs", {})

    def _style_size(name: str, fallback: float) -> float:
        # 优先实测的众数，其次 styles
        if name in measured and "size_pt_mode" in measured[name]:
            return measured[name]["size_pt_mode"]
        if name in styles and styles[name].get("size_pt"):
            return styles[name]["size_pt"]
        return fallback

    def _style_font(name: str, fallback: str) -> str:
        if name in measured and "font_name_mode" in measured[name]:
            return measured[name]["font_name_mode"]
        if name in styles and styles[name].get("font_name"):
            return styles[name]["font_name"]
        return fallback

    def _style_color(name: str, fallback: str = "000000") -> str:
        if name in measured and "color_mode" in measured[name]:
            return measured[name]["color_mode"]
        if name in styles and styles[name].get("color_rgb"):
            return styles[name]["color_rgb"]
        return fallback

    return {
        "profile_name": "general_paper_reference",
        "source": derived.get("source_file"),
        "page": {
            "size": page.get("size", "A4"),
            "margin_top_cm": page.get("margin_top_cm", 2.5),
            "margin_bottom_cm": page.get("margin_bottom_cm", 2.5),
            "margin_left_cm": page.get("margin_left_cm", 2.5),
            "margin_right_cm": page.get("margin_right_cm", 2.5),
        },
        "fonts": {
            "chinese_heading": "SimHei",
            "chinese_body": "SimSun",
            "latin": _style_font("Normal", "Times New Roman"),
            "monospace": "Consolas",
        },
        "title": {
            "font_size_pt": _style_size("Heading 1", 18),
            "bold": True,
            "alignment": "center",
            "color_rgb": _style_color("Heading 1"),
        },
        "headings": {
            "level1_pt": _style_size("Heading 2", 15),
            "level2_pt": _style_size("Heading 3", 14),
            "level3_pt": _style_size("Heading 4", 12),
            "level4_pt": _style_size("Heading 5", 11),
            "bold": True,
            "level1_alignment": "center",
            "level2_alignment": "left",
            "level3_alignment": "left",
            "level4_alignment": "left",
            "level1_color_rgb": _style_color("Heading 2"),
            "level2_color_rgb": _style_color("Heading 3"),
            "level3_color_rgb": _style_color("Heading 4"),
        },
        "body": {
            "font_size_pt": _style_size("Normal", 12),
            "line_spacing": derived.get("body_line_spacing_avg") or 1.5,
            "first_line_indent_chars": 2,
            "space_before_pt": 0,
            "space_after_pt": 0,
            "color_rgb": _style_color("Normal"),
        },
        "abstract": {
            "label_size_pt": _style_size("Heading 2", 14),
            "body_size_pt": _style_size("Normal", 12),
            "label_bold": True,
            "label_alignment": "center",
        },
        "keywords": {
            "label_size_pt": _style_size("Normal", 12),
            "label_bold": True,
            "font": "SimHei",
        },
        "table": {
            "top_border_pt": 1.5,
            "header_border_pt": 0.75,
            "bottom_border_pt": 1.5,
            "font_size_pt": 10.5,
            "header_bold": True,
            "cell_alignment": "center",
        },
        "image": {
            "max_width_cm": 14,
            "alignment": "center",
        },
        "code_block": {
            "font_size_pt": 9,
            "line_spacing": 1.0,
            "background_color": "F5F5F5",
        },
        "references": {
            "hanging_indent_cm": 0.74,
            "font_size_pt": 10.5,
            "numbering_style": "bracket",
        },
        "_derivation_note": (
            "此 profile 派生自一份完整论文样本，仅作为 LLM 自检对照参考使用。"
            "不参与渲染、不当模板。具体的字号字体由参考论文实测得出，"
            "章节大纲与具体内容已被剥离，保留通用骨架。"
        ),
    }


def build_general_reference_md(derived: dict) -> str:
    """生成给 LLM 看的"格式自检对照表"，markdown 格式"""
    profile = build_general_profile(derived)
    page = profile["page"]
    h = profile["headings"]
    b = profile["body"]
    tbl = profile["table"]

    # 章节大纲通用化
    generic_outline = derived.get("outline_generic", [])
    # 按通用 slot 去重，保留每种 slot 的第一个；section 类合并为一行示意
    seen_slots = set()
    cleaned_outline = []
    section_count_at_lvl = Counter()
    for item in generic_outline:
        slot = item["slot"]
        if slot == "title":
            cleaned_outline.append(item)
            seen_slots.add(slot)
        elif slot == "section":
            # section 不去重但限制每级最多展示 2 条作为示意
            section_count_at_lvl[item["level"]] += 1
            if section_count_at_lvl[item["level"]] <= 2:
                cleaned_outline.append(item)
        elif slot not in seen_slots:
            seen_slots.add(slot)
            cleaned_outline.append(item)

    # 章节骨架文字（占位语义化）
    SLOT_LABEL = {
        "title": "论文题目",
        "abstract": "摘要",
        "keywords": "关键词",
        "intro": "引言/问题重述",
        "analysis": "问题分析",
        "assumptions": "模型假设与符号说明",
        "methodology": "模型建立",
        "experiments": "求解/实验/结果",
        "conclusion": "结论",
        "references": "参考文献",
        "appendix": "附录",
        "acknowledgement": "致谢",
        "section": "（其它章节）",
    }
    outline_lines = []
    for item in cleaned_outline:
        depth = item["level"] - 1
        prefix = "  " * depth + "#" * item["level"] + " "
        label = SLOT_LABEL.get(item["slot"], "章节")
        outline_lines.append(f"{prefix}{label}（H{item['level']}）")

    captions = derived.get("captions_samples", [])
    formula_styles = derived.get("formula_number_styles", {})
    formula_desc = (
        "、".join(f"`{k}` ({v} 处)" for k, v in formula_styles.items())
        if formula_styles else "未识别（OMML 公式，文本中不含编号字符串）"
    )
    bold_patterns = derived.get("bold_patterns_samples", [])

    md = f"""# 论文 Markdown 格式自检对照表（通用参考）

> 派生自一份完整论文样本：`{profile.get('source')}`
> 此文件仅作为 docx-format-check 步骤的**自检对照参考**，不参与渲染、不当模板。
> 具体格式要求由 docx-export 渲染引擎按对应 profile 应用。

## 一、页面与正文规范

- 纸张：{page['size']}
- 页边距：上 {page['margin_top_cm']} cm / 下 {page['margin_bottom_cm']} cm / 左 {page['margin_left_cm']} cm / 右 {page['margin_right_cm']} cm
- 正文字号：{b['font_size_pt']} pt
- 正文行距：{b['line_spacing']} 倍
- 首行缩进：{b['first_line_indent_chars']} 个汉字字符宽度
- 中文正文字体：{profile['fonts']['chinese_body']}
- 中文标题字体：{profile['fonts']['chinese_heading']}
- 西文字体：{profile['fonts']['latin']}

## 二、标题层级与字号

| 层级 | Markdown 写法 | 字号(pt) | 加粗 | 对齐 | 备注 |
|------|--------------|---------|------|------|------|
| 标题（论文题目） | `# 论文题目` | {profile['title']['font_size_pt']} | 是 | center | 全文唯一 |
| 一级章节 | `## 一、章节名` 或 `## 1 章节名` | {h['level1_pt']} | 是 | {h['level1_alignment']} | 摘要/章节用 |
| 二级章节 | `### 1.1 小节名` | {h['level2_pt']} | 是 | {h['level2_alignment']} | |
| 三级章节 | `#### 1.1.1 子节名` | {h['level3_pt']} | 是 | {h['level3_alignment']} | |
| 四级章节 | `##### xxx` | {h['level4_pt']} | 是 | left | 慎用，超过 3 级建议改用列表 |

⛔ **铁律：标题必须用 `#` 开头**，禁止用 `- 章节名`（list 项）伪装标题，否则 docx 引擎会渲染成"• 章节名"。

## 三、章节骨架（仅供结构对照，不要照抄具体内容）

参考论文的章节骨架如下（已剥离具体内容，仅保留通用结构槽位）：

```
{chr(10).join(outline_lines[:25])}
```

⛔ **不要照搬"问题一/问题二"等竞赛特定章节**。本骨架只是说明：
- 完整论文应包含「摘要 → 关键词 → 引言/问题分析 → 模型假设 → 模型建立 → 结果 → 结论 → 参考文献」这条主线
- H2 是一级章节；H3 是二级小节；H4 是三级子节

## 四、图/表注释规范

参考论文中的题注样例：

```
{chr(10).join(captions[:6]) if captions else '（未识别到题注）'}
```

⛔ **题注铁律**：
- 图：`图 X：说明文字` 或 `图 X-Y：说明文字`（X 是章节号或全文流水号）
- 表：`表 X：说明文字`
- **题注独占一行**，紧跟在图/表的下面（图在题注上方，表在题注下方）
- ⛔ 不要写成 `- 图 X：xxx`（list 项 → 渲染出 "•"）
- ⛔ 不要写成 `**图 X**：xxx`（加粗 + 冒号会被识别成正文加粗段）

## 五、公式编号规范

参考论文使用的公式编号风格：{formula_desc}

⛔ **公式铁律**：
- 块公式必须独占行：`$$ y = ax + b $$`
- 编号必须**与公式同行**，不能写在下一行：
  - 对：`$$ y = ax + b $$ (1)` 或 `$$ y = ax + b \\tag{{1}} $$`
  - 错：`$$ y = ax + b $$\\n- (1)` （编号被渲染成 list bullet）
- 全文按出现顺序连续编号 `(1)(2)(3)...`，不跳号、不重复
- 行内公式用单 `$..$`：`$E=mc^2$`，**不用** `$$..$$`

## 六、加粗使用模式

参考论文中识别到的加粗模式（结构化标签）：

```
{chr(10).join(bold_patterns) if bold_patterns else '（未识别到结构化加粗模式）'}
```

⛔ **加粗使用规则**：
- 关键词标签 `**关键词**：` 加粗
- 假设/定理标签 `**假设一**：` 加粗
- 章节标题已自带加粗，不要再用 `**`
- ⛔ 不要整段加粗（一段全是 `**...**`）
- ⛔ 不要用加粗当强调，过多加粗反而稀释重点

## 七、表格规范

- 全文表格数：参考论文 {derived.get('table_count', 0)} 张
- 三线表样式：顶线 {tbl['top_border_pt']} pt + 表头线 {tbl['header_border_pt']} pt + 底线 {tbl['bottom_border_pt']} pt
- 表头加粗：{tbl['header_bold']}
- 单元格对齐：{tbl['cell_alignment']}
- 字号：{tbl['font_size_pt']} pt

⛔ **Markdown 表格铁律**：
- 用纯 Markdown 表格语法 `| ... | ... |`，**禁止** LaTeX `\\begin{{tabular}}`
- 第二行必须是分隔行 `|---|---|...|`
- 每行列数一致
- 数据来源优先 `figures/TABLE_*.md`（已生成的真实数据），不要凭记忆手抄

## 八、图片规范

- 引用语法：`![图 X：图题](figures/fig_xxx.png)`
- ⛔ 必须 `.png`，不能 `.pdf`（Word 不能嵌 PDF）
- alt 文本同时作为图题（自动渲染）
- 最大宽度：{profile['image']['max_width_cm']} cm
- 对齐：{profile['image']['alignment']}

## 九、参考文献规范

- 上标引用：`[1]` `[2,3]` `[4-6]`
- 不要用 `\\cite{{}}` `\\ref{{}}`
- 文献格式：`[1] 作者. 题名[J]. 期刊, 年, 卷(期): 页码.`
- 行间距：悬挂缩进 {profile['references']['hanging_indent_cm']} cm
- 字号：{profile['references']['font_size_pt']} pt

---

## 自检清单

按本对照表检查源 markdown 时，重点关注：

1. ☐ 所有章节标题用 `#` 开头，没有 `- xxx` 伪标题
2. ☐ 题注独占行，没有 `- ` 前缀
3. ☐ 块公式编号与公式同行，没有独立的 `- (1)` 行
4. ☐ Markdown 表格分隔行存在，每行列数一致
5. ☐ 图片用 `.png`，不引用 `.pdf`
6. ☐ 章节层级合理（H1 唯一，H2/H3/H4 按通用骨架排布）
7. ☐ 加粗只用在结构化标签（关键词/假设/定理），不滥用
8. ☐ 没有 LaTeX 残留（`\\begin{{}}` `\\section{{}}` `\\cite{{}}`）
"""
    return md


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    p = argparse.ArgumentParser(description="从一份完整 docx 论文派生通用参考资料")
    p.add_argument("input", type=Path, help="输入 .docx 文件")
    p.add_argument("--out-dir", type=Path,
                   default=Path(__file__).resolve().parent / "docx_style_profiles",
                   help="输出目录（默认 tools/docx_style_profiles/）")
    p.add_argument("--profile-name", default="reference_paper_general.json",
                   help="JSON 文件名")
    p.add_argument("--md-name", default="reference_structure_general.md",
                   help="Markdown 文件名")
    args = p.parse_args()

    if not args.input.exists():
        print(f"ERROR: 输入文件不存在: {args.input}", file=sys.stderr)
        sys.exit(1)

    args.out_dir.mkdir(parents=True, exist_ok=True)

    print(f"派生中: {args.input.name}")
    derived = derive_reference(args.input)

    # 输出 JSON profile
    profile = build_general_profile(derived)
    profile_path = args.out_dir / args.profile_name
    profile_path.write_text(json.dumps(profile, ensure_ascii=False, indent=2),
                            encoding="utf-8")
    print(f"  → {profile_path.relative_to(Path.cwd())} ({profile_path.stat().st_size} bytes)")

    # 输出 markdown 自检对照
    md = build_general_reference_md(derived)
    md_path = args.out_dir / args.md_name
    md_path.write_text(md, encoding="utf-8")
    print(f"  → {md_path.relative_to(Path.cwd())} ({md_path.stat().st_size} bytes)")

    # 简要统计
    print()
    print("派生概要：")
    print(f"  - 页边距: 上{derived['page']['margin_top_cm']} 下{derived['page']['margin_bottom_cm']} "
          f"左{derived['page']['margin_left_cm']} 右{derived['page']['margin_right_cm']} cm")
    print(f"  - 正文字号(实测): {profile['body']['font_size_pt']} pt")
    print(f"  - H1/H2/H3/H4 字号: {profile['title']['font_size_pt']}/{profile['headings']['level1_pt']}/"
          f"{profile['headings']['level2_pt']}/{profile['headings']['level3_pt']} pt")
    print(f"  - 章节大纲(原始): {len(derived['outline_raw'])} 条 → 通用化骨架")
    print(f"  - 题注样本: {len(derived['captions_samples'])} 条")
    print(f"  - 公式编号风格: {derived['formula_number_styles']}")
    print(f"  - 表格数: {derived['table_count']}")


if __name__ == "__main__":
    main()
