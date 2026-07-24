"""Markdown → DOCX converter used by the product export path.

This is the clean-room Python engine behind ``get_markdown_to_docx()``.  The
Node ``docx-cn-engine`` remains preferred for Chinese academic profiles when
Node is available; this module provides a deterministic local fallback that
still produces real ``.docx`` artifacts with headings, paragraphs, tables and
embedded images.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_IMAGE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)$")
_TABLE_SEP = re.compile(r"^\|?(?:\s*:?-+:?\s*\|)+\s*:?-+:?\s*\|?$")
_ORDERED = re.compile(r"^(\d+)[.)]\s+(.*)$")
_UNORDERED = re.compile(r"^[-*+]\s+(.*)$")


def _set_run_font(run, *, east_asia: str = "SimSun", ascii_font: str = "Times New Roman", size_pt: float = 12, bold: bool = False):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)


def _add_paragraph(document: Document, text: str, *, style: str | None = None, bold: bool = False, size_pt: float = 12, first_line_indent: bool = True):
    paragraph = document.add_paragraph(style=style) if style else document.add_paragraph()
    if first_line_indent and style is None:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    _set_run_font(run, bold=bold, size_pt=size_pt)
    return paragraph


def _add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_heading(level=min(level, 3))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    if paragraph.runs:
        run = paragraph.runs[0]
        run.text = text
    else:
        run = paragraph.add_run(text)
    sizes = {1: 16, 2: 14, 3: 12}
    _set_run_font(run, east_asia="SimHei", bold=True, size_pt=sizes.get(level, 12))


def _split_table_row(line: str) -> list[str]:
    body = line.strip().strip("|")
    return [cell.strip() for cell in body.split("|")]


def _is_table_block(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    current = lines[index].strip()
    nxt = lines[index + 1].strip()
    return current.startswith("|") and _TABLE_SEP.match(nxt) is not None


def _add_table(document: Document, header: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    for col, value in enumerate(header):
        cell = table.rows[0].cells[col]
        cell.text = value
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                _set_run_font(run, bold=True, size_pt=10.5)
    for row_index, row in enumerate(rows, start=1):
        for col, value in enumerate(row):
            if col >= len(header):
                continue
            cell = table.rows[row_index].cells[col]
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, size_pt=10.5)


def _resolve_image(path_value: str, workspace: Path | None, source: Path) -> Path | None:
    candidate = Path(path_value)
    search: list[Path] = []
    if candidate.is_absolute():
        search.append(candidate)
    else:
        search.append((source.parent / candidate).resolve())
        if workspace is not None:
            search.append((workspace / candidate).resolve())
            search.append((workspace / path_value.lstrip("./")).resolve())
    for item in search:
        try:
            if item.is_file():
                return item
        except OSError:
            continue
    return None


def _iter_blocks(lines: Iterable[str]):
    buffer: list[str] = []
    for line in lines:
        if not line.strip():
            if buffer:
                yield buffer
                buffer = []
            continue
        buffer.append(line.rstrip())
    if buffer:
        yield buffer


def markdown_to_docx(
    source: str | Path,
    output_path: str | Path,
    style_profile: str | Path | None = None,
    workspace: str | Path | None = None,
    engine: str = "python",
) -> Path:
    """Convert Markdown to a real DOCX artifact.

    Signature matches the historical product converter contract used by
    ``backend/routers/docx_export.py`` and ``docx_tool_loader``.
    """
    source_path = Path(source)
    output = Path(output_path)
    workspace_path = Path(workspace).resolve() if workspace else source_path.parent.resolve()
    text = source_path.read_text(encoding="utf-8")
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    for edge in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, edge, Cm(2.5))

    # style_profile is reserved for profile-aware layout; absence must not fail export.
    _ = style_profile, engine

    lines = text.splitlines()
    index = 0
    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()
        if not line.strip():
            index += 1
            continue

        heading = _HEADING.match(line.strip())
        if heading:
            _add_heading(document, heading.group(2).strip(), len(heading.group(1)))
            index += 1
            continue

        image = _IMAGE.match(line.strip())
        if image:
            image_path = _resolve_image(image.group(2).strip().strip("<>"), workspace_path, source_path)
            if image_path is not None:
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(str(image_path), width=Inches(5.2))
                caption = image.group(1).strip()
                if caption:
                    cap = document.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = cap.add_run(caption)
                    _set_run_font(cap_run, size_pt=10.5, bold=True)
            else:
                _add_paragraph(document, line.strip(), first_line_indent=False)
            index += 1
            continue

        if _is_table_block(lines, index):
            header = _split_table_row(lines[index])
            index += 2
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(_split_table_row(lines[index]))
                index += 1
            _add_table(document, header, rows)
            continue

        ordered = _ORDERED.match(line.strip())
        if ordered:
            paragraph = document.add_paragraph(style="List Number")
            run = paragraph.add_run(ordered.group(2))
            _set_run_font(run, size_pt=12)
            index += 1
            continue

        unordered = _UNORDERED.match(line.strip())
        if unordered:
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(unordered.group(1))
            _set_run_font(run, size_pt=12)
            index += 1
            continue

        # accumulate a plain paragraph until blank/special line
        chunk = [line.strip()]
        index += 1
        while index < len(lines):
            nxt = lines[index].rstrip()
            if not nxt.strip():
                break
            if (
                _HEADING.match(nxt.strip())
                or _IMAGE.match(nxt.strip())
                or _is_table_block(lines, index)
                or _ORDERED.match(nxt.strip())
                or _UNORDERED.match(nxt.strip())
            ):
                break
            chunk.append(nxt.strip())
            index += 1
        _add_paragraph(document, " ".join(chunk), first_line_indent=True)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    if not output.is_file() or output.stat().st_size == 0:
        raise RuntimeError(f"DOCX export produced no artifact at {output}")
    return output


__all__ = ["markdown_to_docx"]
